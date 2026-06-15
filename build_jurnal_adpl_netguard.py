from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


OUT = Path(r"C:\Users\Asus_\OneDrive\Documents\NetGuard-AI\Jurnal_ADPL_NetGuard_AI_dengan_gambar.docx")
IMAGE_DASHBOARD = Path(r"C:\Users\Asus_\Downloads\NetGuard-AI_Dashboard.png")


def clear_document(doc: Document) -> None:
    body = doc._body._element
    for child in list(body):
        if child.tag.endswith("}sectPr"):
            continue
        body.remove(child)


def set_cell_shading(cell, fill="FFFFFF"):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=110, bottom=90, end=110):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "000000")


def force_times_new_roman(doc: Document) -> None:
    for style in doc.styles:
        if style.type in (1, 2):
            try:
                style.font.name = "Times New Roman"
                style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
                style.font.color.rgb = RGBColor(0, 0, 0)
            except Exception:
                pass


def p(doc, text="", *, align=None, bold=False, italic=False, size=10, before=0, after=0, line=1.15, first_line=False):
    para = doc.add_paragraph()
    para.alignment = align if align is not None else WD_ALIGN_PARAGRAPH.JUSTIFY
    fmt = para.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line
    if first_line:
        fmt.first_line_indent = Cm(0.75)
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor(0, 0, 0)
    return para


def add_mixed_paragraph(doc, parts, *, align=WD_ALIGN_PARAGRAPH.JUSTIFY, size=10, before=0, after=0, line=1.15, first_line=False):
    para = doc.add_paragraph()
    para.alignment = align
    fmt = para.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line
    if first_line:
        fmt.first_line_indent = Cm(0.75)
    for text, bold, italic in parts:
        run = para.add_run(text)
        run.bold = bold
        run.italic = italic
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        run.font.size = Pt(size)
        run.font.color.rgb = RGBColor(0, 0, 0)
    return para


def heading(doc, text):
    return p(doc, text, align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=10, before=8, after=6, line=1.15)


def subheading(doc, number, text):
    return add_mixed_paragraph(
        doc,
        [(number + " ", True, False), (text, True, False)],
        align=WD_ALIGN_PARAGRAPH.JUSTIFY,
        size=10,
        before=6,
        after=2,
        line=1.15,
    )


def caption(doc, text):
    return p(doc, text, align=WD_ALIGN_PARAGRAPH.CENTER, size=10, before=2, after=6, line=1.15)


def add_dashboard_image(doc):
    if not IMAGE_DASHBOARD.exists():
        return
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_before = Pt(6)
    para.paragraph_format.space_after = Pt(2)
    run = para.add_run()
    run.add_picture(str(IMAGE_DASHBOARD), width=Cm(15.6))
    caption(doc, "Gambar 1. Tampilan dashboard monitoring NetGuard AI")


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table)
    hdr = table.rows[0].cells
    for i, text in enumerate(headers):
        hdr[i].width = Cm(widths[i])
        set_cell_shading(hdr[i], "FFFFFF")
        set_cell_margins(hdr[i])
        hdr[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        hdr[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = hdr[i].paragraphs[0].add_run(text)
        r.bold = True
        r.font.name = "Times New Roman"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(0, 0, 0)
    for row in rows:
        cells = table.add_row().cells
        for i, text in enumerate(row):
            cells[i].width = Cm(widths[i])
            set_cell_margins(cells[i])
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            para = cells[i].paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER if i == 0 else WD_ALIGN_PARAGRAPH.JUSTIFY
            para.paragraph_format.line_spacing = 1.05
            para.paragraph_format.space_after = Pt(0)
            r = para.add_run(text)
            r.font.name = "Times New Roman"
            r._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
            r.font.size = Pt(9)
            r.font.color.rgb = RGBColor(0, 0, 0)
    p(doc, "", after=0)
    return table


def add_architecture_box(doc):
    rows = [
        ["Sumber Data", "Router MikroTik, log firewall, DHCP lease, trafik SNMP, laporan pengguna"],
        ["Akuisisi", "Collector membaca log dan metrik jaringan secara berkala melalui API/agent ringan"],
        ["Analitik", "Rule engine dan model anomali memberi skor risiko, pola serangan, dan rekomendasi"],
        ["Layanan", "Dashboard, notifikasi WhatsApp/e-mail, laporan SLA, dan tiket penanganan"],
        ["Pengguna", "Admin jaringan, pemilik UMKM, operator kampus/lab, dan teknisi lapangan"],
    ]
    add_table(doc, ["Lapisan", "Komponen Blueprint NetGuard AI"], rows, [3.2, 12.2])
    caption(doc, "Gambar 2. Blueprint arsitektur konseptual NetGuard AI")


def build():
    doc = Document()
    force_times_new_roman(doc)

    sec = doc.sections[0]
    sec.page_width = Cm(21)
    sec.page_height = Cm(29.7)
    sec.top_margin = Cm(2.54)
    sec.bottom_margin = Cm(2.54)
    sec.left_margin = Cm(2.54)
    sec.right_margin = Cm(2.54)

    p(
        doc,
        "Analisis dan Desain Perangkat Lunak NetGuard AI sebagai Sistem Monitoring Keamanan Jaringan Berbasis Kecerdasan Buatan untuk UMKM dan Lingkungan Kampus",
        align=WD_ALIGN_PARAGRAPH.CENTER,
        bold=True,
        size=14,
        after=8,
        line=1.15,
    )
    p(doc, "Nama Mahasiswa1*, Dosen Pembimbing2", align=WD_ALIGN_PARAGRAPH.CENTER, size=10, line=1.2)
    p(doc, "1,2 Program Studi Informatika, UIN Sultan Maulana Hasanuddin Banten", align=WD_ALIGN_PARAGRAPH.CENTER, size=10, line=1.2)
    p(doc, "* E-mail: nama.mahasiswa@uinbanten.ac.id", align=WD_ALIGN_PARAGRAPH.CENTER, size=10, after=8, line=1.2)

    heading(doc, "ABSTRAK")
    p(
        doc,
        "Perkembangan digitalisasi UMKM, laboratorium sekolah, dan lingkungan kampus kecil membuat kebutuhan pemantauan jaringan semakin penting, tetapi tidak semua organisasi memiliki administrator jaringan khusus. Penelitian ini membahas analisis dan desain perangkat lunak NetGuard AI, yaitu sistem monitoring keamanan jaringan yang dirancang untuk membaca log router, trafik, serta status perangkat kemudian menyajikan peringatan dan rekomendasi tindakan secara sederhana. Metode yang digunakan adalah pendekatan Analisis dan Desain Perangkat Lunak (ADPL) melalui identifikasi masalah, analisis kebutuhan, pemodelan proses bisnis, perancangan use case, perancangan arsitektur, dan rancangan basis data. Hasil rancangan menunjukkan bahwa NetGuard AI dapat diposisikan sebagai produk digital yang relevan untuk pasar UMKM dan institusi pendidikan karena menawarkan visibilitas jaringan, klasifikasi risiko, laporan berkala, serta alur penanganan insiden yang mudah dipahami. Desain ini menekankan prinsip keamanan, privasi data, skalabilitas, dan keberlanjutan produk agar solusi tidak hanya layak secara teknis, tetapi juga memiliki arah pengembangan bisnis yang etis dan visioner.",
        first_line=True,
    )
    add_mixed_paragraph(doc, [("Kata Kunci : ", True, False), ("ADPL, desain perangkat lunak, keamanan jaringan, NetGuard AI, UMKM", False, False)], after=8)

    heading(doc, "ABSTRACT")
    p(
        doc,
        "The digital transformation of small businesses, school laboratories, and campus environments increases the need for network monitoring, yet many organizations do not have dedicated network administrators. This study discusses the software analysis and design of NetGuard AI, a network security monitoring system designed to read router logs, traffic patterns, and device status, then present alerts and practical recommendations. The method follows a Software Analysis and Design approach by identifying problems, analyzing requirements, modeling business processes, designing use cases, defining system architecture, and preparing a database design. The proposed design indicates that NetGuard AI can be positioned as a relevant digital product for small businesses and educational institutions by providing network visibility, risk classification, periodic reporting, and incident-handling workflows. The design emphasizes security, data privacy, scalability, and product longevity so the solution is not only technically feasible but also ethically market-oriented and future-ready.",
        first_line=True,
    )
    add_mixed_paragraph(doc, [("Keywords : ", True, False), ("software analysis and design, network security, NetGuard AI, SMEs, monitoring system", False, False)], after=8)

    heading(doc, "PENDAHULUAN")
    for text in [
        "Jaringan komputer menjadi tulang punggung aktivitas digital pada UMKM, laboratorium sekolah, organisasi kampus, dan layanan publik kecil. Koneksi internet dipakai untuk transaksi, penyimpanan data, sistem kasir, pembelajaran daring, administrasi, dan komunikasi pelanggan. Namun, peningkatan ketergantungan tersebut sering tidak diimbangi dengan kemampuan memantau keamanan jaringan secara terstruktur. Banyak perangkat router sudah menghasilkan log dan metrik yang bermanfaat, tetapi data tersebut jarang dibaca karena tampilannya teknis, tersebar, dan tidak langsung memberi rekomendasi tindakan.",
        "Latar belakang pengguna yang memahami jaringan, seperti pengalaman SMK TKJ dan peminatan Informatika, memberi peluang untuk mengembangkan solusi yang dekat dengan masalah lapangan. Masalah sederhana seperti perangkat asing masuk Wi-Fi, bandwidth tiba-tiba habis, brute force ke router, atau server lokal tidak merespons dapat berdampak besar bagi usaha kecil. Di sisi lain, pasar membutuhkan produk yang tidak sekadar canggih, tetapi mudah digunakan, terjangkau, dan mampu mengubah data teknis menjadi keputusan praktis.",
        "NetGuard AI dirancang sebagai gagasan perangkat lunak yang menggabungkan monitoring jaringan, klasifikasi risiko, dan rekomendasi berbasis aturan serta kecerdasan buatan. Tujuannya bukan menggantikan administrator jaringan, melainkan membantu pemilik usaha atau operator non-ahli memahami kondisi jaringan dan mengambil tindakan awal. Dalam konteks ADPL, penelitian ini berfokus pada bagaimana kebutuhan pengguna diterjemahkan menjadi rancangan sistem, modul, basis data, dan alur interaksi yang dapat dikembangkan menjadi produk.",
        "Rumusan masalah dalam artikel ini adalah bagaimana menganalisis kebutuhan dan merancang perangkat lunak NetGuard AI agar mampu membantu pemantauan keamanan jaringan pada UMKM dan lingkungan kampus kecil. Batasan penelitian berada pada tahap analisis dan desain, belum sampai implementasi penuh. Kontribusi artikel ini adalah rancangan blueprint awal yang dapat menjadi dasar pengembangan prototype, validasi pasar, dan penyusunan model bisnis berkelanjutan.",
    ]:
        p(doc, text, first_line=True)

    heading(doc, "METODE")
    subheading(doc, "2.1", "Pendekatan Analisis dan Desain Perangkat Lunak")
    for text in [
        "Penelitian ini menggunakan pendekatan ADPL yang menempatkan kebutuhan pengguna sebagai dasar rancangan. Tahapan yang digunakan meliputi observasi masalah, identifikasi stakeholder, analisis kebutuhan fungsional dan nonfungsional, pemodelan proses, perancangan use case, arsitektur sistem, rancangan basis data, serta evaluasi kelayakan awal. Pendekatan ini sejalan dengan prinsip rekayasa kebutuhan yang menuntut kebutuhan sistem dinyatakan jelas, dapat ditelusuri, dan dapat diuji.",
        "Objek rancangan adalah NetGuard AI, yaitu sistem berbasis web yang membaca data jaringan dari router atau agent ringan, menyimpan log, menghitung indikator risiko, dan menyajikan rekomendasi kepada pengguna. Fokus desain diarahkan pada organisasi dengan sumber daya terbatas, sehingga sistem harus sederhana, hemat biaya, mudah dipasang, dan tidak memerlukan konfigurasi rumit.",
    ]:
        p(doc, text, first_line=True)

    subheading(doc, "2.2", "Teknik Pengumpulan dan Analisis Kebutuhan")
    p(doc, "Kebutuhan sistem disusun melalui analisis skenario lapangan dari sudut pandang admin jaringan, pemilik UMKM, teknisi freelance, dan operator kampus/lab. Analisis juga mempertimbangkan praktik keamanan aplikasi, manajemen risiko, serta kebutuhan produk digital yang dapat dipasarkan secara berlangganan. Kebutuhan kemudian diklasifikasikan menjadi kebutuhan fungsional, kebutuhan nonfungsional, dan batasan operasional.", first_line=True)
    add_table(
        doc,
        ["Kode", "Kebutuhan Utama", "Prioritas"],
        [
            ["F-01", "Sistem dapat menerima log router, status perangkat, dan metrik trafik jaringan.", "Tinggi"],
            ["F-02", "Sistem menampilkan dashboard kondisi jaringan, perangkat aktif, dan anomali.", "Tinggi"],
            ["F-03", "Sistem memberi peringatan saat ada perangkat asing, trafik tidak wajar, atau login gagal berulang.", "Tinggi"],
            ["F-04", "Sistem menyediakan rekomendasi tindakan awal yang mudah dipahami pengguna non-ahli.", "Tinggi"],
            ["F-05", "Sistem menghasilkan laporan mingguan untuk evaluasi keamanan dan kualitas layanan.", "Sedang"],
        ],
        [2.0, 10.8, 2.6],
    )
    caption(doc, "Tabel 1. Kebutuhan fungsional NetGuard AI")

    subheading(doc, "2.3", "Kriteria Desain")
    for text in [
        "Kriteria desain ditetapkan agar rancangan tidak hanya memenuhi fungsi, tetapi juga memiliki nilai produk. Kriteria tersebut meliputi kemudahan instalasi, keamanan data, privasi, skalabilitas, integrasi dengan perangkat jaringan umum, dan kemampuan sistem untuk menghasilkan insight yang actionable. Dengan demikian, desain tidak berhenti pada tampilan monitoring, melainkan diarahkan menjadi layanan bernilai pasar.",
        "Dalam aspek keamanan, rancangan mengikuti prinsip minimum privilege, pencatatan audit, validasi input, enkripsi komunikasi, dan pemisahan akses pengguna. Dalam aspek bisnis, rancangan mempertimbangkan model private market atau SaaS ringan dengan paket untuk UMKM, sekolah, dan teknisi jaringan independen.",
    ]:
        p(doc, text, first_line=True)

    heading(doc, "HASIL DAN PEMBAHASAN")
    subheading(doc, "3.1", "Gambaran Umum Sistem")
    for text in [
        "NetGuard AI dirancang sebagai platform monitoring keamanan jaringan yang mengubah data teknis menjadi informasi yang mudah dipahami. Pengguna utama adalah pemilik UMKM, admin lab, pengelola jaringan kampus kecil, dan teknisi jaringan. Sistem bekerja dengan mengumpulkan data dari router atau agent, menyimpan data ke server, melakukan analisis risiko, lalu menampilkan hasil pada dashboard.",
        "Nilai utama sistem adalah membantu pengguna mengetahui apa yang sedang terjadi pada jaringan, mengapa kondisi tersebut berisiko, dan tindakan apa yang sebaiknya dilakukan. Pendekatan ini berbeda dari monitoring tradisional yang sering hanya menampilkan angka dan grafik tanpa konteks keputusan.",
    ]:
        p(doc, text, first_line=True)
    add_dashboard_image(doc)
    add_architecture_box(doc)

    subheading(doc, "3.2", "Aktor dan Use Case")
    p(doc, "Aktor pada sistem terdiri dari Admin Sistem, Operator Jaringan, Pemilik Usaha, dan Teknisi Lapangan. Admin Sistem mengelola tenant, paket layanan, dan konfigurasi global. Operator Jaringan memantau kondisi perangkat dan menindaklanjuti alert. Pemilik Usaha membaca ringkasan risiko dan laporan. Teknisi Lapangan menerima tiket serta catatan troubleshooting.", first_line=True)
    add_table(
        doc,
        ["Aktor", "Use Case", "Deskripsi"],
        [
            ["Admin Sistem", "Kelola tenant", "Mendaftarkan organisasi, paket layanan, pengguna, dan batas penyimpanan log."],
            ["Operator Jaringan", "Pantau dashboard", "Melihat status jaringan, perangkat aktif, kualitas koneksi, dan anomali."],
            ["Operator Jaringan", "Validasi alert", "Menentukan apakah peringatan merupakan insiden, false positive, atau perlu eskalasi."],
            ["Pemilik Usaha", "Lihat laporan", "Membaca ringkasan risiko, uptime, perangkat bermasalah, dan rekomendasi investasi."],
            ["Teknisi Lapangan", "Tangani tiket", "Menerima detail masalah dan mencatat tindakan perbaikan."],
        ],
        [3.1, 4.0, 8.3],
    )
    caption(doc, "Tabel 2. Rancangan aktor dan use case")

    subheading(doc, "3.3", "Rancangan Proses Bisnis")
    for text in [
        "Proses bisnis dimulai ketika organisasi mendaftarkan perangkat jaringan. Setelah konfigurasi awal selesai, collector membaca log dan metrik secara periodik. Data mentah disaring untuk mengurangi noise, kemudian dianalisis oleh rule engine dan modul anomali. Jika ditemukan kondisi mencurigakan, sistem membuat alert dengan tingkat risiko rendah, sedang, atau tinggi. Pengguna menerima notifikasi dan dapat membuka tiket penanganan.",
        "Alur ini menjaga agar pengguna tidak tenggelam dalam detail teknis. Informasi teknis tetap tersimpan untuk audit, tetapi tampilan utama menonjolkan prioritas tindakan. Dengan cara ini, sistem dapat mendukung pengambilan keputusan yang cepat pada organisasi yang tidak memiliki tim keamanan khusus.",
    ]:
        p(doc, text, first_line=True)
    add_table(
        doc,
        ["Tahap", "Input", "Output"],
        [
            ["Registrasi", "Data organisasi, perangkat, kredensial API terbatas", "Tenant dan perangkat aktif"],
            ["Akuisisi", "Log firewall, DHCP, trafik, status interface", "Data mentah tervalidasi"],
            ["Analisis", "Data mentah dan baseline trafik", "Skor risiko dan klasifikasi kejadian"],
            ["Respons", "Alert dan rekomendasi", "Tiket, notifikasi, dan catatan tindakan"],
            ["Evaluasi", "Riwayat insiden dan uptime", "Laporan mingguan/bulanan"],
        ],
        [2.8, 6.5, 6.1],
    )
    caption(doc, "Tabel 3. Alur proses bisnis NetGuard AI")

    subheading(doc, "3.4", "Rancangan Arsitektur Sistem")
    for text in [
        "Arsitektur sistem menggunakan pola modular agar mudah dikembangkan. Modul collector bertugas mengambil data dari router atau agent. Modul ingestion melakukan validasi, normalisasi, dan penyimpanan. Modul analytics menjalankan rule engine dan model anomali. Modul notification mengirim peringatan. Modul dashboard menampilkan ringkasan kondisi jaringan. Pemisahan modul memudahkan sistem ditingkatkan dari prototype lokal menjadi layanan cloud.",
        "Untuk tahap awal, sistem dapat berjalan pada server kecil atau virtual private server. Pada tahap lanjut, ingestion dan analytics dapat dipisahkan menggunakan message queue agar mampu menerima data dari banyak tenant. Strategi ini menjaga longevity produk karena sistem bisa dimulai sederhana, tetapi tetap memiliki jalur pertumbuhan ketika jumlah pengguna meningkat.",
    ]:
        p(doc, text, first_line=True)

    subheading(doc, "3.5", "Rancangan Data")
    p(doc, "Rancangan basis data memuat entitas utama yaitu User, Tenant, Device, NetworkLog, Alert, Recommendation, Ticket, dan Report. Relasi antarentitas dirancang untuk mendukung multi-tenant sehingga satu platform dapat melayani banyak organisasi tanpa mencampur data. Setiap log memiliki relasi ke perangkat, setiap alert memiliki relasi ke log pemicu, dan setiap tiket memiliki status penanganan.", first_line=True)
    add_table(
        doc,
        ["Entitas", "Atribut Utama", "Fungsi"],
        [
            ["Tenant", "tenant_id, nama, paket, status", "Menyimpan identitas organisasi pengguna layanan."],
            ["Device", "device_id, tenant_id, tipe, ip_address, lokasi", "Mewakili router, access point, server, atau perangkat jaringan."],
            ["NetworkLog", "log_id, device_id, timestamp, event_type, payload", "Menyimpan log jaringan yang sudah dinormalisasi."],
            ["Alert", "alert_id, severity, status, score, created_at", "Mencatat kejadian berisiko dan tingkat prioritasnya."],
            ["Ticket", "ticket_id, alert_id, assignee, action_note, status", "Mengatur tindak lanjut insiden hingga selesai."],
        ],
        [3.0, 5.9, 6.5],
    )
    caption(doc, "Tabel 4. Rancangan entitas data utama")

    subheading(doc, "3.6", "Kebutuhan Nonfungsional")
    p(doc, "Kebutuhan nonfungsional menentukan kualitas sistem. NetGuard AI harus aman, stabil, mudah digunakan, dan dapat dipercaya. Sistem perlu menyimpan kredensial perangkat secara terenkripsi, membatasi hak akses berdasarkan peran, menyediakan audit trail, serta menjaga data tenant tetap terisolasi. Dari sisi performa, dashboard harus menampilkan ringkasan terbaru tanpa menunggu pemrosesan log yang berat.", first_line=True)
    add_table(
        doc,
        ["Aspek", "Rancangan Kualitas"],
        [
            ["Keamanan", "TLS, enkripsi token, role-based access control, audit log, validasi input, dan pembatasan API."],
            ["Privasi", "Isolasi data tenant, retensi log terukur, masking data sensitif, dan persetujuan integrasi perangkat."],
            ["Usability", "Dashboard ringkas, bahasa rekomendasi non-teknis, indikator risiko sederhana, dan laporan otomatis."],
            ["Reliabilitas", "Retry pada collector, health check, backup berkala, dan pemantauan status service."],
            ["Skalabilitas", "Arsitektur modular, pemisahan ingestion-analytics, dan dukungan multi-tenant."],
        ],
        [3.2, 12.2],
    )
    caption(doc, "Tabel 5. Kebutuhan nonfungsional")

    subheading(doc, "3.7", "Analisis Kelayakan Produk dan Pasar")
    for text in [
        "Dari sisi pasar, NetGuard AI memiliki peluang karena banyak UMKM dan institusi kecil membutuhkan keamanan jaringan tetapi tidak memiliki anggaran untuk solusi enterprise. Produk dapat diposisikan sebagai layanan monitoring ringan dengan paket bertingkat: paket starter untuk satu lokasi, paket growth untuk beberapa cabang, dan paket technician untuk penyedia jasa jaringan. Model ini memungkinkan kapitalisasi pasar tanpa merugikan pengguna karena nilai produk berasal dari transparansi, pencegahan gangguan, dan efisiensi operasional.",
        "Blueprint jangka panjang NetGuard AI dapat berkembang dari monitoring jaringan lokal menjadi platform observability untuk aset digital kecil. Perkembangan tersebut dapat mencakup integrasi perangkat IoT, prediksi kebutuhan bandwidth, rekomendasi topologi, marketplace teknisi, dan laporan kepatuhan keamanan. Dengan arah tersebut, ide sederhana dari pengalaman jaringan dapat naik kelas menjadi produk yang memiliki proyeksi berkelanjutan.",
    ]:
        p(doc, text, first_line=True)

    subheading(doc, "3.8", "Risiko dan Mitigasi")
    add_table(
        doc,
        ["Risiko", "Dampak", "Mitigasi Desain"],
        [
            ["False positive alert", "Pengguna lelah menerima notifikasi", "Gunakan threshold adaptif, feedback pengguna, dan kategori prioritas."],
            ["Kebocoran kredensial perangkat", "Akses tidak sah ke jaringan klien", "Enkripsi rahasia, rotasi token, dan prinsip minimum privilege."],
            ["Data log terlalu besar", "Biaya penyimpanan meningkat", "Retensi bertingkat, agregasi metrik, dan arsip otomatis."],
            ["Kesulitan instalasi", "Adopsi pasar rendah", "Wizard konfigurasi, template MikroTik, dan dokumentasi singkat."],
            ["Ketergantungan model AI", "Rekomendasi tidak akurat", "Gabungkan rule engine, validasi manusia, dan evaluasi berkala."],
        ],
        [4.2, 4.5, 6.7],
    )
    caption(doc, "Tabel 6. Risiko dan mitigasi rancangan")

    heading(doc, "KESIMPULAN")
    for text in [
        "Artikel ini menghasilkan rancangan awal NetGuard AI sebagai perangkat lunak monitoring keamanan jaringan berbasis kecerdasan buatan untuk UMKM dan lingkungan kampus kecil. Melalui pendekatan ADPL, kebutuhan pengguna diterjemahkan menjadi rancangan fungsional, use case, arsitektur modular, rancangan data, kebutuhan nonfungsional, serta analisis kelayakan produk. Hasil pembahasan menunjukkan bahwa permasalahan jaringan yang sering dianggap teknis dapat dikemas menjadi layanan digital yang praktis dan bernilai pasar.",
        "NetGuard AI memiliki potensi dikembangkan menjadi prototype karena berangkat dari kebutuhan nyata: keterbatasan sumber daya keamanan jaringan pada organisasi kecil. Pengembangan berikutnya dapat dilakukan melalui pembuatan minimum viable product, integrasi awal dengan router MikroTik, pengujian pada jaringan lab, serta validasi pengalaman pengguna. Dengan menjaga keamanan, privasi, dan manfaat nyata, sistem ini dapat menjadi langkah awal menuju produk visioner yang membantu banyak orang menyelesaikan masalah jaringan tanpa mengeksploitasi pengguna.",
    ]:
        p(doc, text, first_line=True)

    heading(doc, "DAFTAR PUSTAKA")
    refs = [
        "[1] ISO/IEC/IEEE, ISO/IEC/IEEE 29148:2018 Systems and software engineering - Life cycle processes - Requirements engineering, 2018.",
        "[2] National Institute of Standards and Technology, The NIST Cybersecurity Framework (CSF) 2.0, 2024.",
        "[3] OWASP Foundation, Application Security Verification Standard (ASVS), OWASP Project, 2024.",
        "[4] I. Sommerville, Software Engineering, 10th ed. Boston: Pearson, 2016.",
        "[5] R. S. Pressman and B. R. Maxim, Software Engineering: A Practitioner's Approach, 9th ed. New York: McGraw-Hill, 2020.",
        "[6] G. Booch, J. Rumbaugh, and I. Jacobson, The Unified Modeling Language User Guide, 2nd ed. Boston: Addison-Wesley, 2005.",
        "[7] A. S. Tanenbaum and D. J. Wetherall, Computer Networks, 5th ed. Boston: Pearson, 2011.",
        "[8] MikroTik, RouterOS Documentation: Logging, Firewall, and API Concepts, MikroTik Documentation, 2024.",
    ]
    for ref in refs:
        para = p(doc, ref, align=WD_ALIGN_PARAGRAPH.JUSTIFY, size=10, line=1.15)
        para.paragraph_format.first_line_indent = Cm(-0.55)
        para.paragraph_format.left_indent = Cm(0.55)

    for para in doc.paragraphs:
        for run in para.runs:
            run.font.name = "Times New Roman"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
            run.font.color.rgb = RGBColor(0, 0, 0)
    for table in doc.tables:
        set_table_borders(table)
        for row in table.rows:
            for cell in row.cells:
                set_cell_shading(cell, "FFFFFF")
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.name = "Times New Roman"
                        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
                        run.font.color.rgb = RGBColor(0, 0, 0)

    doc.save(OUT)


if __name__ == "__main__":
    build()
    print(OUT)
