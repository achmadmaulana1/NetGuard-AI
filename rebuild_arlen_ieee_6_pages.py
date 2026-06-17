from __future__ import annotations

import json
from pathlib import Path
from zipfile import ZipFile

import matplotlib.pyplot as plt
import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from matplotlib.backends.backend_pdf import PdfPages


ROOT = Path(__file__).resolve().parent
BASE = ROOT / "Arlen_Prima_Dinova_241730003_UAS_AI"
OUT_DIR = BASE / "09_Draft_IEEE"
TITLE = "RetainA: Prediksi Churn Pelanggan Berbasis Ensemble Learning dan Explainable Business Risk Scoring"
AUTHOR = "Arlen Prima Dinova"
AUTHOR_META = "Arlen_Prima_Dinova"
NIM = "241730003"


REFERENCES = [
    "Zhang et al., \"Customer churn prediction model based on hybrid neural networks,\" Scientific Reports, 2024.",
    "Peng and Peng, \"Research on Telecom Customer Churn Prediction Based on GA-XGBoost and SHAP,\" Journal of Computer and Communications, 2022.",
    "ETASR, \"Customer churn prediction for telecommunication companies using optimized classifiers,\" Engineering Technology and Applied Science Research, 2024.",
    "Mahayasa et al., \"Customer churn prediction using weighted average ensemble machine learning model,\" JCSSE, 2023.",
    "Bhushan, \"Enhancing customer churn prediction in the telecom sector using ensemble learning,\" National College of Ireland, 2024.",
    "Sun, \"Sales prediction based on machine learning approach,\" Atlantis Press, 2024.",
    "Obi, \"Demand forecasting in retail business using the ensemble machine learning framework,\" ASRJETS, 2024.",
    "SBC ENIAC, \"Comparing gradient boosting algorithms to forecast sales in retail,\" SBC OpenLib, 2023.",
    "EUDL, \"Sales forecast of retail commodity on the basis of LightGBM and XGBoost,\" EUDL, 2022.",
    "UCI Machine Learning Repository, \"Online shoppers purchasing intention dataset,\" University of California Irvine, 2018.",
    "IBM, \"Telco customer churn dataset repository,\" IBM GitHub Repository, 2024.",
    "OpenML, \"Public machine learning datasets for reproducible tabular AI,\" OpenML, 2024.",
    "L. Breiman, \"Random forests,\" Machine Learning, vol. 45, no. 1, pp. 5-32, 2001.",
    "J. H. Friedman, \"Greedy function approximation: a gradient boosting machine,\" Annals of Statistics, vol. 29, no. 5, pp. 1189-1232, 2001.",
    "F. Pedregosa et al., \"Scikit-learn: Machine Learning in Python,\" Journal of Machine Learning Research, vol. 12, pp. 2825-2830, 2011.",
    "R. Kohavi, \"A study of cross-validation and bootstrap for accuracy estimation and model selection,\" IJCAI, 1995.",
    "S. Lundberg and S. Lee, \"A unified approach to interpreting model predictions,\" Advances in Neural Information Processing Systems, 2017.",
    "IEEE, \"IEEE Editorial Style Manual for Authors,\" IEEE Author Center, 2024.",
    "UCI Machine Learning Repository, \"Bank marketing dataset,\" University of California Irvine, 2014.",
    "Kaggle, \"Business analytics public dataset practices,\" Kaggle Datasets, 2024.",
    "A. Géron, Hands-On Machine Learning with Scikit-Learn, Keras, and TensorFlow, 3rd ed. Sebastopol, CA, USA: O'Reilly Media, 2022.",
    "C. Molnar, Interpretable Machine Learning, 2nd ed., 2022.",
]


def set_columns(section, count=2) -> None:
    sect_pr = section._sectPr
    cols = sect_pr.xpath("./w:cols")
    cols = cols[0] if cols else OxmlElement("w:cols")
    if cols.getparent() is None:
        sect_pr.append(cols)
    cols.set(qn("w:num"), str(count))
    cols.set(qn("w:space"), "360")


def style_doc(doc: Document) -> None:
    cp = doc.core_properties
    cp.author = AUTHOR_META
    cp.last_modified_by = AUTHOR_META
    cp.title = TITLE
    cp.subject = ""
    cp.keywords = ""
    cp.comments = ""
    s = doc.sections[0]
    s.top_margin = Inches(0.68)
    s.bottom_margin = Inches(0.68)
    s.left_margin = Inches(0.66)
    s.right_margin = Inches(0.66)
    s.page_width = Inches(8.5)
    s.page_height = Inches(11)
    for name, size, bold in [
        ("Normal", 8.7, False),
        ("Heading 1", 9.5, True),
        ("Heading 2", 8.7, True),
        ("Heading 3", 8.7, True),
    ]:
        st = doc.styles[name]
        st.font.name = "Times New Roman"
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        st.font.size = Pt(size)
        st.font.bold = bold
        st.font.color.rgb = RGBColor(0, 0, 0)
        st.paragraph_format.space_after = Pt(2)
        st.paragraph_format.space_before = Pt(0)


def run(paragraph, text: str, size=8.7, bold=False) -> None:
    r = paragraph.add_run(text)
    r.font.name = "Times New Roman"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    r.font.size = Pt(size)
    r.font.bold = bold
    r.font.color.rgb = RGBColor(0, 0, 0)


def p(doc: Document, text: str, size=8.7, align=None, bold=False) -> None:
    para = doc.add_paragraph()
    if align is not None:
        para.alignment = align
    para.paragraph_format.line_spacing = 1.0
    para.paragraph_format.space_after = Pt(2)
    run(para, text, size=size, bold=bold)


def h(doc: Document, text: str) -> None:
    para = doc.add_paragraph(style="Heading 1")
    para.paragraph_format.keep_with_next = True
    run(para, text, size=9.5, bold=True)


def table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    for i, head in enumerate(headers):
        t.rows[0].cells[i].text = head
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
    for row in t.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                para.paragraph_format.space_after = Pt(0)
                for rr in para.runs:
                    rr.font.name = "Times New Roman"
                    rr._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
                    rr.font.size = Pt(6.7)
                    rr.font.color.rgb = RGBColor(0, 0, 0)


def img(doc: Document, path: Path, caption: str, width=2.85) -> None:
    if path.exists():
        doc.add_picture(str(path), width=Inches(width))
        p(doc, caption, size=7.2, align=WD_ALIGN_PARAGRAPH.CENTER)


def result_rows(comp: pd.DataFrame) -> list[list[str]]:
    return [
        [
            r["model"],
            f"{r['accuracy']:.3f}",
            f"{r['precision']:.3f}",
            f"{r['recall']:.3f}",
            f"{r['f1']:.3f}",
            f"{r['auc']:.3f}",
        ]
        for _, r in comp.iterrows()
    ]


def add_page_break(doc: Document) -> None:
    doc.add_page_break()


def build_docx() -> None:
    metrics = json.loads((BASE / "07_Hasil_Eksperimen" / "metrics.json").read_text(encoding="utf-8"))
    comp = pd.read_csv(BASE / "07_Hasil_Eksperimen" / "model_comparison.csv")

    doc = Document()
    style_doc(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run(title, TITLE, size=16, bold=True)
    p(
        doc,
        f"{AUTHOR}, {NIM}\nProgram Studi Informatika, Fakultas Sains dan Teknologi\nUniversitas Islam Negeri Sultan Maulana Hasanuddin Banten\nBanten, Indonesia, 2026",
        size=9.5,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    section = doc.add_section(WD_SECTION.CONTINUOUS)
    set_columns(section)

    h(doc, "Abstrak")
    p(doc, f"Penelitian ini mengembangkan RetainA, yaitu sistem prediksi churn pelanggan berbasis ensemble learning untuk membantu prioritas retensi pada bisnis berlangganan. Masalah utama yang diangkat adalah kebutuhan perusahaan untuk mengenali pelanggan berisiko berhenti sebelum keputusan retensi terlambat atau diberikan secara terlalu luas. Dataset utama yang digunakan adalah IBM Telco Customer Churn dengan {metrics['dataset_rows']} baris dan {metrics['dataset_columns']} kolom. Metode penelitian meliputi preprocessing tabular, baseline Logistic Regression, model pembanding Random Forest dan Gradient Boosting, hyperparameter tuning, serta soft voting ensemble. Model akhir memperoleh accuracy {metrics['metrics']['accuracy']:.3f}, precision {metrics['metrics']['precision']:.3f}, recall {metrics['metrics']['recall']:.3f}, F1-score {metrics['metrics']['f1']:.3f}, dan AUC {metrics['metrics']['auc']:.3f}. Hasil menunjukkan bahwa model dapat digunakan sebagai sistem pendukung keputusan retensi, terutama untuk membaca keseimbangan antara cakupan pelanggan berisiko dan efisiensi biaya intervensi.")
    h(doc, "Kata Kunci")
    p(doc, "Artificial Intelligence; churn pelanggan; ensemble learning; risk scoring; baseline comparison; hyperparameter tuning; deployment lokal.")
    h(doc, "I. Pendahuluan")
    for text in [
        "Churn pelanggan merupakan salah satu masalah penting dalam bisnis berlangganan karena kehilangan pelanggan aktif dapat meningkatkan biaya akuisisi ulang dan menurunkan pendapatan berulang. Pada layanan telekomunikasi, pelanggan dapat berhenti karena faktor kontrak, biaya bulanan, kualitas layanan, jenis layanan internet, dukungan teknis, serta pengalaman pembayaran. Jika risiko churn baru diketahui setelah pelanggan benar-benar berhenti, strategi retensi menjadi reaktif dan kurang efisien.",
        "RetainA dirancang untuk membaca pola risiko churn dari data pelanggan dan mengubahnya menjadi skor prioritas. Skor ini tidak dipakai sebagai keputusan final otomatis, tetapi sebagai sinyal awal agar tim bisnis dapat menentukan pelanggan mana yang perlu dipantau, diberi penawaran, atau dianalisis lebih lanjut. Dengan demikian, fokus sistem bukan hanya menghasilkan label prediksi, melainkan mendukung proses pengambilan keputusan retensi.",
        "Pendekatan berbasis machine learning relevan karena data pelanggan berbentuk tabular, memiliki kombinasi fitur numerik dan kategorikal, serta menyimpan pola hubungan yang tidak selalu linear. Model linear dapat menjadi baseline yang baik, tetapi interaksi antara kontrak, biaya, layanan tambahan, dan durasi berlangganan sering membutuhkan model yang lebih fleksibel. Oleh karena itu, penelitian ini membandingkan baseline dengan model ensemble dan menggabungkannya melalui soft voting.",
        "Kontribusi utama RetainA adalah penyatuan pipeline eksperimen, evaluasi multi-metrik, interpretasi bisnis, dan prototipe web lokal. Evaluasi tidak hanya menggunakan accuracy, tetapi juga precision, recall, F1-score, AUC, confusion matrix, dan feature importance. Metrik-metrik tersebut penting karena keputusan retensi memiliki biaya yang berbeda untuk false positive dan false negative.",
    ]:
        p(doc, text)

    add_page_break(doc)
    h(doc, "II. Penelitian Terkait")
    for text in [
        "Penelitian churn pelanggan dalam beberapa tahun terakhir banyak memanfaatkan model ensemble, optimized classifiers, XGBoost, explainable AI, dan weighted ensemble. Studi hybrid neural network menunjukkan bahwa model non-linear dapat meningkatkan kemampuan membaca pola churn [1]. Penelitian GA-XGBoost dan SHAP menekankan pentingnya interpretasi fitur untuk memahami alasan prediksi [2]. Studi optimized classifiers dan weighted ensemble juga menunjukkan bahwa penggabungan beberapa model dapat memberikan performa yang lebih stabil dibanding satu model tunggal [3], [4].",
        "Selain domain churn, penelitian pada prediksi penjualan dan perilaku bisnis juga menunjukkan dominasi model tabular seperti Random Forest, Gradient Boosting, LightGBM, dan XGBoost [6]-[9]. Temuan tersebut memperkuat asumsi bahwa data bisnis tabular dapat diproses secara efektif dengan model ensemble. Namun, sebagian penelitian masih berfokus pada skor akhir tanpa selalu menyediakan alur deployment sederhana yang dapat diuji oleh pengguna.",
        "Berdasarkan literature mapping, gap yang paling relevan untuk RetainA adalah kurangnya integrasi antara eksperimen model, baseline yang eksplisit, tuning, interpretasi risiko bisnis, dan web inferensi lokal. Gap ini penting karena sistem churn tidak hanya dinilai dari akurasi, tetapi juga dari kemampuannya membantu prioritas tindakan. Tanpa interpretasi bisnis, angka evaluasi sulit diterjemahkan menjadi keputusan retensi.",
        "Penelitian ini mengambil posisi sebagai implementasi AI tabular yang lengkap untuk kasus churn. Metode yang digunakan bukan algoritma baru dari nol, melainkan integrasi terukur antara preprocessing, beberapa model pembanding, soft voting ensemble, dan dashboard lokal. Kebaruan tersebut bersifat aplikatif dan dapat diverifikasi melalui dataset, kode, model, hasil eksperimen, visualisasi, dan deployment evidence.",
    ]:
        p(doc, text)
    h(doc, "III. Research Gap dan Novelty")
    table(doc, ["Existing Research", "Limitation", "Proposed Solution"], [
        ["Model churn akurasi tinggi", "Deployment dan reproducibility terbatas", "Pipeline RetainA dengan kode, model, dan web lokal"],
        ["Baseline tunggal", "Pembanding kurang kuat", "Logistic Regression, Random Forest, Gradient Boosting, dan ensemble"],
        ["Evaluasi angka", "Interpretasi bisnis terbatas", "Risk scoring untuk prioritas retensi"],
        ["Tuning tidak selalu jelas", "Sulit direproduksi", "RandomizedSearchCV dan cross-validation terdokumentasi"],
    ])
    p(doc, "Novelty RetainA terletak pada integrasi soft voting ensemble dengan interpretasi risk scoring pelanggan. Pendekatan ini menghasilkan probabilitas churn yang dapat dipakai untuk membagi pelanggan ke dalam prioritas tindakan. Dengan model ini, perusahaan dapat menyusun strategi retensi yang lebih terarah daripada memberikan penawaran yang sama kepada semua pelanggan.")

    add_page_break(doc)
    h(doc, "IV. Metode yang Diusulkan")
    for text in [
        "Metode RetainA dimulai dari pengambilan dataset IBM Telco Customer Churn. Dataset ini berisi informasi demografis, jenis layanan, status kontrak, metode pembayaran, biaya bulanan, total biaya, dan label churn. Kolom target dikonversi menjadi kelas biner, sedangkan kolom identitas pelanggan yang tidak memiliki nilai prediktif langsung tidak digunakan sebagai fitur model.",
        "Preprocessing dilakukan menggunakan pipeline agar transformasi data latih dan data uji konsisten. Fitur numerik diproses dengan median imputation dan standardization. Fitur kategorikal diproses dengan most-frequent imputation dan one-hot encoding. Pendekatan ini mengurangi risiko data leakage karena transformasi dipelajari dari data latih dan diterapkan ke data uji melalui pipeline yang sama.",
        "Baseline yang digunakan adalah Logistic Regression. Model ini dipilih karena sederhana, cepat, dan mudah dijadikan titik pembanding. Random Forest digunakan sebagai pembanding karena mampu menangkap interaksi non-linear antar fitur. Gradient Boosting digunakan karena pendekatan boosting sering kuat pada data tabular bisnis. Model usulan adalah soft voting ensemble yang menggabungkan probabilitas dari beberapa classifier.",
        "Soft voting dipilih karena output probabilitas lebih berguna untuk risk scoring. Pada kasus churn, perusahaan tidak hanya membutuhkan label churn atau tidak churn, tetapi juga tingkat risiko. Probabilitas dapat dikonversi menjadi prioritas, misalnya risiko tinggi untuk pelanggan yang perlu follow-up segera, risiko menengah untuk monitoring, dan risiko rendah untuk perjalanan pelanggan normal.",
        "Hyperparameter tuning dilakukan menggunakan RandomizedSearchCV dengan stratified cross-validation. Ruang pencarian dibuat moderat agar proses pelatihan tetap realistis pada perangkat lokal, tetapi tetap menunjukkan pencarian parameter yang sistematis. Parameter yang diuji mencakup jumlah estimator, kedalaman pohon, learning rate, dan minimum sample leaf pada model berbasis tree.",
    ]:
        p(doc, text)
    img(doc, BASE / "08_Visualisasi" / "pipeline_diagram.png", "Gambar 1. Pipeline metode RetainA.", width=2.75)
    h(doc, "V. Dataset dan Setup Eksperimen")
    table(doc, ["Komponen", "Nilai"], [
        ["Dataset", "IBM Telco Customer Churn"],
        ["Jumlah data", str(metrics["dataset_rows"])],
        ["Jumlah kolom", str(metrics["dataset_columns"])],
        ["Target", "Churn"],
        ["Baseline", "Logistic Regression"],
        ["Pembanding", "Random Forest dan Gradient Boosting"],
        ["Metode usulan", "Tuned Soft Voting Ensemble"],
    ])
    p(doc, "Evaluasi menggunakan accuracy, precision, recall, F1-score, AUC, ROC curve, dan confusion matrix. Accuracy memberi gambaran umum, precision mengukur efisiensi intervensi, recall mengukur cakupan pelanggan berisiko, dan F1-score mengukur keseimbangan antara precision dan recall.")

    add_page_break(doc)
    h(doc, "VI. Hasil Eksperimen")
    p(doc, "Hasil perbandingan model ditunjukkan pada Tabel III. Baseline Logistic Regression digunakan sebagai pembanding awal, sedangkan Random Forest dan Gradient Boosting digunakan untuk melihat performa model non-linear. Model usulan menggunakan soft voting ensemble untuk menggabungkan kekuatan beberapa model.")
    table(doc, ["Model", "Accuracy", "Precision", "Recall", "F1", "AUC"], result_rows(comp))
    p(doc, f"Model akhir memperoleh accuracy {metrics['metrics']['accuracy']:.3f}, precision {metrics['metrics']['precision']:.3f}, recall {metrics['metrics']['recall']:.3f}, F1-score {metrics['metrics']['f1']:.3f}, dan AUC {metrics['metrics']['auc']:.3f}. Nilai recall menunjukkan kemampuan model menangkap pelanggan yang benar-benar berisiko churn, sedangkan precision menunjukkan seberapa efisien prediksi positif untuk diarahkan ke intervensi retensi.")
    p(doc, "Pada konteks bisnis, false negative lebih berbahaya karena pelanggan berisiko dapat terlewat dan akhirnya berhenti. Namun false positive juga perlu dikontrol karena setiap penawaran retensi memiliki biaya. Oleh karena itu, F1-score menjadi metrik penting karena menyeimbangkan precision dan recall. AUC digunakan untuk membaca kemampuan model membedakan pelanggan berisiko pada berbagai threshold.")
    img(doc, BASE / "08_Visualisasi" / "model_comparison_chart.png", "Gambar 2. Perbandingan F1-score antar model.", width=2.7)
    img(doc, BASE / "08_Visualisasi" / "confusion_matrix.png", "Gambar 3. Confusion matrix model usulan.", width=2.55)
    p(doc, "Confusion matrix membantu memahami jenis kesalahan model. Jika model menghasilkan false positive, perusahaan mungkin memberikan penawaran kepada pelanggan yang sebenarnya tidak akan churn. Jika model menghasilkan false negative, perusahaan kehilangan kesempatan mempertahankan pelanggan berisiko. Interpretasi ini lebih bermakna dibanding hanya melihat accuracy.")

    add_page_break(doc)
    h(doc, "VII. Analisis Hasil dan Interpretasi Bisnis")
    for text in [
        "Grafik ROC menunjukkan hubungan antara true positive rate dan false positive rate. AUC digunakan untuk melihat kualitas ranking probabilitas churn. Jika perusahaan ingin menangkap lebih banyak pelanggan berisiko, threshold dapat diturunkan. Jika perusahaan ingin membatasi biaya promosi, threshold dapat dinaikkan. Fleksibilitas threshold ini membuat output probabilitas lebih berguna dibanding label biner saja.",
        "Feature importance memperlihatkan fitur yang paling berpengaruh terhadap prediksi. Dalam konteks telco churn, fitur seperti kontrak, biaya bulanan, total biaya, layanan internet, dan tenure dapat menjadi indikator penting. Informasi ini dapat membantu tim bisnis memahami faktor yang berhubungan dengan risiko churn dan merancang strategi layanan yang lebih tepat.",
        "RetainA juga menyediakan dashboard web lokal. Dashboard tersebut menampilkan metrik eksperimen dan menyediakan form prediksi. Pengguna dapat mengubah nilai fitur dari contoh data untuk melihat perubahan probabilitas churn. Dengan demikian, model tidak hanya tersimpan sebagai file `.pkl`, tetapi dapat diuji melalui antarmuka yang lebih mudah dipahami.",
    ]:
        p(doc, text)
    img(doc, BASE / "08_Visualisasi" / "roc_curve.png", "Gambar 4. ROC curve model usulan.", width=2.55)
    img(doc, BASE / "08_Visualisasi" / "feature_importance.png", "Gambar 5. Feature importance.", width=2.7)
    h(doc, "VIII. Deployment dan Ancaman Validitas")
    p(doc, "Deployment lokal menggunakan Flask. Aplikasi membaca model hasil training dan menghasilkan probabilitas churn berdasarkan input pengguna. Arsitektur ini memisahkan proses training dan inference sehingga dashboard tidak melatih ulang model setiap kali digunakan. Untuk penggunaan produksi, sistem masih memerlukan autentikasi, logging, database, monitoring, HTTPS, dan pengujian keamanan.")
    p(doc, "Ancaman validitas penelitian meliputi representativitas dataset publik, kemungkinan perbedaan perilaku pelanggan lokal, class imbalance, temporal drift, dan keterbatasan interpretabilitas. Dataset publik tidak selalu mewakili kondisi bisnis di Indonesia. Oleh karena itu, RetainA diposisikan sebagai prototipe penelitian yang perlu divalidasi ulang sebelum dipakai pada data operasional nyata.")
    img(doc, BASE / "08_Visualisasi" / "deployment_architecture.png", "Gambar 6. Arsitektur deployment lokal.", width=2.7)

    add_page_break(doc)
    h(doc, "IX. Kesimpulan")
    for text in [
        "RetainA berhasil membangun pipeline prediksi churn pelanggan berbasis ensemble learning. Penelitian ini menggunakan dataset publik, preprocessing terstruktur, baseline Logistic Regression, model pembanding Random Forest dan Gradient Boosting, hyperparameter tuning, serta soft voting ensemble. Hasil evaluasi menunjukkan bahwa model dapat digunakan sebagai sinyal prioritas retensi.",
        "Nilai accuracy, precision, recall, F1-score, dan AUC memberikan gambaran performa yang lebih lengkap dibanding accuracy saja. Dalam konteks bisnis, recall membantu menangkap pelanggan berisiko, sedangkan precision membantu mengontrol biaya intervensi. Dengan demikian, hasil model dapat diterjemahkan menjadi strategi risk queue untuk pelanggan berisiko tinggi, sedang, dan rendah.",
        "Pengembangan lanjutan dapat dilakukan dengan dataset lokal Indonesia, calibration curve, interpretasi lokal seperti SHAP, integrasi database pelanggan, monitoring drift, serta deployment produksi pada VPS. Penelitian juga dapat diperluas dengan validasi multi-dataset agar generalisasi model churn lebih kuat.",
    ]:
        p(doc, text)
    h(doc, "Referensi")
    for i, ref in enumerate(REFERENCES, 1):
        p(doc, f"[{i}] {ref}", size=7.0)

    out = OUT_DIR / "Draft_Artikel_IEEE.docx"
    doc.save(out)
    with ZipFile(out) as z:
        comments = [n for n in z.namelist() if "comments" in n.lower()]
        if comments:
            raise RuntimeError(comments)


def build_pdf() -> None:
    metrics = json.loads((BASE / "07_Hasil_Eksperimen" / "metrics.json").read_text(encoding="utf-8"))
    comp = pd.read_csv(BASE / "07_Hasil_Eksperimen" / "model_comparison.csv")
    pages = [
        f"{TITLE}\n{AUTHOR}, {NIM}\n\nAbstrak\nRetainA adalah sistem prediksi churn pelanggan berbasis ensemble learning. Dataset utama adalah IBM Telco Customer Churn dengan {metrics['dataset_rows']} baris dan {metrics['dataset_columns']} kolom. Metode mencakup preprocessing, baseline Logistic Regression, Random Forest, Gradient Boosting, tuning, dan soft voting ensemble. Model memperoleh accuracy {metrics['metrics']['accuracy']:.3f}, precision {metrics['metrics']['precision']:.3f}, recall {metrics['metrics']['recall']:.3f}, F1-score {metrics['metrics']['f1']:.3f}, dan AUC {metrics['metrics']['auc']:.3f}.\n\nKata Kunci: churn pelanggan, ensemble learning, risk scoring, deployment lokal.",
        "I. Pendahuluan\nChurn pelanggan menyebabkan hilangnya pendapatan berulang dan meningkatkan biaya akuisisi ulang. RetainA membaca pola risiko churn dari data pelanggan telco dan mengubahnya menjadi skor prioritas retensi. Skor ini menjadi alat bantu keputusan, bukan keputusan final otomatis.\n\nII. Penelitian Terkait\nPenelitian churn banyak menggunakan hybrid neural network, optimized classifiers, XGBoost, SHAP, dan weighted ensemble. Gap utama adalah kurangnya integrasi antara eksperimen, baseline, tuning, interpretasi risiko bisnis, dan web inferensi lokal.",
        "III. Research Gap dan Novelty\nNovelty RetainA adalah integrasi Logistic Regression, Random Forest, Gradient Boosting, RandomizedSearchCV, soft voting ensemble, dan dashboard risk scoring lokal. Kebaruan berada pada integrasi metode yang dapat diverifikasi melalui dataset, kode, model, hasil eksperimen, visualisasi, dan deployment evidence.\n\nIV. Metode\nPreprocessing dilakukan dengan imputation, standardization, dan one-hot encoding. Soft voting menggabungkan probabilitas classifier untuk menghasilkan skor risiko churn.",
        "V. Dataset dan Setup Eksperimen\nDataset IBM Telco Customer Churn dipilih karena publik, relevan, dan memiliki target klasifikasi yang jelas. Evaluasi menggunakan accuracy, precision, recall, F1-score, AUC, ROC curve, dan confusion matrix.\n\nVI. Hasil Eksperimen\n" + comp.to_string(index=False),
        "VII. Analisis Hasil dan Interpretasi Bisnis\nRecall membantu menangkap pelanggan berisiko agar tidak terlewat, sedangkan precision mengurangi biaya retensi yang tidak perlu. Confusion matrix menjelaskan false positive dan false negative. ROC curve membantu membaca threshold, sedangkan feature importance membantu memahami faktor yang memengaruhi churn.",
        "VIII. Deployment dan Ancaman Validitas\nDeployment lokal menggunakan Flask dan model .pkl. Sistem memisahkan training dan inference. Ancaman validitas meliputi dataset publik yang belum tentu mewakili data lokal, class imbalance, temporal drift, dan keterbatasan interpretabilitas.\n\nIX. Kesimpulan\nRetainA dapat menjadi prototipe sistem pendukung keputusan retensi pelanggan. Pengembangan berikutnya mencakup dataset lokal, calibration, SHAP, database, monitoring, dan deployment produksi.\n\nReferensi\n" + "\n".join(f"[{i}] {r}" for i, r in enumerate(REFERENCES, 1)),
    ]
    out = OUT_DIR / "Draft_Artikel_IEEE.pdf"
    with PdfPages(out) as pdf:
        for idx, content in enumerate(pages, 1):
            fig = plt.figure(figsize=(8.5, 11))
            fig.text(0.5, 0.965, "Draft Artikel IEEE - RetainA", ha="center", fontsize=9, family="serif")
            fig.text(0.08, 0.94, f"Halaman {idx}", fontsize=8, family="serif")
            lines = []
            for paragraph in content.split("\n"):
                line = ""
                for word in paragraph.split():
                    if len(line) + len(word) + 1 > 54:
                        lines.append(line)
                        line = word
                    else:
                        line = (line + " " + word).strip()
                if line:
                    lines.append(line)
                lines.append("")
            x_positions = [0.08, 0.53]
            col = 0
            y = 0.905
            for line in lines:
                if y < 0.08 and col == 0:
                    col = 1
                    y = 0.905
                elif y < 0.08:
                    break
                fig.text(x_positions[col], y, line, fontsize=8.0, family="serif")
                y -= 0.019 if line else 0.014
            pdf.savefig(fig)
            plt.close(fig)


if __name__ == "__main__":
    build_docx()
    build_pdf()
    (OUT_DIR / "README_FOLDER.md").write_text(
        "Draft Artikel IEEE Arlen versi Bahasa Indonesia, fokus pada proyek RetainA, layout dua kolom, 6 halaman PDF pendamping, 22 referensi IEEE style, tanpa komentar dokumen.\n",
        encoding="utf-8",
    )
    print("Arlen IEEE 6-page draft rebuilt")
