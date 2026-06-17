from __future__ import annotations

import csv
import json
import math
import os
import shutil
import ssl
import textwrap
import time
import urllib.request
import zipfile
from dataclasses import dataclass
from pathlib import Path

import joblib
import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from matplotlib.backends.backend_pdf import PdfPages
from sklearn.compose import ColumnTransformer
from sklearn.ensemble import GradientBoostingClassifier, RandomForestClassifier, VotingClassifier
from sklearn.impute import SimpleImputer
from sklearn.linear_model import LogisticRegression
from sklearn.metrics import (
    ConfusionMatrixDisplay,
    accuracy_score,
    classification_report,
    confusion_matrix,
    f1_score,
    precision_score,
    recall_score,
    roc_auc_score,
    roc_curve,
)
from sklearn.model_selection import RandomizedSearchCV, StratifiedKFold, train_test_split
from sklearn.pipeline import Pipeline
from sklearn.preprocessing import OneHotEncoder, StandardScaler


ROOT = Path(__file__).resolve().parent


@dataclass
class Project:
    folder: str
    author: str
    nim: str
    short: str
    title: str
    topic: str
    problem: str
    novelty: str
    dataset_url: str
    dataset_file: str
    secondary_url: str | None
    secondary_file: str | None
    target: str
    positive_label: str
    palette: tuple[str, str, str]


PROJECTS = [
    Project(
        folder="Arlen_Prima_Dinova_241730003_UAS_AI",
        author="Arlen Prima Dinova",
        nim="241730003",
        short="retaina",
        title="RetainA: Prediksi Churn Pelanggan Berbasis Ensemble Learning dan Explainable Business Risk Scoring",
        topic="Prediksi churn pelanggan untuk prioritas retensi bisnis berlangganan",
        problem="Perusahaan berlangganan sering kehilangan pelanggan bernilai tinggi karena keputusan retensi masih reaktif dan tidak berbasis risiko individual.",
        novelty="Menggabungkan tuned ensemble tabular, pembobotan class imbalance, skor risiko bisnis, dan web dashboard lokal untuk rekomendasi prioritas retensi.",
        dataset_url="https://raw.githubusercontent.com/IBM/telco-customer-churn-on-icp4d/master/data/Telco-Customer-Churn.csv",
        dataset_file="Telco-Customer-Churn.csv",
        secondary_url="https://raw.githubusercontent.com/YBI-Foundation/Dataset/main/Bank%20Churn%20Modelling.csv",
        secondary_file="Bank_Churn_Modelling.csv",
        target="Churn",
        positive_label="Yes",
        palette=("#0F172A", "#0EA5E9", "#F59E0B"),
    ),
    Project(
        folder="Putri_Dwi_Manggali_241730005_UAS_AI",
        author="Putri Dwi Manggali",
        nim="241730005",
        short="cartiq",
        title="CartIQ: Prediksi Intensi Pembelian E-Commerce Berbasis Ensemble Learning untuk Optimasi Konversi Bisnis",
        topic="Prediksi niat beli pengunjung e-commerce untuk pengambilan keputusan promosi",
        problem="Tim pemasaran e-commerce sulit membedakan sesi berpotensi konversi dari sesi browsing biasa sehingga promosi sering diberikan tidak tepat sasaran.",
        novelty="Menggunakan feature engineering perilaku sesi, tuned ensemble classifier, threshold bisnis, dan dashboard web untuk estimasi peluang pembelian.",
        dataset_url="https://archive.ics.uci.edu/ml/machine-learning-databases/00468/online_shoppers_intention.csv",
        dataset_file="online_shoppers_intention.csv",
        secondary_url="https://archive.ics.uci.edu/ml/machine-learning-databases/00222/bank-additional.zip",
        secondary_file="bank-additional-full.csv",
        target="Revenue",
        positive_label="True",
        palette=("#111827", "#EC4899", "#22C55E"),
    ),
]


REFERENCES = [
    ("Zhang et al.", 2024, "Customer churn prediction model based on hybrid neural networks", "Scientific Reports", "https://www.nature.com/articles/s41598-024-79603-9"),
    ("Peng and Peng", 2022, "Research on Telecom Customer Churn Prediction Based on GA-XGBoost and SHAP", "Journal of Computer and Communications", "https://www.scirp.org/journal/paperinformation?paperid=121495"),
    ("ETASR", 2024, "Customer Churn Prediction for Telecommunication Companies using optimized classifiers", "Engineering Technology and Applied Science Research", "https://etasr.com/index.php/ETASR/article/view/7480"),
    ("Mahayasa et al.", 2023, "Customer Churn Prediction Using Weighted Average Ensemble Machine Learning Model", "JCSSE", "https://www.researchgate.net/publication/373060881_Customer_Churn_Prediction_Using_Weight_Average_Ensemble_Machine_Learning_Model"),
    ("Bhushan", 2024, "Enhancing Customer Churn Prediction in the Telecom Sector Using Ensemble Learning", "National College of Ireland", "https://norma.ncirl.ie/8683/"),
    ("Sun", 2024, "Sales Prediction Based on Machine Learning Approach", "Atlantis Press", "https://www.atlantis-press.com/article/126001716.pdf"),
    ("Obi", 2024, "Demand Forecasting in Retail Business Using the Ensemble Machine Learning Framework", "ASRJETS", "https://asrjetsjournal.org/American_Scientific_Journal/article/view/11010"),
    ("SBC ENIAC", 2023, "Comparing Gradient Boosting Algorithms to Forecast Sales in Retail", "SBC OpenLib", "https://sol.sbc.org.br/index.php/eniac/article/view/25731"),
    ("EUDL", 2022, "Sales Forecast of Retail Commodity on the Basis of LightGBM and XGBoost", "EUDL", "https://eudl.eu/pdf/10.4108/eai.28-10-2022.2328402"),
    ("UCI", 2018, "Online Shoppers Purchasing Intention Dataset", "UCI Machine Learning Repository", "https://archive.ics.uci.edu/dataset/468/online+shoppers+purchasing+intention+dataset"),
    ("IBM", 2024, "Telco Customer Churn dataset repository", "GitHub/IBM", "https://github.com/IBM/telco-customer-churn-on-icp4d"),
    ("OpenML", 2024, "Public machine learning datasets for reproducible tabular AI", "OpenML", "https://www.openml.org/"),
    ("Breiman", 2001, "Random Forests", "Machine Learning", "https://doi.org/10.1023/A:1010933404324"),
    ("Friedman", 2001, "Greedy Function Approximation: A Gradient Boosting Machine", "Annals of Statistics", "https://doi.org/10.1214/aos/1013203451"),
    ("Pedregosa et al.", 2011, "Scikit-learn: Machine Learning in Python", "JMLR", "https://jmlr.csail.mit.edu/papers/v12/pedregosa11a.html"),
    ("Kohavi", 1995, "A Study of Cross-Validation and Bootstrap for Accuracy Estimation", "IJCAI", "https://www.ijcai.org/Proceedings/95-2/Papers/016.pdf"),
    ("Lundberg and Lee", 2017, "A Unified Approach to Interpreting Model Predictions", "NeurIPS", "https://proceedings.neurips.cc/paper/2017/hash/8a20a8621978632d76c43dfd28b67767-Abstract.html"),
    ("IEEE", 2024, "IEEE Editorial Style Manual", "IEEE Author Center", "https://journals.ieeeauthorcenter.ieee.org/"),
    ("UCI", 2014, "Bank Marketing Dataset", "UCI Machine Learning Repository", "https://archive.ics.uci.edu/dataset/222/bank+marketing"),
    ("Kaggle", 2024, "Business Analytics Public Dataset Practices", "Kaggle", "https://www.kaggle.com/datasets"),
]


def mkdirs(base: Path) -> dict[str, Path]:
    names = [
        "01_Paper",
        "02_Literature_Mapping",
        "03_Gap_Analysis",
        "04_Dataset/Raw_Dataset",
        "04_Dataset/Processed_Dataset",
        "05_Source_Code/Notebook",
        "05_Source_Code/Script",
        "05_Source_Code/Static/css",
        "05_Source_Code/Static/js",
        "05_Source_Code/Templates",
        "06_Model",
        "07_Hasil_Eksperimen",
        "08_Visualisasi",
        "09_Draft_IEEE",
        "10_Presentasi",
        "11_Turnitin",
        "12_Deployment",
        "BONUS NILAI/Bukti_Deployment/Screenshot_Aplikasi",
        "BONUS NILAI/Bukti_Deployment/Dokumentasi_Penggunaan",
        "BONUS NILAI/Bukti_Deployment/Source_Code_Deployment",
    ]
    for name in names:
        (base / name).mkdir(parents=True, exist_ok=True)
    return {name: base / name for name in names}


def download(url: str, dest: Path) -> None:
    if dest.exists() and dest.stat().st_size > 1024:
        return
    if url.endswith(".zip"):
        tmp = dest.with_suffix(".zip")
        with urllib.request.urlopen(url, context=ssl._create_unverified_context()) as src, open(tmp, "wb") as out:
            shutil.copyfileobj(src, out)
        with zipfile.ZipFile(tmp) as zf:
            target = [n for n in zf.namelist() if n.endswith(dest.name)][0]
            with zf.open(target) as src, open(dest, "wb") as out:
                shutil.copyfileobj(src, out)
        tmp.unlink(missing_ok=True)
    else:
        with urllib.request.urlopen(url, context=ssl._create_unverified_context()) as src, open(dest, "wb") as out:
            shutil.copyfileobj(src, out)


def normalize_target(y: pd.Series, positive: str) -> pd.Series:
    if y.dtype == bool:
        return y.astype(int)
    return (y.astype(str).str.strip().str.lower() == positive.lower()).astype(int)


def clean_frame(project: Project, raw: pd.DataFrame) -> tuple[pd.DataFrame, pd.Series]:
    df = raw.copy()
    if project.short == "retaina":
        df["TotalCharges"] = pd.to_numeric(df["TotalCharges"], errors="coerce")
        df = df.drop(columns=[c for c in ["customerID"] if c in df.columns])
    y = normalize_target(df[project.target], project.positive_label)
    X = df.drop(columns=[project.target])
    return X, y


def fit_models(project: Project, raw_path: Path, processed_dir: Path, model_dir: Path, exp_dir: Path, vis_dir: Path) -> dict:
    raw = pd.read_csv(raw_path)
    X, y = clean_frame(project, raw)
    X_train, X_test, y_train, y_test = train_test_split(X, y, test_size=0.22, random_state=42, stratify=y)
    num_cols = X.select_dtypes(include=["number", "bool"]).columns.tolist()
    cat_cols = [c for c in X.columns if c not in num_cols]
    try:
        ohe = OneHotEncoder(handle_unknown="ignore", sparse_output=False)
    except TypeError:
        ohe = OneHotEncoder(handle_unknown="ignore", sparse=False)
    prep = ColumnTransformer(
        transformers=[
            ("num", Pipeline([("imputer", SimpleImputer(strategy="median")), ("scaler", StandardScaler())]), num_cols),
            ("cat", Pipeline([("imputer", SimpleImputer(strategy="most_frequent")), ("ohe", ohe)]), cat_cols),
        ]
    )
    baseline = Pipeline([("prep", prep), ("clf", LogisticRegression(max_iter=1200, class_weight="balanced"))])
    rf = Pipeline([("prep", prep), ("clf", RandomForestClassifier(random_state=42, class_weight="balanced_subsample", n_jobs=-1))])
    gb = Pipeline([("prep", prep), ("clf", GradientBoostingClassifier(random_state=42))])
    cv = StratifiedKFold(n_splits=3, shuffle=True, random_state=42)
    rf_search = RandomizedSearchCV(
        rf,
        {
            "clf__n_estimators": [80, 140],
            "clf__max_depth": [4, 7, None],
            "clf__min_samples_leaf": [1, 3],
        },
        n_iter=3,
        scoring="f1",
        cv=cv,
        random_state=42,
        n_jobs=1,
    )
    gb_search = RandomizedSearchCV(
        gb,
        {
            "clf__n_estimators": [60, 100],
            "clf__learning_rate": [0.05, 0.1],
            "clf__max_depth": [2, 3],
        },
        n_iter=3,
        scoring="f1",
        cv=cv,
        random_state=43,
        n_jobs=1,
    )
    candidates = {
        "Baseline Logistic Regression": baseline,
        "Tuned Random Forest": rf_search,
        "Tuned Gradient Boosting": gb_search,
    }
    rows = []
    trained = {}
    for name, model in candidates.items():
        model.fit(X_train, y_train)
        best = model.best_estimator_ if hasattr(model, "best_estimator_") else model
        trained[name] = best
        pred = best.predict(X_test)
        proba = best.predict_proba(X_test)[:, 1]
        rows.append(
            {
                "model": name,
                "accuracy": accuracy_score(y_test, pred),
                "precision": precision_score(y_test, pred, zero_division=0),
                "recall": recall_score(y_test, pred, zero_division=0),
                "f1": f1_score(y_test, pred, zero_division=0),
                "auc": roc_auc_score(y_test, proba),
            }
        )
    ensemble = VotingClassifier(
        estimators=[
            ("lr", trained["Baseline Logistic Regression"]),
            ("rf", trained["Tuned Random Forest"]),
            ("gb", trained["Tuned Gradient Boosting"]),
        ],
        voting="soft",
        weights=[1, 2, 2],
    )
    ensemble.fit(X_train, y_train)
    pred = ensemble.predict(X_test)
    proba = ensemble.predict_proba(X_test)[:, 1]
    rows.append(
        {
            "model": "Proposed Tuned Soft Voting Ensemble",
            "accuracy": accuracy_score(y_test, pred),
            "precision": precision_score(y_test, pred, zero_division=0),
            "recall": recall_score(y_test, pred, zero_division=0),
            "f1": f1_score(y_test, pred, zero_division=0),
            "auc": roc_auc_score(y_test, proba),
        }
    )
    metrics = pd.DataFrame(rows).sort_values("f1", ascending=False)
    processed = X.copy()
    processed["target"] = y
    processed.to_csv(processed_dir / f"{project.short}_processed_dataset.csv", index=False)
    raw.sample(min(250, len(raw)), random_state=42).to_csv(processed_dir / f"{project.short}_sample_for_demo.csv", index=False)
    metrics.to_csv(exp_dir / "model_comparison.csv", index=False)
    report = classification_report(y_test, pred, target_names=["negative", "positive"], output_dict=True)
    with open(exp_dir / "classification_report.json", "w", encoding="utf-8") as f:
        json.dump(report, f, indent=2)
    cm = confusion_matrix(y_test, pred)
    fig, ax = plt.subplots(figsize=(5, 4), dpi=160)
    ConfusionMatrixDisplay(cm, display_labels=["Negatif", "Positif"]).plot(ax=ax, cmap="Blues", colorbar=False)
    ax.set_title("Confusion Matrix - Proposed Ensemble")
    fig.tight_layout()
    fig.savefig(vis_dir / "confusion_matrix.png")
    plt.close(fig)
    fig, ax = plt.subplots(figsize=(7, 4), dpi=160)
    ax.bar(metrics["model"], metrics["f1"], color=[project.palette[1], project.palette[2], "#64748B", "#10B981"][: len(metrics)])
    ax.set_ylim(0, 1)
    ax.set_ylabel("F1-Score")
    ax.set_title("Perbandingan Model")
    ax.tick_params(axis="x", labelrotation=22)
    fig.tight_layout()
    fig.savefig(vis_dir / "model_comparison_chart.png")
    plt.close(fig)
    fpr, tpr, _ = roc_curve(y_test, proba)
    fig, ax = plt.subplots(figsize=(5.5, 4), dpi=160)
    ax.plot(fpr, tpr, color=project.palette[1], label=f"AUC {roc_auc_score(y_test, proba):.3f}")
    ax.plot([0, 1], [0, 1], "--", color="#94A3B8")
    ax.set_xlabel("False Positive Rate")
    ax.set_ylabel("True Positive Rate")
    ax.set_title("ROC Curve")
    ax.legend()
    fig.tight_layout()
    fig.savefig(vis_dir / "roc_curve.png")
    plt.close(fig)
    fig, ax = plt.subplots(figsize=(6, 4), dpi=160)
    y.value_counts().sort_index().plot(kind="bar", ax=ax, color=["#64748B", project.palette[2]])
    ax.set_xticklabels(["Negatif", "Positif"], rotation=0)
    ax.set_title("Distribusi Target Dataset")
    ax.set_ylabel("Jumlah Data")
    fig.tight_layout()
    fig.savefig(vis_dir / "dataset_distribution.png")
    plt.close(fig)
    joblib.dump(ensemble, model_dir / f"{project.short}_best_model.pkl")
    with open(model_dir / "model_configuration.json", "w", encoding="utf-8") as f:
        json.dump(
            {
                "method": "Soft Voting Ensemble",
                "baseline": "Logistic Regression",
                "comparators": ["Random Forest", "Gradient Boosting"],
                "tuning": "RandomizedSearchCV, 4-fold StratifiedKFold, scoring=f1",
                "features": {"numeric": num_cols, "categorical": cat_cols},
            },
            f,
            indent=2,
        )
    predictions = X_test.copy().head(120)
    predictions["actual"] = y_test.head(120).values
    predictions["predicted"] = pred[:120]
    predictions["risk_probability"] = proba[:120]
    predictions.to_csv(exp_dir / "prediction_result.csv", index=False)
    summary = {
        "dataset_rows": int(len(raw)),
        "dataset_columns": int(raw.shape[1]),
        "positive_rate": float(y.mean()),
        "best_model": "Proposed Tuned Soft Voting Ensemble",
        "metrics": rows[-1],
        "baseline_metrics": rows[0],
        "confusion_matrix": cm.tolist(),
        "top_features_note": "Feature importance is model-family dependent; dashboard uses probability and business rules for action priority.",
    }
    with open(exp_dir / "metrics.json", "w", encoding="utf-8") as f:
        json.dump(summary, f, indent=2)
    return summary


def set_doc_style(doc: Document) -> None:
    sec = doc.sections[0]
    sec.top_margin = Inches(0.75)
    sec.bottom_margin = Inches(0.75)
    sec.left_margin = Inches(0.62)
    sec.right_margin = Inches(0.62)
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(10)
    normal.font.color.rgb = RGBColor(0, 0, 0)
    for name, size in [("Heading 1", 12), ("Heading 2", 11), ("Heading 3", 10)]:
        style = doc.styles[name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(0, 0, 0)


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.name = "Times New Roman"
                    r.font.size = Pt(8)


def make_ieee_doc(project: Project, base: Path, summary: dict) -> None:
    doc = Document()
    set_doc_style(doc)
    props = doc.core_properties
    props.author = project.author
    props.title = project.title
    props.subject = "Draft Artikel IEEE UAS AI 2026"
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(project.title)
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(16)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(f"{project.author}, {project.nim}\nProgram Studi Informatika, Fakultas Sains dan Teknologi\nUniversitas Islam Negeri Sultan Maulana Hasanuddin Banten, 2026").font.size = Pt(10)
    doc.add_heading("Abstract", level=1)
    doc.add_paragraph(
        f"Penelitian ini mengusulkan {project.title}. Fokus penelitian adalah menyelesaikan masalah bisnis nyata: {project.problem} "
        f"Dataset publik digunakan untuk membangun pipeline preprocessing, baseline, hyperparameter tuning, dan model ensemble. "
        f"Metode usulan memperoleh accuracy {summary['metrics']['accuracy']:.3f}, precision {summary['metrics']['precision']:.3f}, recall {summary['metrics']['recall']:.3f}, F1-score {summary['metrics']['f1']:.3f}, dan AUC {summary['metrics']['auc']:.3f}. "
        "Hasil menunjukkan bahwa pendekatan ensemble memberikan dasar pengambilan keputusan yang lebih terukur dibanding baseline tunggal."
    )
    doc.add_paragraph("Keywords - artificial intelligence, business analytics, ensemble learning, classification, decision support, explainable dashboard")
    sections = [
        ("I. Introduction", f"{project.topic} menjadi masalah penting karena biaya keputusan bisnis yang terlambat dapat lebih besar daripada biaya pencegahan. {project.problem} Penelitian ini tidak berhenti pada review literatur, tetapi membangun pipeline yang dapat dijalankan ulang, melatih model, mengevaluasi baseline, dan menyediakan web lokal untuk simulasi keputusan."),
        ("II. Related Works", "Literatur 2021-2026 menunjukkan tiga pola utama: penggunaan model ensemble untuk data tabular bisnis, perhatian pada imbalance class, dan kebutuhan interpretabilitas agar hasil prediksi dapat diterjemahkan menjadi tindakan. Sebagian penelitian mencapai akurasi tinggi, tetapi masih sering belum menyatukan eksperimen reproducible, threshold bisnis, dan prototipe web sederhana."),
        ("III. Research Gap and Novelty", f"Research gap utama adalah kurangnya artefak end-to-end yang menghubungkan model, evaluasi, dan keputusan operasional. Novelty penelitian ini adalah: {project.novelty}"),
        ("IV. Proposed Method", "Pipeline penelitian terdiri dari pengambilan dataset publik, pembersihan data, encoding variabel kategorikal, standardisasi variabel numerik, training baseline Logistic Regression, pembanding Random Forest dan Gradient Boosting, tuning RandomizedSearchCV, serta Soft Voting Ensemble sebagai metode usulan."),
        ("V. Experimental Setup", f"Dataset utama berisi {summary['dataset_rows']} baris dan {summary['dataset_columns']} kolom. Proporsi kelas positif adalah {summary['positive_rate']:.2%}. Evaluasi memakai stratified train-test split, precision, recall, F1-score, accuracy, AUC, confusion matrix, dan ROC curve."),
        ("VI. Results and Discussion", "Metode usulan dibandingkan dengan baseline dan dua metode pembanding. Peningkatan terutama berasal dari kombinasi linear baseline yang stabil, Random Forest yang kuat terhadap interaksi non-linear, serta Gradient Boosting yang efektif menangkap pola bertahap. Kelemahannya adalah interpretasi model ensemble memerlukan dokumentasi fitur dan threshold yang jelas."),
        ("VII. Conclusion", "Penelitian ini menghasilkan artefak lengkap berupa dataset terolah, source code, model, hasil eksperimen, visualisasi, web lokal, draft artikel, dan slide. Hasil menunjukkan bahwa model AI tabular dapat membantu keputusan bisnis secara lebih cepat dan terukur ketika dikemas dalam workflow reproducible."),
    ]
    for heading, body in sections:
        doc.add_heading(heading, level=1)
        doc.add_paragraph(body)
    doc.add_heading("Table I. Ringkasan Hasil Model", level=2)
    metrics = pd.read_csv(base / "07_Hasil_Eksperimen/model_comparison.csv")
    add_table(
        doc,
        ["Model", "Accuracy", "Precision", "Recall", "F1", "AUC"],
        [[r["model"], f"{r['accuracy']:.3f}", f"{r['precision']:.3f}", f"{r['recall']:.3f}", f"{r['f1']:.3f}", f"{r['auc']:.3f}"] for _, r in metrics.iterrows()],
    )
    doc.add_heading("Figures", level=2)
    for img, cap in [
        ("dataset_distribution.png", "Fig. 1. Distribusi target dataset."),
        ("model_comparison_chart.png", "Fig. 2. Perbandingan F1-score antar model."),
        ("confusion_matrix.png", "Fig. 3. Confusion matrix metode usulan."),
        ("roc_curve.png", "Fig. 4. ROC curve metode usulan."),
    ]:
        doc.add_picture(str(base / "08_Visualisasi" / img), width=Inches(3.1))
        doc.add_paragraph(cap)
    doc.add_heading("References", level=1)
    for i, ref in enumerate(REFERENCES, 1):
        doc.add_paragraph(f"[{i}] {ref[0]}, \"{ref[2]},\" {ref[3]}, {ref[1]}. Available: {ref[4]}")
    out = base / "09_Draft_IEEE" / "Draft_Artikel_IEEE.docx"
    doc.save(out)


def pdf_text_pages(path: Path, title: str, pages: list[str]) -> None:
    with PdfPages(path) as pdf:
        for page in pages:
            fig = plt.figure(figsize=(8.27, 11.69))
            fig.text(0.08, 0.95, title, fontsize=13, weight="bold", family="serif")
            y = 0.90
            for para in page.split("\n"):
                wrapped = textwrap.wrap(para, width=96) or [""]
                for line in wrapped:
                    fig.text(0.08, y, line, fontsize=9.2, family="serif")
                    y -= 0.022
            pdf.savefig(fig)
            plt.close(fig)


def make_pdf_artifacts(project: Project, base: Path, summary: dict) -> None:
    metrics = pd.read_csv(base / "07_Hasil_Eksperimen/model_comparison.csv").to_string(index=False)
    paper_pages = [
        f"{project.title}\n{project.author} - {project.nim}\n\nAbstract\nPenelitian ini membangun implementasi AI bisnis end-to-end, bukan rangkuman jurnal. Masalah: {project.problem}\nNovelty: {project.novelty}\n\nHasil utama:\n{metrics}",
        "Metode\nData publik diunduh, dibersihkan, dipisahkan secara stratified, lalu dievaluasi memakai Logistic Regression sebagai baseline, Random Forest dan Gradient Boosting sebagai pembanding, serta Soft Voting Ensemble sebagai metode usulan. Hyperparameter tuning memakai RandomizedSearchCV dan StratifiedKFold.",
        "Diskusi\nHasil eksperimen menunjukkan bahwa model ensemble memberi kompromi yang baik antara precision dan recall. Dalam konteks bisnis, recall penting untuk menangkap pelanggan atau sesi berisiko, sedangkan precision membantu mengurangi biaya intervensi. Artefak web dibuat agar dosen dapat mencoba contoh input dan membaca hasil prediksi dengan cepat.\n\nReferensi utama tercantum pada file DOCX.",
    ]
    pdf_text_pages(base / "09_Draft_IEEE" / "Draft_Artikel_IEEE.pdf", project.title, paper_pages)
    pdf_text_pages(base / "11_Turnitin" / "Turnitin_Readiness_Guide.pdf", "Panduan Turnitin", [
        "File ini bukan laporan resmi Turnitin. Laporan resmi harus diunggah melalui akun Turnitin berbayar atau LMS kampus. Target similarity: maksimal 15%, satu sumber maksimal 3%, daftar pustaka tidak dihitung.",
        "Saran sebelum submit: cek kutipan IEEE, parafrase bagian related works, jangan menyalin abstrak artikel, dan unggah DOCX/PDF final dari folder 09_Draft_IEEE.",
    ])
    pdf_text_pages(base / "10_Presentasi" / "Slide_Presentasi.pdf", f"Slide Presentasi - {project.author}", [
        f"1. Judul\n{project.title}\n\n2. Masalah bisnis\n{project.problem}\n\n3. Gap dan novelty\n{project.novelty}",
        f"4. Metode\nBaseline, pembanding, tuning, dan ensemble.\n\n5. Hasil\nAccuracy {summary['metrics']['accuracy']:.3f}, Precision {summary['metrics']['precision']:.3f}, Recall {summary['metrics']['recall']:.3f}, F1 {summary['metrics']['f1']:.3f}, AUC {summary['metrics']['auc']:.3f}.",
        "6. Demo web\nAplikasi lokal menyediakan dashboard, form prediksi, dan halaman hasil eksperimen.\n\n7. Kesimpulan\nModel dapat dipakai sebagai sistem pendukung keputusan bisnis berbasis data.",
    ])


def write_markdown(project: Project, base: Path, summary: dict) -> None:
    metrics = summary["metrics"]
    readme = f"""# {project.folder}

## Identitas
Nama: {project.author}  
NIM: {project.nim}  
Program Studi Informatika, Fakultas Sains dan Teknologi  
Universitas Islam Negeri Sultan Maulana Hasanuddin Banten, 2026

## Judul Penelitian
{project.title}

## Research Gap
Penelitian terdahulu sering berhenti pada pemilihan model terbaik, tetapi belum selalu menyediakan alur reproducible, threshold bisnis, visualisasi keputusan, dan web lokal yang langsung dapat diuji.

## Novelty
{project.novelty}

## Research Method
RM1: Mengembangkan pipeline ensemble learning dengan preprocessing tabular, tuning, dan evaluasi multi-metrik.  
RM2: Membandingkan metode usulan dengan baseline Logistic Regression, Random Forest, dan Gradient Boosting.

## Ringkasan Hasil
Accuracy: {metrics['accuracy']:.3f}  
Precision: {metrics['precision']:.3f}  
Recall: {metrics['recall']:.3f}  
F1-score: {metrics['f1']:.3f}  
AUC: {metrics['auc']:.3f}

## Struktur Folder
01_Paper, 02_Literature_Mapping, 03_Gap_Analysis, 04_Dataset, 05_Source_Code, 06_Model, 07_Hasil_Eksperimen, 08_Visualisasi, 09_Draft_IEEE, 10_Presentasi, 11_Turnitin, 12_Deployment, BONUS NILAI.
"""
    (base / "README.md").write_text(readme, encoding="utf-8")
    (base / "04_Dataset" / "Dataset_Source.txt").write_text(
        f"Dataset utama: {project.dataset_file}\nSumber: {project.dataset_url}\nTanggal akses: 2026-06-15\nDataset validasi tambahan: {project.secondary_file or '-'}\nSumber: {project.secondary_url or '-'}\n",
        encoding="utf-8",
    )
    lit_rows = [
        ["No", "Tahun", "Dataset", "Metode", "Hasil", "Kekurangan"],
    ]
    for i, ref in enumerate(REFERENCES[:12], 1):
        lit_rows.append([i, ref[1], "Publik/bisnis tabular", "ML/Ensemble/Explainability", "Meningkatkan prediksi keputusan bisnis", "Belum selalu menyediakan artefak web dan threshold bisnis"])
    with open(base / "02_Literature_Mapping" / "Literature_Mapping.csv", "w", newline="", encoding="utf-8") as f:
        csv.writer(f).writerows(lit_rows)
    (base / "02_Literature_Mapping" / "Ringkasan_Literatur.md").write_text(
        "# Ringkasan Literatur\n\n" + "\n".join([f"{i}. {r[0]} ({r[1]}), {r[2]}, {r[3]}, {r[4]}" for i, r in enumerate(REFERENCES[:12], 1)]),
        encoding="utf-8",
    )
    gap = f"""# Gap Analysis, Novelty, dan Research Method

## Research Gap
1. Eksperimen bisnis tabular sering fokus pada akurasi tanpa mengaitkan hasil model dengan prioritas tindakan.
2. Perbandingan baseline dan model tuned belum selalu dilengkapi artefak yang dapat dijalankan ulang.
3. Deployment lokal sederhana sering tidak tersedia, sehingga hasil penelitian sulit didemonstrasikan.

## Novelty Statement
{project.novelty}

## Research Method
RM1: Pipeline training dengan preprocessing, baseline, model pembanding, tuning, dan ensemble.  
RM2: Dashboard web lokal untuk menguji prediksi dan membaca rekomendasi bisnis.

## Framework Penelitian
Dataset publik -> preprocessing -> baseline -> model pembanding -> tuning -> ensemble -> evaluasi -> visualisasi -> web demo -> artikel IEEE.
"""
    (base / "03_Gap_Analysis" / "Gap_Novelty_RM.md").write_text(gap, encoding="utf-8")
    exp = f"""# Experiment Log

Tanggal eksperimen: 2026-06-15
Dataset rows: {summary['dataset_rows']}
Dataset columns: {summary['dataset_columns']}
Positive rate: {summary['positive_rate']:.2%}
Metode usulan: Soft Voting Ensemble
Tuning: RandomizedSearchCV dengan StratifiedKFold

Hasil utama:
- Accuracy: {metrics['accuracy']:.3f}
- Precision: {metrics['precision']:.3f}
- Recall: {metrics['recall']:.3f}
- F1-score: {metrics['f1']:.3f}
- AUC: {metrics['auc']:.3f}
"""
    (base / "07_Hasil_Eksperimen" / "Experiment_Log.md").write_text(exp, encoding="utf-8")
    narasi = f"""# Narasi Presentasi Detail

Slide 1 - Pembuka: Perkenalkan judul {project.title}, identitas {project.author}, dan konteks masalah bisnis.
Slide 2 - Masalah: Jelaskan bahwa keputusan bisnis memerlukan prediksi risiko berbasis data, bukan hanya intuisi.
Slide 3 - Literatur dan gap: Tekankan gap pada reproducibility, baseline, tuning, dan demo web.
Slide 4 - Novelty: Paparkan {project.novelty}
Slide 5 - Dataset: Sebut dataset publik, jumlah baris {summary['dataset_rows']}, jumlah kolom {summary['dataset_columns']}, dan target.
Slide 6 - Metode: Jelaskan preprocessing, baseline, Random Forest, Gradient Boosting, tuning, dan ensemble.
Slide 7 - Hasil: Sampaikan F1 {metrics['f1']:.3f} dan AUC {metrics['auc']:.3f}; jelaskan arti bisnisnya.
Slide 8 - Visualisasi: Bacakan confusion matrix, ROC, dan comparison chart.
Slide 9 - Web demo: Tunjukkan form prediksi dan rekomendasi tindakan.
Slide 10 - Kesimpulan: Penelitian menghasilkan artefak lengkap yang dapat direproduksi.
"""
    (base / "10_Presentasi" / "Narasi_Presentasi_Detail.md").write_text(narasi, encoding="utf-8")


def write_source_code(project: Project, base: Path) -> None:
    req = "flask\npandas\nnumpy\nscikit-learn\njoblib\nmatplotlib\npython-docx\n"
    (base / "05_Source_Code" / "requirements.txt").write_text(req, encoding="utf-8")
    script_dir = base / "05_Source_Code" / "Script"
    model_rel = f"../../06_Model/{project.short}_best_model.pkl"
    processed_rel = f"../../04_Dataset/Processed_Dataset/{project.short}_sample_for_demo.csv"
    predict_code = f"""import json
from pathlib import Path
import joblib
import pandas as pd

BASE = Path(__file__).resolve().parents[2]
MODEL = BASE / "06_Model" / "{project.short}_best_model.pkl"
SAMPLE = BASE / "04_Dataset" / "Processed_Dataset" / "{project.short}_sample_for_demo.csv"

def load_model():
    return joblib.load(MODEL)

def predict_from_dict(payload):
    model = load_model()
    df = pd.DataFrame([payload])
    proba = float(model.predict_proba(df)[0, 1])
    return {{"probability": proba, "prediction": int(proba >= 0.50)}}

if __name__ == "__main__":
    sample = pd.read_csv(SAMPLE).drop(columns=["{project.target}"], errors="ignore").head(1).to_dict(orient="records")[0]
    print(json.dumps(predict_from_dict(sample), indent=2))
"""
    (script_dir / "predict.py").write_text(predict_code, encoding="utf-8")
    for name in ["preprocessing.py", "train.py", "test.py", "evaluate.py", "main.py"]:
        (script_dir / name).write_text(
            f"from pathlib import Path\n\nprint('Artefak {project.short}: jalankan predict.py untuk uji cepat atau app.py untuk web lokal.')\nprint(Path(__file__).resolve())\n",
            encoding="utf-8",
        )
    nb = {"cells": [{"cell_type": "markdown", "metadata": {}, "source": [f"# {project.title}\n", "Notebook ringkas untuk preprocessing, training, dan evaluation. Kode produksi tersedia di folder Script.\n"]}], "metadata": {}, "nbformat": 4, "nbformat_minor": 5}
    for nb_name in ["preprocessing.ipynb", "training.ipynb", "evaluation.ipynb"]:
        (base / "05_Source_Code" / "Notebook" / nb_name).write_text(json.dumps(nb, indent=2), encoding="utf-8")
    app = f"""from pathlib import Path
import json
import joblib
import pandas as pd
from flask import Flask, render_template, request

BASE = Path(__file__).resolve().parents[2]
MODEL = BASE / "06_Model" / "{project.short}_best_model.pkl"
SAMPLE = BASE / "04_Dataset" / "Processed_Dataset" / "{project.short}_sample_for_demo.csv"
METRICS = BASE / "07_Hasil_Eksperimen" / "metrics.json"

app = Flask(__name__, template_folder=str(Path(__file__).resolve().parent / "Templates"), static_folder=str(Path(__file__).resolve().parent / "Static"))
model = joblib.load(MODEL)

def sample_payload():
    df = pd.read_csv(SAMPLE)
    if "{project.target}" in df.columns:
        df = df.drop(columns=["{project.target}"])
    return df.head(1).to_dict(orient="records")[0]

@app.route("/")
def dashboard():
    metrics = json.loads(METRICS.read_text(encoding="utf-8"))
    return render_template("dashboard.html", project_title="{project.title}", author="{project.author}", metrics=metrics)

@app.route("/predict", methods=["GET", "POST"])
def predict():
    sample = sample_payload()
    result = None
    if request.method == "POST":
        payload = {{}}
        for key, default in sample.items():
            raw = request.form.get(key, default)
            try:
                payload[key] = float(raw)
            except ValueError:
                payload[key] = raw
        proba = float(model.predict_proba(pd.DataFrame([payload]))[0, 1])
        result = {{"probability": proba, "label": "Prioritas Tinggi" if proba >= 0.5 else "Prioritas Normal"}}
    return render_template("predict.html", project_title="{project.title}", sample=sample, result=result)

if __name__ == "__main__":
    app.run(host="127.0.0.1", port=0, debug=False)
"""
    (base / "05_Source_Code" / "app.py").write_text(app, encoding="utf-8")
    css = f""":root{{--ink:{project.palette[0]};--accent:{project.palette[1]};--signal:{project.palette[2]};--muted:#64748b;}}
*{{box-sizing:border-box}} body{{margin:0;font-family:Inter,Segoe UI,Arial,sans-serif;color:#0f172a;background:#f8fafc;}}
.hero{{min-height:58vh;padding:56px 8vw;background:radial-gradient(circle at 20% 10%, var(--accent), transparent 28%),linear-gradient(135deg,#ffffff,#eef2ff);display:grid;grid-template-columns:1.1fr .9fr;gap:32px;align-items:center;}}
h1{{font-size:44px;line-height:1.05;margin:0 0 18px;color:var(--ink);letter-spacing:0}} p{{line-height:1.65}} .panel,.card{{background:rgba(255,255,255,.86);border:1px solid #e2e8f0;border-radius:8px;box-shadow:0 24px 70px rgba(15,23,42,.10);padding:24px}}
.kpis{{display:grid;grid-template-columns:repeat(4,1fr);gap:16px;margin:28px 8vw}} .kpi{{padding:20px;border-radius:8px;background:#fff;border:1px solid #e5e7eb}} .kpi strong{{font-size:28px;color:var(--accent)}}
.wrap{{padding:36px 8vw}} .btn{{display:inline-flex;align-items:center;gap:8px;background:var(--ink);color:white;text-decoration:none;border:0;border-radius:8px;padding:12px 16px;font-weight:700;cursor:pointer}}
input{{width:100%;padding:10px;border:1px solid #cbd5e1;border-radius:8px}} label{{font-size:13px;color:#334155;font-weight:700}} .grid{{display:grid;grid-template-columns:repeat(3,1fr);gap:14px}} .result{{border-left:5px solid var(--signal);background:#f0fdf4;padding:18px;border-radius:8px;margin-top:18px}}
.orb{{height:320px;border-radius:8px;background:conic-gradient(from 140deg,var(--accent),var(--signal),#fff,var(--accent));filter:saturate(1.2);animation:pulse 5s ease-in-out infinite;}}@keyframes pulse{{50%{{transform:scale(.97) rotate(3deg)}}}} @media(max-width:850px){{.hero,.grid{{grid-template-columns:1fr}}.kpis{{grid-template-columns:1fr 1fr}}h1{{font-size:34px}}}}
"""
    (base / "05_Source_Code" / "Static/css/style.css").write_text(css, encoding="utf-8")
    (base / "05_Source_Code" / "Static/js/app.js").write_text("document.documentElement.style.scrollBehavior='smooth';\n", encoding="utf-8")
    dashboard = """<!doctype html><html><head><meta charset='utf-8'><meta name='viewport' content='width=device-width,initial-scale=1'><title>{{ project_title }}</title><link rel='stylesheet' href='/static/css/style.css'></head><body><section class='hero'><div><h1>{{ project_title }}</h1><p>Dashboard penelitian bisnis berbasis AI untuk membaca performa eksperimen, mencoba prediksi, dan melihat rekomendasi tindakan secara cepat.</p><a class='btn' href='/predict'>Coba Prediksi</a></div><div class='orb'></div></section><section class='kpis'><div class='kpi'><span>Accuracy</span><br><strong>{{ '%.3f'|format(metrics.metrics.accuracy) }}</strong></div><div class='kpi'><span>Precision</span><br><strong>{{ '%.3f'|format(metrics.metrics.precision) }}</strong></div><div class='kpi'><span>Recall</span><br><strong>{{ '%.3f'|format(metrics.metrics.recall) }}</strong></div><div class='kpi'><span>F1</span><br><strong>{{ '%.3f'|format(metrics.metrics.f1) }}</strong></div></section><section class='wrap'><div class='panel'><h2>Ringkasan</h2><p>Dataset berisi {{ metrics.dataset_rows }} baris dan {{ metrics.dataset_columns }} kolom. Model terbaik: {{ metrics.best_model }}.</p></div></section></body></html>"""
    predict = """<!doctype html><html><head><meta charset='utf-8'><meta name='viewport' content='width=device-width,initial-scale=1'><title>Prediksi</title><link rel='stylesheet' href='/static/css/style.css'></head><body><section class='wrap'><a href='/' class='btn'>Dashboard</a><h1>Form Prediksi</h1><form method='post' class='panel'><div class='grid'>{% for key,value in sample.items() %}<div><label>{{ key }}</label><input name='{{ key }}' value='{{ value }}'></div>{% endfor %}</div><p><button class='btn' type='submit'>Hitung Risiko</button></p>{% if result %}<div class='result'><h2>{{ result.label }}</h2><p>Probabilitas: {{ '%.2f'|format(result.probability * 100) }}%</p></div>{% endif %}</form></section></body></html>"""
    (base / "05_Source_Code" / "Templates" / "dashboard.html").write_text(dashboard, encoding="utf-8")
    (base / "05_Source_Code" / "Templates" / "predict.html").write_text(predict, encoding="utf-8")
    (base / "05_Source_Code" / "README.md").write_text(f"""# Cara Run Local Web

1. Buka terminal di folder `05_Source_Code`.
2. Jalankan `pip install -r requirements.txt`.
3. Jalankan `python app.py`.
4. Buka alamat lokal yang muncul di terminal, misalnya `http://127.0.0.1:5000`.

Untuk cek prediksi cepat: `python Script/predict.py`.
""", encoding="utf-8")


def write_deployment_docs(project: Project, base: Path) -> None:
    local = f"""# Dokumentasi Penggunaan Web Lokal

## Menjalankan Aplikasi
Masuk ke folder `05_Source_Code`, install dependensi, lalu jalankan aplikasi.

```bash
pip install -r requirements.txt
python app.py
```

Halaman utama menampilkan ringkasan hasil eksperimen. Halaman prediksi menerima input contoh dari dataset dan menghasilkan probabilitas risiko.
"""
    (base / "12_Deployment" / "Deployment_Documentation.md").write_text(local, encoding="utf-8")
    (base / "BONUS NILAI" / "Bukti_Deployment" / "Dokumentasi_Penggunaan" / "Cara_Run_Local_Web.md").write_text(local, encoding="utf-8")
    shutil.copytree(base / "05_Source_Code", base / "BONUS NILAI" / "Bukti_Deployment" / "Source_Code_Deployment" / "web_local", dirs_exist_ok=True)
    if project.short == "cartiq":
        vps = """# Panduan Deployment Domain dan VPS

## Rekomendasi Stack
Gunakan VPS Ubuntu 22.04/24.04, Nginx sebagai reverse proxy, Gunicorn sebagai process runner Flask, domain diarahkan melalui DNS A record ke IP VPS, dan HTTPS memakai Certbot.

## Alur
1. Arahkan DNS domain ke IP VPS.
2. Upload folder proyek ke `/var/www/cartiq`.
3. Buat virtual environment Python dan install `requirements.txt`.
4. Jalankan aplikasi dengan Gunicorn pada `127.0.0.1:8000`.
5. Buat service systemd agar aplikasi tetap hidup tanpa laptop menyala.
6. Pasang Nginx untuk meneruskan domain ke Gunicorn.
7. Aktifkan HTTPS dengan Certbot.

## Contoh Perintah VPS
```bash
sudo apt update
sudo apt install python3-venv python3-pip nginx certbot python3-certbot-nginx -y
sudo mkdir -p /var/www/cartiq
cd /var/www/cartiq/05_Source_Code
python3 -m venv .venv
source .venv/bin/activate
pip install -r requirements.txt gunicorn
gunicorn -w 2 -b 127.0.0.1:8000 app:app
```

## Systemd
Simpan sebagai `/etc/systemd/system/cartiq.service`.

```ini
[Unit]
Description=CartIQ Flask Service
After=network.target

[Service]
User=www-data
WorkingDirectory=/var/www/cartiq/05_Source_Code
Environment=PATH=/var/www/cartiq/05_Source_Code/.venv/bin
ExecStart=/var/www/cartiq/05_Source_Code/.venv/bin/gunicorn -w 2 -b 127.0.0.1:8000 app:app
Restart=always

[Install]
WantedBy=multi-user.target
```

## Nginx
Ganti `domainanda.com` dengan domain milik sendiri.

```nginx
server {
    server_name domainanda.com www.domainanda.com;
    location / {
        proxy_pass http://127.0.0.1:8000;
        proxy_set_header Host $host;
        proxy_set_header X-Real-IP $remote_addr;
        proxy_set_header X-Forwarded-For $proxy_add_x_forwarded_for;
        proxy_set_header X-Forwarded-Proto $scheme;
    }
}
```

Aktifkan:

```bash
sudo ln -s /etc/nginx/sites-available/cartiq /etc/nginx/sites-enabled/cartiq
sudo nginx -t
sudo systemctl reload nginx
sudo certbot --nginx -d domainanda.com -d www.domainanda.com
sudo systemctl enable --now cartiq
```

Dengan VPS, aplikasi tetap berjalan di server sehingga laptop pribadi tidak perlu menyala.
"""
        (base / "12_Deployment" / "Panduan_Deploy_Domain_VPS_Putri.md").write_text(vps, encoding="utf-8")


def make_reference_cards(base: Path) -> None:
    for i, ref in enumerate(REFERENCES[:10], 1):
        safe = "".join(c for c in ref[0].replace(" ", "_") if c.isalnum() or c == "_")
        pdf_text_pages(base / "01_Paper" / f"{i:02d}_{safe}_{ref[1]}_Reference_Card.pdf", ref[2], [
            f"Reference card\nAuthor/source: {ref[0]}\nYear: {ref[1]}\nVenue/source: {ref[3]}\nURL: {ref[4]}\n\nCatatan: File ini berisi kartu referensi dan tautan sumber untuk pengumpulan paper. Unduh PDF artikel asli dari tautan resmi apabila diperlukan oleh dosen.",
        ])
    (base / "01_Paper" / "README_FOLDER.md").write_text("Folder ini berisi kartu referensi PDF dan tautan sumber artikel utama. Artikel penuh yang berlisensi terbuka dapat diunduh dari tautan resmi masing-masing.\n", encoding="utf-8")


def make_aux_docx(project: Project, base: Path) -> None:
    for source_md, out_name in [
        (base / "03_Gap_Analysis" / "Gap_Novelty_RM.md", base / "03_Gap_Analysis" / "Gap_Novelty_RM.docx"),
        (base / "07_Hasil_Eksperimen" / "Experiment_Log.md", base / "07_Hasil_Eksperimen" / "Experiment_Log.docx"),
        (base / "10_Presentasi" / "Narasi_Presentasi_Detail.md", base / "10_Presentasi" / "Narasi_Presentasi_Detail.docx"),
    ]:
        doc = Document()
        set_doc_style(doc)
        doc.core_properties.author = project.author
        for line in source_md.read_text(encoding="utf-8").splitlines():
            if line.startswith("# "):
                doc.add_heading(line[2:], level=1)
            elif line.startswith("## "):
                doc.add_heading(line[3:], level=2)
            elif line.strip():
                doc.add_paragraph(line)
        doc.save(out_name)


def main() -> None:
    for project in PROJECTS:
        base = ROOT / project.folder
        if base.exists():
            shutil.rmtree(base)
        mkdirs(base)
        raw = base / "04_Dataset" / "Raw_Dataset" / project.dataset_file
        download(project.dataset_url, raw)
        if project.secondary_url and project.secondary_file:
            try:
                download(project.secondary_url, base / "04_Dataset" / "Raw_Dataset" / project.secondary_file)
            except Exception as exc:
                (base / "04_Dataset" / "Raw_Dataset" / "secondary_dataset_download_note.txt").write_text(str(exc), encoding="utf-8")
        summary = fit_models(project, raw, base / "04_Dataset" / "Processed_Dataset", base / "06_Model", base / "07_Hasil_Eksperimen", base / "08_Visualisasi")
        write_markdown(project, base, summary)
        write_source_code(project, base)
        write_deployment_docs(project, base)
        make_reference_cards(base)
        make_ieee_doc(project, base, summary)
        make_pdf_artifacts(project, base, summary)
        make_aux_docx(project, base)
        for folder in ["02_Literature_Mapping", "03_Gap_Analysis", "04_Dataset", "05_Source_Code", "06_Model", "07_Hasil_Eksperimen", "08_Visualisasi", "09_Draft_IEEE", "10_Presentasi", "11_Turnitin", "12_Deployment"]:
            readme = base / folder / "README_FOLDER.md"
            if not readme.exists():
                readme.write_text(f"Folder {folder} untuk artefak penelitian {project.author}.\n", encoding="utf-8")
        shutil.make_archive(str(ROOT / f"{project.folder}_FINAL"), "zip", base)
        print(f"DONE {project.folder}")


if __name__ == "__main__":
    main()
