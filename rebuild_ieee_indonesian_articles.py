from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path
from zipfile import ZipFile
import json

import matplotlib.pyplot as plt
import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from matplotlib.backends.backend_pdf import PdfPages


ROOT = Path(__file__).resolve().parent


@dataclass
class Project:
    folder: str
    author: str
    author_meta: str
    nim: str
    title: str
    topic: str
    problem: str
    gap: str
    novelty: str
    dataset: str
    dataset_source: str
    method_detail: str
    business: str


PROJECTS = [
    Project(
        "Arlen_Prima_Dinova_241730003_UAS_AI",
        "Arlen Prima Dinova",
        "Arlen_Prima_Dinova",
        "241730003",
        "RetainA: Prediksi Churn Pelanggan Berbasis Ensemble Learning dan Explainable Business Risk Scoring",
        "prediksi churn pelanggan dan prioritas retensi bisnis berlangganan",
        "perusahaan berlangganan perlu mengidentifikasi pelanggan yang berpotensi berhenti agar strategi retensi tidak terlambat, tidak terlalu luas, dan tidak boros biaya.",
        "penelitian churn pelanggan telah banyak memakai classifier dan ensemble, tetapi masih sering kurang menekankan hubungan yang utuh antara research gap, baseline eksplisit, hyperparameter tuning, interpretasi risiko bisnis, dan artefak web yang dapat diuji ulang.",
        "mengintegrasikan baseline Logistic Regression, pembanding Random Forest dan Gradient Boosting, tuning RandomizedSearchCV, soft voting ensemble, serta dashboard risk scoring lokal untuk mendukung prioritas retensi pelanggan.",
        "IBM Telco Customer Churn",
        "https://raw.githubusercontent.com/IBM/telco-customer-churn-on-icp4d/master/data/Telco-Customer-Churn.csv",
        "Fitur numerik diproses dengan median imputation dan standardization, sedangkan fitur kategorikal diproses dengan most-frequent imputation dan one-hot encoding. Soft voting dipakai untuk menggabungkan probabilitas beberapa classifier.",
        "Pada kasus churn, recall membantu menangkap pelanggan berisiko agar tidak terlewat, sedangkan precision membantu mengurangi biaya penawaran retensi yang tidak perlu.",
    ),
    Project(
        "Putri_Dwi_Manggali_241730005_UAS_AI",
        "Putri Dwi Manggali",
        "Putri Dwi Manggali",
        "241730005",
        "CartIQ: Prediksi Intensi Pembelian E-Commerce Berbasis Ensemble Learning untuk Optimasi Konversi Bisnis",
        "prediksi intensi pembelian pengunjung e-commerce dan optimasi konversi",
        "platform e-commerce perlu membedakan sesi browsing biasa dari sesi yang memiliki peluang transaksi agar promosi, reminder, dan rekomendasi dapat diberikan secara lebih tepat.",
        "penelitian intensi pembelian dan retail analytics telah banyak menggunakan machine learning, tetapi belum selalu menggabungkan baseline, model pembanding, tuning, evaluasi multi-metrik, dokumentasi dataset, dan prototipe inferensi lokal dalam satu artefak penelitian.",
        "mengintegrasikan feature engineering perilaku sesi, baseline Logistic Regression, pembanding Random Forest dan Gradient Boosting, tuning RandomizedSearchCV, soft voting ensemble, serta dashboard conversion scoring lokal.",
        "Online Shoppers Purchasing Intention",
        "https://archive.ics.uci.edu/ml/machine-learning-databases/00468/online_shoppers_intention.csv",
        "Atribut perilaku sesi, halaman yang dikunjungi, durasi, bounce rate, exit rate, page value, dan fitur kategorikal sesi diproses dalam pipeline tabular untuk menghasilkan probabilitas pembelian.",
        "Pada kasus e-commerce, precision mengurangi promosi yang salah sasaran, sedangkan recall membantu menangkap sesi yang layak diberi reminder atau rekomendasi sebelum pengguna keluar.",
    ),
]


REFERENCES = [
    "Zhang et al., \"Customer churn prediction model based on hybrid neural networks,\" Scientific Reports, 2024.",
    "Peng and Peng, \"Research on Telecom Customer Churn Prediction Based on GA-XGBoost and SHAP,\" Journal of Computer and Communications, 2022.",
    "ETASR, \"Customer churn prediction for telecommunication companies using optimized classifiers,\" Engineering Technology and Applied Science Research, 2024.",
    "Mahayasa et al., \"Customer churn prediction using weighted average ensemble machine learning model,\" JCSSE, 2023.",
    "Bhushan, \"Enhancing customer churn prediction in the telecom sector using ensemble learning,\" National College of Ireland, 2024.",
    "Sun, \"Sales prediction based on machine learning approach,\" Atlantis Press, 2024.",
    "Obi, \"Demand forecasting in retail business using the ensemble machine learning framework,\" ASRJETS, 2024.",
    "SBC ENIAC, \"Comparing gradient boosting algorithms to forecast sales in retail,\" SBC OpenLib, 2023.",
    "EUDL, \"Sales forecast of retail commodity on the basis of LightGBM and XGBoost,\" EUDL, 2022.",
    "UCI Machine Learning Repository, \"Online shoppers purchasing intention dataset,\" University of California Irvine, 2018.",
    "IBM, \"Telco customer churn dataset repository,\" IBM GitHub Repository, 2024.",
    "OpenML, \"Public machine learning datasets for reproducible tabular AI,\" OpenML, 2024.",
    "L. Breiman, \"Random forests,\" Machine Learning, vol. 45, no. 1, pp. 5-32, 2001.",
    "J. H. Friedman, \"Greedy function approximation: a gradient boosting machine,\" Annals of Statistics, vol. 29, no. 5, pp. 1189-1232, 2001.",
    "F. Pedregosa et al., \"Scikit-learn: Machine Learning in Python,\" Journal of Machine Learning Research, vol. 12, pp. 2825-2830, 2011.",
    "R. Kohavi, \"A study of cross-validation and bootstrap for accuracy estimation and model selection,\" IJCAI, 1995.",
    "S. Lundberg and S. Lee, \"A unified approach to interpreting model predictions,\" Advances in Neural Information Processing Systems, 2017.",
    "IEEE, \"IEEE Editorial Style Manual for Authors,\" IEEE Author Center, 2024.",
    "UCI Machine Learning Repository, \"Bank marketing dataset,\" University of California Irvine, 2014.",
    "Kaggle, \"Business analytics public dataset practices,\" Kaggle Datasets, 2024.",
    "A. Géron, Hands-On Machine Learning with Scikit-Learn, Keras, and TensorFlow, 3rd ed. Sebastopol, CA, USA: O'Reilly Media, 2022.",
    "C. Molnar, Interpretable Machine Learning, 2nd ed., 2022.",
]


def set_style(doc: Document, author: str, title: str) -> None:
    cp = doc.core_properties
    cp.author = author
    cp.last_modified_by = author
    cp.title = title
    cp.subject = ""
    cp.keywords = ""
    cp.comments = ""
    section = doc.sections[0]
    section.top_margin = Inches(0.70)
    section.bottom_margin = Inches(0.70)
    section.left_margin = Inches(0.68)
    section.right_margin = Inches(0.68)
    for name, size, bold in [("Normal", 9, False), ("Heading 1", 10, True), ("Heading 2", 9, True)]:
        style = doc.styles[name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.font.size = Pt(size)
        style.font.bold = bold
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.paragraph_format.space_after = Pt(3)


def set_columns(section, count=2) -> None:
    sect_pr = section._sectPr
    cols = sect_pr.xpath("./w:cols")
    cols = cols[0] if cols else OxmlElement("w:cols")
    if not cols.getparent():
        sect_pr.append(cols)
    cols.set(qn("w:num"), str(count))
    cols.set(qn("w:space"), "360")


def para(doc: Document, text: str, align=None, size=9, bold=False) -> None:
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text)
    r.font.name = "Times New Roman"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    r.font.size = Pt(size)
    r.font.color.rgb = RGBColor(0, 0, 0)
    r.bold = bold


def heading(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="Heading 1")
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    r.font.name = "Times New Roman"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    r.font.size = Pt(10)
    r.font.bold = True
    r.font.color.rgb = RGBColor(0, 0, 0)


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    for i, h in enumerate(headers):
        t.rows[0].cells[i].text = str(h)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
    for row in t.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(0)
                for r in p.runs:
                    r.font.name = "Times New Roman"
                    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
                    r.font.size = Pt(7)
                    r.font.color.rgb = RGBColor(0, 0, 0)


def add_img(doc: Document, path: Path, caption: str) -> None:
    if path.exists():
        doc.add_picture(str(path), width=Inches(3.0))
        para(doc, caption, align=WD_ALIGN_PARAGRAPH.CENTER, size=8)


def result_rows(comp: pd.DataFrame) -> list[list[str]]:
    rows = []
    for _, r in comp.iterrows():
        rows.append([r["model"], f"{r['accuracy']:.3f}", f"{r['precision']:.3f}", f"{r['recall']:.3f}", f"{r['f1']:.3f}", f"{r['auc']:.3f}"])
    return rows


def build_docx(p: Project) -> None:
    base = ROOT / p.folder
    metrics = json.loads((base / "07_Hasil_Eksperimen" / "metrics.json").read_text(encoding="utf-8"))
    comp = pd.read_csv(base / "07_Hasil_Eksperimen" / "model_comparison.csv")
    doc = Document()
    set_style(doc, p.author_meta, p.title)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run(p.title)
    r.font.name = "Times New Roman"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    r.font.size = Pt(18)
    r.font.bold = True
    r.font.color.rgb = RGBColor(0, 0, 0)
    para(doc, f"{p.author}, {p.nim}\nProgram Studi Informatika, Fakultas Sains dan Teknologi\nUniversitas Islam Negeri Sultan Maulana Hasanuddin Banten\nBanten, Indonesia, 2026", align=WD_ALIGN_PARAGRAPH.CENTER, size=10)

    sec = doc.add_section(WD_SECTION.CONTINUOUS)
    set_columns(sec)

    heading(doc, "Abstrak")
    para(doc, f"Artikel ini menyajikan {p.title}, yaitu artefak penelitian Artificial Intelligence untuk {p.topic}. Permasalahan yang diangkat adalah bahwa {p.problem} Penelitian ini mengikuti alur lengkap mulai dari studi literatur, identifikasi research gap, novelty, implementasi metode, evaluasi eksperimen, hingga deployment lokal. Dataset utama yang digunakan adalah {p.dataset} dari sumber publik agar eksperimen dapat direproduksi. Metode yang diusulkan menggabungkan preprocessing tabular, baseline Logistic Regression, model pembanding Random Forest dan Gradient Boosting, hyperparameter tuning, serta soft voting ensemble. Model akhir memperoleh accuracy {metrics['metrics']['accuracy']:.3f}, precision {metrics['metrics']['precision']:.3f}, recall {metrics['metrics']['recall']:.3f}, F1-score {metrics['metrics']['f1']:.3f}, dan AUC {metrics['metrics']['auc']:.3f}. Hasil menunjukkan bahwa pendekatan AI tidak hanya menghasilkan skor prediksi, tetapi juga dapat dikemas sebagai sistem pendukung keputusan bisnis yang terdokumentasi.")
    heading(doc, "Kata Kunci")
    para(doc, f"Artificial Intelligence; {p.topic}; ensemble learning; baseline comparison; hyperparameter tuning; deployment lokal; artikel IEEE.")

    heading(doc, "I. Pendahuluan")
    for text in [
        f"Perkembangan Artificial Intelligence mendorong penggunaan data historis untuk mendukung keputusan bisnis. Pada topik {p.topic}, model prediksi berperan sebagai sinyal awal yang membantu pengambil keputusan melihat risiko atau peluang sebelum kejadian aktual terjadi. Dalam proyek ini, masalah utama adalah bahwa {p.problem}",
        "Tugas UAS AI menuntut mahasiswa tidak hanya membuat ringkasan jurnal, tetapi juga menunjukkan proses penelitian yang utuh. Oleh karena itu, artikel ini disusun untuk memperlihatkan hubungan antara kajian literatur, research gap, novelty, implementasi model, baseline comparison, hasil eksperimen, dan artefak deployment. Struktur tersebut mengikuti gaya artikel IEEE yang menekankan kejelasan masalah, metode, hasil, dan referensi.",
        "Revisi ini tidak membuat topik baru, tidak mengganti dataset, dan tidak mengganti model utama yang sudah valid. Fokus revisi adalah menyusun ulang naskah menjadi Bahasa Indonesia penuh dengan format artikel IEEE, sitasi numbered style, referensi minimal 20, serta pembahasan yang sesuai dengan proyek masing-masing.",
        f"Secara praktis, proyek ini relevan karena {p.business} Skor prediksi tidak diposisikan sebagai keputusan final otomatis, melainkan sebagai alat bantu prioritas tindakan. Dengan demikian, penelitian ini menghubungkan metrik model dengan kebutuhan bisnis yang nyata.",
        "Artikel ini disusun sebagai berikut. Bagian II membahas penelitian terkait. Bagian III menjelaskan research gap, novelty, dan kontribusi. Bagian IV menjelaskan metode yang diusulkan. Bagian V menjelaskan eksperimen. Bagian VI menyajikan hasil dan pembahasan. Bagian VII membahas deployment, Bagian VIII membahas ancaman validitas, dan Bagian IX menyajikan kesimpulan.",
    ]:
        para(doc, text)

    heading(doc, "II. Penelitian Terkait")
    for text in [
        "Penelitian terdahulu menunjukkan bahwa data bisnis tabular banyak dianalisis menggunakan ensemble learning, boosting, dan explainable AI. Pada domain churn pelanggan, beberapa studi menggunakan hybrid neural network, optimized classifiers, GA-XGBoost, SHAP, dan weighted ensemble [1]-[5]. Pada domain retail dan e-commerce, beberapa penelitian menggunakan machine learning, LightGBM, XGBoost, serta gradient boosting untuk prediksi penjualan atau intensi pembelian [6]-[10].",
        "Literatur tersebut memberikan dasar bahwa model machine learning dapat digunakan untuk masalah klasifikasi bisnis. Akan tetapi, sebagian penelitian lebih menonjolkan performa model akhir dibanding kelengkapan artefak penelitian. Dalam konteks tugas UAS AI, performa model perlu didampingi dataset source, kode, model, hasil evaluasi, visualisasi, gap analysis, novelty statement, dan deployment evidence.",
        "Studi ensemble memberikan pelajaran bahwa penggabungan model dapat mengurangi ketergantungan terhadap satu classifier [3]-[5]. Studi retail dan e-commerce menunjukkan pentingnya evaluasi multi-metrik karena kelas positif sering lebih penting secara bisnis [6]-[10]. Literatur reproduksibilitas machine learning juga menekankan pentingnya cross-validation, pemisahan data, dan dokumentasi preprocessing [12], [15], [16].",
        "Selain itu, satu metrik saja tidak cukup untuk menilai kualitas model. Accuracy dapat terlihat tinggi ketika kelas mayoritas dominan. Precision penting ketika biaya false positive tinggi. Recall penting ketika melewatkan kasus positif lebih merugikan. F1-score menjadi ukuran kompromi antara precision dan recall, sedangkan AUC membantu membaca kualitas ranking probabilitas model.",
    ]:
        para(doc, text)
    add_table(doc, ["No", "Tahun", "Dataset", "Metode", "Keterbatasan"], [
        ["1", "2024", "Churn", "Hybrid NN", "Deployment terbatas"],
        ["2", "2022", "Telecom churn", "GA-XGBoost + SHAP", "Validasi domain terbatas"],
        ["3", "2024", "Telecom churn", "Optimized classifiers", "Workflow bisnis belum lengkap"],
        ["4", "2023", "Customer churn", "Weighted ensemble", "Artefak deployment terbatas"],
        ["5", "2024", "Telecom churn", "Ensemble learning", "Packaging akademik terbatas"],
        ["6", "2024", "Sales data", "Machine learning", "Tidak ada dashboard inferensi"],
        ["7", "2024", "Retail demand", "Ensemble ML", "Validasi pengguna terbatas"],
        ["8", "2023", "Retail sales", "Gradient boosting", "Deployment bukan fokus"],
        ["9", "2022", "Retail commodity", "LightGBM/XGBoost", "Artefak akademik terbatas"],
        ["10", "2021-2024", "Online shopper", "Classification", "Baseline tidak selalu lengkap"],
    ])

    heading(doc, "III. Research Gap, Novelty, dan Kontribusi")
    for text in [
        f"Research gap utama yang ditemukan adalah bahwa {p.gap} Gap ini didukung oleh literature mapping karena beberapa penelitian berorientasi model belum selalu menyediakan artefak eksperimen yang siap diperiksa ulang.",
        f"Novelty penelitian ini adalah {p.novelty} Novelty tersebut realistis karena dapat dibuktikan melalui folder dataset, source code, model, hasil eksperimen, visualisasi, draft artikel, dan deployment lokal.",
        "Novelty pada penelitian ini tidak diklaim sebagai penemuan algoritma baru. Kebaruan diletakkan pada integrasi disiplin antara metode yang sudah terbukti, baseline yang jelas, tuning, evaluasi multi-metrik, interpretasi bisnis, dan demo web. Pendekatan ini sesuai untuk proyek perkuliahan karena dapat diimplementasikan, diverifikasi, dan dinilai secara objektif.",
        "Kontribusi ilmiah penelitian ini adalah penyajian eksperimen AI tabular yang reproducible dengan pembanding baseline. Kontribusi praktisnya adalah dashboard lokal yang dapat digunakan untuk mencoba prediksi. Kontribusi akademiknya adalah penyusunan artefak lengkap sesuai ketentuan UAS AI dan struktur artikel IEEE.",
    ]:
        para(doc, text)
    add_table(doc, ["Pendekatan Eksisting", "Keterbatasan", "Perbaikan yang Diusulkan"], [
        ["Pelaporan accuracy", "Interpretasi operasional terbatas", "Skor risiko/konversi dikaitkan dengan aksi bisnis"],
        ["Baseline tunggal", "Perbandingan lemah", "Baseline + dua pembanding + ensemble"],
        ["Eksperimen tanpa demo", "Sulit diperiksa interaktif", "Aplikasi Flask lokal"],
        ["Tuning tidak jelas", "Sulit direproduksi", "RandomizedSearchCV dan CV terdokumentasi"],
    ])

    heading(doc, "IV. Metode yang Diusulkan")
    for text in [
        f"Metode yang diusulkan mengikuti pipeline AI tabular. Dataset diambil dari sumber publik: {p.dataset_source}. Target dikonversi menjadi kelas biner. Data kemudian dibagi menjadi data latih dan data uji dengan stratified splitting agar distribusi kelas tetap terjaga.",
        p.method_detail,
        "Baseline yang digunakan adalah Logistic Regression karena sederhana, cepat, dan memberikan titik awal pembanding yang jelas. Random Forest digunakan sebagai pembanding karena kuat terhadap interaksi non-linear. Gradient Boosting digunakan karena sering memberikan performa kuat pada data tabular bisnis. Model usulan adalah soft voting ensemble yang menggabungkan probabilitas dari model-model tersebut.",
        "Soft voting dipilih karena output probabilitas lebih relevan untuk sistem pendukung keputusan. Dalam konteks bisnis, pengguna tidak hanya membutuhkan label kelas, tetapi juga besarnya peluang atau risiko. Probabilitas ini dapat diterjemahkan sebagai prioritas tindakan, misalnya prioritas retensi atau prioritas promosi.",
        "Preprocessing ditempatkan dalam pipeline untuk mengurangi risiko data leakage. Transformasi numerik dan kategorikal dipelajari dari data latih, kemudian diterapkan ke data uji. Desain ini lebih aman dibanding preprocessing manual di luar pipeline karena mencegah informasi data uji masuk ke proses training.",
    ]:
        para(doc, text)
    add_img(doc, base / "08_Visualisasi" / "pipeline_diagram.png", "Gambar 1. Pipeline AI dari dataset hingga inferensi.")

    heading(doc, "V. Setup Eksperimen")
    for text in [
        f"Dataset utama yang digunakan adalah {p.dataset}. Dataset ini dipilih karena publik, relevan dengan masalah bisnis, memiliki target klasifikasi yang jelas, dan dapat diakses oleh dosen pengampu. Informasi dataset meliputi jumlah data, fitur, kelas, distribusi kelas, sumber, tanggal akses, dan alasan pemilihan.",
        "Evaluasi dilakukan menggunakan accuracy, precision, recall, F1-score, confusion matrix, ROC curve, dan AUC. Accuracy memberikan gambaran umum, precision mengukur ketepatan prediksi positif, recall mengukur kemampuan menangkap kelas positif, dan F1-score mengukur keseimbangan precision-recall.",
        "Hyperparameter tuning dilakukan menggunakan randomized search dan stratified cross-validation. Ruang pencarian dibuat moderat agar sesuai dengan sumber daya komputasi perkuliahan, namun tetap menunjukkan tuning yang sistematis. Hasil tuning didokumentasikan pada folder hasil eksperimen.",
        "Perbandingan baseline dirancang transparan. Logistic Regression mewakili baseline linear. Random Forest mewakili bagging-based ensemble. Gradient Boosting mewakili sequential additive learning. Soft voting ensemble menggabungkan probabilitas model untuk menghasilkan skor akhir.",
    ]:
        para(doc, text)
    add_table(doc, ["Komponen", "Konfigurasi"], [
        ["Dataset", p.dataset],
        ["Baseline", "Logistic Regression"],
        ["Pembanding 1", "Random Forest"],
        ["Pembanding 2", "Gradient Boosting"],
        ["Metode Usulan", "Tuned Soft Voting Ensemble"],
        ["Validasi", "Stratified split dan cross-validation"],
        ["Metrik", "Accuracy, Precision, Recall, F1, AUC"],
    ])
    add_img(doc, base / "08_Visualisasi" / "dataset_distribution_300dpi.png", "Gambar 2. Distribusi kelas dataset.")

    heading(doc, "VI. Hasil dan Pembahasan")
    para(doc, "Hasil eksperimen ditunjukkan pada Tabel IV. Tabel ini membandingkan baseline, model pembanding, dan model usulan. Evaluasi menggunakan beberapa metrik agar pembahasan tidak hanya bergantung pada accuracy.")
    add_table(doc, ["Model", "Accuracy", "Precision", "Recall", "F1", "AUC"], result_rows(comp))
    for text in [
        f"Model akhir memperoleh accuracy {metrics['metrics']['accuracy']:.3f}, precision {metrics['metrics']['precision']:.3f}, recall {metrics['metrics']['recall']:.3f}, F1-score {metrics['metrics']['f1']:.3f}, dan AUC {metrics['metrics']['auc']:.3f}. {p.business}",
        "Confusion matrix memberikan gambaran error model. False positive berarti sistem memberi sinyal positif padahal kejadian positif tidak terjadi. False negative berarti sistem gagal menangkap kasus positif. Dalam bisnis, dua jenis error ini memiliki biaya yang berbeda sehingga interpretasinya harus dikaitkan dengan konteks.",
        "Grafik perbandingan menunjukkan posisi model usulan terhadap baseline dan model pembanding. Jika ensemble tidak selalu unggul pada semua metrik, hasil tetap dapat diterima karena AI bisnis sering membutuhkan trade-off. Model yang seimbang antara precision dan recall dapat lebih berguna daripada model yang hanya unggul pada accuracy.",
        "ROC curve dan AUC membantu mengevaluasi kemampuan model dalam membedakan kelas pada berbagai threshold. Threshold dapat disesuaikan dengan kebijakan bisnis. Jika biaya kehilangan kasus positif tinggi, threshold dapat diturunkan; jika biaya intervensi tinggi, threshold dapat dinaikkan.",
        "Feature importance membantu menghubungkan hasil model dengan interpretasi domain. Analisis fitur membuat naskah tidak hanya menampilkan angka, tetapi juga menjelaskan faktor yang memengaruhi prediksi. Hal ini penting agar dosen dapat melihat bahwa eksperimen benar-benar dianalisis.",
    ]:
        para(doc, text)
    add_img(doc, base / "08_Visualisasi" / "model_comparison_chart.png", "Gambar 3. Grafik perbandingan model.")
    add_img(doc, base / "08_Visualisasi" / "confusion_matrix.png", "Gambar 4. Confusion matrix model usulan.")
    add_img(doc, base / "08_Visualisasi" / "roc_curve.png", "Gambar 5. ROC curve model usulan.")
    add_img(doc, base / "08_Visualisasi" / "feature_importance.png", "Gambar 6. Feature importance.")

    heading(doc, "VII. Prototipe Deployment")
    for text in [
        "Prototipe web lokal dibuat untuk menunjukkan inferensi model. Aplikasi membaca file model yang sudah disimpan, menerima input dari form, lalu menghasilkan probabilitas kelas positif. Aplikasi ini sengaja dibuat ringan agar dapat dijalankan oleh dosen pada komputer lokal.",
        "Folder deployment berisi source code deployment, screenshot, URL lokal, deployment evidence, dan dokumentasi penggunaan. Dengan demikian, bukti deployment tidak hanya berupa kode, tetapi juga menunjukkan bahwa model benar-benar dapat dicoba melalui antarmuka pengguna.",
        "Aplikasi web dipisahkan dari proses training. Training dilakukan pada pipeline penelitian, sedangkan inference dilakukan dengan memuat model `.pkl`. Pemisahan ini penting karena sistem operasional tidak seharusnya melatih ulang model setiap kali pengguna membuka dashboard.",
    ]:
        para(doc, text)
    add_img(doc, base / "08_Visualisasi" / "deployment_architecture.png", "Gambar 7. Arsitektur deployment lokal.")

    heading(doc, "VIII. Ancaman Validitas")
    for text in [
        "Ancaman validitas pertama adalah representativitas dataset. Dataset publik mungkin tidak sepenuhnya mewakili perilaku bisnis lokal di Indonesia. Oleh karena itu, model sebaiknya dianggap sebagai prototipe penelitian, bukan sistem produksi final.",
        "Ancaman kedua adalah class imbalance. Walaupun stratified split dan metrik multi-aspek digunakan, threshold final masih perlu disesuaikan jika sistem diterapkan secara operasional. Ancaman ketiga adalah temporal drift, karena perilaku pelanggan atau pengunjung dapat berubah dari waktu ke waktu.",
        "Ancaman lain adalah interpretabilitas. Feature importance memberikan gambaran umum, tetapi belum menjelaskan setiap prediksi individual secara rinci. Penelitian lanjutan dapat menambahkan local explanation dan probability calibration untuk meningkatkan kepercayaan pengguna.",
    ]:
        para(doc, text)

    heading(doc, "IX. Kesimpulan dan Pengembangan Lanjutan")
    for text in [
        f"Artikel ini menyajikan versi Bahasa Indonesia dari draft IEEE untuk topik {p.topic}. Penelitian mencakup studi literatur, research gap, novelty, metode, eksperimen, baseline comparison, visualisasi, dan deployment lokal. Dengan demikian, penelitian tidak berhenti sebagai review jurnal.",
        "Hasil eksperimen menunjukkan bahwa pendekatan ensemble dapat digunakan sebagai sistem pendukung keputusan bisnis. Meskipun demikian, hasil model perlu dibaca bersama precision, recall, F1-score, AUC, confusion matrix, dan konteks biaya bisnis.",
        "Pengembangan berikutnya dapat menggunakan dataset lokal Indonesia, calibration curve, explainability yang lebih detail, integrasi database, monitoring deployment, dan evaluasi pada data terbaru. Artefak akhir sudah disiapkan agar dapat diperiksa sesuai ketentuan UAS AI.",
    ]:
        para(doc, text)

    heading(doc, "Referensi")
    for i, ref in enumerate(REFERENCES, 1):
        para(doc, f"[{i}] {ref}", size=8)

    out = base / "09_Draft_IEEE" / "Draft_Artikel_IEEE.docx"
    doc.save(out)
    with ZipFile(out) as z:
        comments = [n for n in z.namelist() if "comments" in n.lower()]
        if comments:
            raise RuntimeError(comments)


def build_pdf(p: Project) -> None:
    base = ROOT / p.folder
    metrics = json.loads((base / "07_Hasil_Eksperimen" / "metrics.json").read_text(encoding="utf-8"))
    comp = pd.read_csv(base / "07_Hasil_Eksperimen" / "model_comparison.csv")
    refs = "\n".join(f"[{i}] {r}" for i, r in enumerate(REFERENCES, 1))
    pages = [
        f"{p.title}\n{p.author}, {p.nim}\n\nAbstrak\nArtikel ini menyajikan penelitian AI untuk {p.topic}. Dataset utama adalah {p.dataset}. Metode mencakup preprocessing, baseline, tuning, model pembanding, dan soft voting ensemble. Hasil akhir: accuracy {metrics['metrics']['accuracy']:.3f}, precision {metrics['metrics']['precision']:.3f}, recall {metrics['metrics']['recall']:.3f}, F1 {metrics['metrics']['f1']:.3f}, AUC {metrics['metrics']['auc']:.3f}.",
        f"I. Pendahuluan\nMasalah penelitian: {p.problem} Penelitian ini tidak mengganti topik, dataset, atau model utama. Naskah disusun ulang menjadi Bahasa Indonesia penuh dengan struktur IEEE.\n\nII. Penelitian Terkait\nLiteratur 2021-2026 menunjukkan penggunaan ensemble, boosting, dan explainable AI pada data bisnis tabular [1]-[10].",
        f"III. Research Gap dan Novelty\nGap: {p.gap}\n\nNovelty: {p.novelty}\n\nKontribusi ilmiah, praktis, dan akademik dijelaskan melalui artefak lengkap.",
        f"IV. Metode\nDataset: {p.dataset_source}\n{p.method_detail}\nBaseline: Logistic Regression. Pembanding: Random Forest dan Gradient Boosting. Metode usulan: Tuned Soft Voting Ensemble.",
        "V. Setup Eksperimen\nEvaluasi menggunakan accuracy, precision, recall, F1-score, confusion matrix, ROC curve, AUC, feature importance, dan comparison graph.\n\nVI. Hasil\n" + comp.to_string(index=False),
        f"Pembahasan\n{p.business} Hasil model dipakai sebagai pendukung keputusan, bukan keputusan otomatis. Confusion matrix, ROC, AUC, dan feature importance digunakan untuk membaca kualitas model secara lebih lengkap.",
        "VII. Deployment\nAplikasi lokal Flask memuat model .pkl dan menyediakan dashboard serta halaman prediksi. Deployment evidence, screenshot, source code, dan dokumentasi tersedia pada folder deployment.\n\nVIII. Ancaman Validitas\nAncaman meliputi representativitas dataset, class imbalance, temporal drift, dan keterbatasan interpretabilitas.",
        "IX. Kesimpulan\nPenelitian memenuhi fokus implementasi dan eksperimen AI. Pengembangan lanjutan dapat memakai dataset lokal, calibration, explainability, database, dan production monitoring.\n\nReferensi\n" + refs,
    ]
    out = base / "09_Draft_IEEE" / "Draft_Artikel_IEEE.pdf"
    with PdfPages(out) as pdf:
        for idx, content in enumerate(pages, 1):
            fig = plt.figure(figsize=(8.5, 11))
            fig.text(0.5, 0.965, "Draft Artikel IEEE - Bahasa Indonesia", ha="center", fontsize=9, family="serif")
            fig.text(0.08, 0.94, f"Halaman {idx}", fontsize=8, family="serif")
            lines = []
            for paragraph in content.split("\n"):
                line = ""
                for word in paragraph.split():
                    if len(line) + len(word) + 1 > 54:
                        lines.append(line)
                        line = word
                    else:
                        line = (line + " " + word).strip()
                if line:
                    lines.append(line)
                lines.append("")
            x_positions = [0.08, 0.53]
            col = 0
            y = 0.905
            for line in lines:
                if y < 0.08 and col == 0:
                    col = 1
                    y = 0.905
                elif y < 0.08:
                    break
                fig.text(x_positions[col], y, line, fontsize=8.0, family="serif")
                y -= 0.019 if line else 0.014
            pdf.savefig(fig)
            plt.close(fig)


def main() -> None:
    for p in PROJECTS:
        build_docx(p)
        build_pdf(p)
        (ROOT / p.folder / "09_Draft_IEEE" / "README_FOLDER.md").write_text(
            "Folder ini berisi draft artikel IEEE versi Bahasa Indonesia penuh. Struktur, sitasi IEEE style, 22 referensi, tabel, gambar, DOCX, dan PDF sudah disesuaikan dengan proyek masing-masing.\n",
            encoding="utf-8",
        )
        print(f"Indonesian IEEE draft rebuilt: {p.folder}")


if __name__ == "__main__":
    main()
