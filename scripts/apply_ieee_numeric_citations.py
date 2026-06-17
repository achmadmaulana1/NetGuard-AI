from __future__ import annotations

import subprocess
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import RGBColor, Pt


ROOT = Path(__file__).resolve().parents[1]
DOCX = ROOT / "Achmad_Maulana_241730016_UAS_AI" / "09_Draft_IEEE" / "Draft_Artikel_IEEE.docx"
PDF = ROOT / "Achmad_Maulana_241730016_UAS_AI" / "09_Draft_IEEE" / "Draft_Artikel_IEEE.pdf"
STUDENT = "Achmad Maulana"
FONT = "Times New Roman"


REPLACEMENTS = {
    "Research on network intrusion detection commonly uses public datasets such as CICIDS2017, CICIDS2018, UNSW-NB15, and NSL-KDD.":
        "Research on network intrusion detection commonly uses public datasets such as CICIDS2017, CICIDS2018, UNSW-NB15, and NSL-KDD [1], [2], [11].",
    "Recent work shows that classical machine learning remains competitive for structured flow features, especially tree-based classifiers.":
        "Recent work shows that classical machine learning remains competitive for structured flow features, especially tree-based classifiers [3], [4], [10].",
    "Several studies also warn that benchmark datasets must be used carefully.":
        "Several studies also warn that benchmark datasets must be used carefully [1], [2].",
    "Dataset leakage, duplicate records, improper train-test separation, and unbalanced class distribution can produce inflated results.":
        "Dataset leakage, duplicate records, improper train-test separation, and unbalanced class distribution can produce inflated results [1], [2].",
    "The literature gap identified in this study is not merely model accuracy.":
        "The literature gap identified in this study is not merely model accuracy [3]-[10].",
    "Many IDS papers present high metrics but provide limited operational packaging for beginners.":
        "Many IDS papers present high metrics but provide limited operational packaging for beginners [4], [7], [8].",
    "The mapping indicates that the selected studies already cover dataset quality, machine learning comparison, feature importance, ensemble learning, deep learning, and robustness.":
        "The mapping indicates that the selected studies already cover dataset quality, machine learning comparison, feature importance, ensemble learning, deep learning, and robustness [1]-[10].",
    "Logistic Regression is used as the first baseline because it represents a simple linear classifier.":
        "Logistic Regression is used as the first baseline because it represents a simple linear classifier [12].",
    "Decision Tree is selected as the second baseline and practical proposed model because it can model non-linear decision rules without requiring heavy computation.":
        "Decision Tree is selected as the second baseline and practical proposed model because it can model non-linear decision rules without requiring heavy computation [10], [12].",
    "Random Forest is selected as the third model because it represents an ensemble of many decision trees.":
        "Random Forest is selected as the third model because it represents an ensemble of many decision trees [5], [10], [12].",
    "The dashboard is implemented using Flask, Bootstrap, and Chart.js.":
        "The dashboard is implemented using Flask, Bootstrap, and Chart.js [13], [14].",
    "The experiment uses a representative subset of CICIDS2017 Friday-WorkingHours-Afternoon-DDos.":
        "The experiment uses a representative subset of CICIDS2017 Friday-WorkingHours-Afternoon-DDos [11].",
}


def style_all_runs(doc: Document) -> None:
    for style_name in ["Normal", "Heading 1", "Heading 2", "Heading 3", "Title", "Subtitle"]:
        if style_name in doc.styles:
            style = doc.styles[style_name]
            style.font.name = FONT
            style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
            style.font.color.rgb = RGBColor(0, 0, 0)
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            run.font.name = FONT
            run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
            run.font.color.rgb = RGBColor(0, 0, 0)
            if run.font.size is None:
                run.font.size = Pt(9)


def export_pdf() -> None:
    ps = f"""
$ErrorActionPreference = 'Stop'
$word = New-Object -ComObject Word.Application
$word.Visible = $false
$doc = $word.Documents.Open('{DOCX}')
$doc.Repaginate()
$doc.SaveAs([ref]'{PDF}', [ref]17)
$doc.Close($false)
$word.Quit()
"""
    subprocess.run(["powershell", "-NoProfile", "-Command", ps], check=True, timeout=120)


def main() -> None:
    doc = Document(DOCX)
    props = doc.core_properties
    props.author = STUDENT
    props.last_modified_by = STUDENT
    props.comments = ""
    props.subject = ""
    props.keywords = ""

    changed = 0
    for paragraph in doc.paragraphs:
        text = paragraph.text
        for old, new in REPLACEMENTS.items():
            if old in text and new not in text:
                text = text.replace(old, new)
        if text != paragraph.text:
            paragraph.clear()
            run = paragraph.add_run(text)
            run.font.name = FONT
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0, 0, 0)
            changed += 1

    style_all_runs(doc)
    doc.save(DOCX)
    export_pdf()
    print(f"Applied IEEE numeric citations to {changed} paragraphs.")


if __name__ == "__main__":
    main()
