from __future__ import annotations

import re
import shutil
import subprocess
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import RGBColor, Pt


ROOT = Path(__file__).resolve().parents[1]
DIR = ROOT / "Achmad_Maulana_241730016_UAS_AI" / "09_Draft_IEEE"
SOURCE = DIR / "Draft_Artikel_IEEE.docx"
EN_DOCX = DIR / "Draft_Artikel_IEEE_EN.docx"
EN_PDF = DIR / "Draft_Artikel_IEEE_EN.pdf"
ID_DOCX = DIR / "Draft_Artikel_IEEE_ID.docx"
ID_PDF = DIR / "Draft_Artikel_IEEE_ID.pdf"
STUDENT = "Achmad Maulana"
FONT = "Times New Roman"


PARAGRAPH_MAP = {
    "NetGuard AI: Lightweight Machine Learning-Based Network Anomaly Detection and Risk Monitoring Dashboard":
        "NetGuard AI: Dashboard Ringan Berbasis Machine Learning untuk Deteksi Anomali Jaringan dan Pemantauan Risiko",
    "Abstract--Network anomaly detection is important for educational institutions that operate computer laboratories, local servers, and internet gateways with limited monitoring budgets. This paper presents NetGuard AI, a low-budget machine learning system for detecting anomalous network traffic and visualizing operational risk through a lightweight dashboard. The study uses a representative subset of CICIDS2017 Friday Afternoon DDoS traffic. The preprocessing stage normalizes column names, converts labels into binary classes, handles missing and infinity values, keeps numeric features, removes duplicates, and exports processed data for reproducible experiments. Logistic Regression, Decision Tree, and Random Forest are compared using accuracy, precision, recall, F1-score, and confusion matrix. The best model is selected using F1-score as the primary criterion, recall as the secondary criterion, and accuracy as the tertiary criterion. The experiment uses 19,935 processed records with an 80:20 train-test split and random_state=42. Decision Tree achieved accuracy 0.9992, precision 0.9995, recall 0.9990, and F1-score 0.9992. The system exports metrics, model comparison, confusion matrix, prediction results, and research summary files. The novelty is the integration of machine learning anomaly detection, automated risk classification, recommended actions, and a student-friendly web dashboard that can be reproduced on a normal laptop without paid APIs.":
        "Abstrak--Deteksi anomali jaringan penting bagi institusi pendidikan yang mengelola laboratorium komputer, server lokal, dan gerbang internet dengan anggaran monitoring terbatas. Artikel ini menyajikan NetGuard AI, yaitu sistem machine learning berbiaya rendah untuk mendeteksi trafik jaringan anomali dan memvisualisasikan risiko operasional melalui dashboard ringan. Penelitian ini menggunakan subset representatif dari trafik DDoS CICIDS2017 Friday Afternoon. Tahap preprocessing menormalkan nama kolom, mengubah label menjadi kelas biner, menangani nilai hilang dan infinity, mempertahankan fitur numerik, menghapus duplikasi, serta mengekspor data terproses untuk eksperimen yang dapat direproduksi. Logistic Regression, Decision Tree, dan Random Forest dibandingkan menggunakan accuracy, precision, recall, F1-score, dan confusion matrix. Model terbaik dipilih menggunakan F1-score sebagai kriteria utama, recall sebagai kriteria kedua, dan accuracy sebagai kriteria ketiga. Eksperimen menggunakan 19.935 record terproses dengan train-test split 80:20 dan random_state=42. Decision Tree memperoleh accuracy 0,9992, precision 0,9995, recall 0,9990, dan F1-score 0,9992. Sistem mengekspor metrik, perbandingan model, confusion matrix, hasil prediksi, dan ringkasan penelitian. Kebaruan penelitian ini adalah integrasi deteksi anomali berbasis machine learning, klasifikasi risiko otomatis, rekomendasi tindakan, dan dashboard web ramah mahasiswa yang dapat direproduksi pada laptop biasa tanpa API berbayar.",
    "Keywords--network anomaly detection, CICIDS2017, machine learning, intrusion detection system, risk dashboard":
        "Kata kunci--deteksi anomali jaringan, CICIDS2017, machine learning, intrusion detection system, dashboard risiko",
    "I. Introduction": "I. Pendahuluan",
    "II. Related Works": "II. Tinjauan Pustaka",
    "III. Research Gap and Novelty": "III. Research Gap dan Novelty",
    "IV. Proposed Method": "IV. Metode yang Diusulkan",
    "V. Experimental Setup": "V. Skenario Eksperimen",
    "VI. Results and Discussion": "VI. Hasil dan Pembahasan",
    "VII. Detailed Analysis of Baseline Selection": "VII. Analisis Detail Pemilihan Baseline",
    "VIII. Detailed Metric Interpretation": "VIII. Interpretasi Detail Metrik Evaluasi",
    "IX. Error Analysis": "IX. Analisis Error",
    "VII. Dashboard Implementation": "X. Implementasi Dashboard",
    "X. Implementation Details": "XI. Detail Implementasi",
    "XI. Threats to Validity": "XII. Ancaman terhadap Validitas",
    "XII. Reproducibility Package": "XIII. Paket Reproducibility",
    "XIII. Practical Deployment Scenario": "XIV. Skenario Deployment Praktis",
    "XIV. Academic and Educational Contribution": "XV. Kontribusi Akademik dan Edukatif",
    "XV. Future Work": "XVI. Pengembangan Selanjutnya",
    "XVI. Conclusion": "XVII. Kesimpulan",
    "References": "Referensi",
}


PHRASE_MAP = {
    "Artificial Intelligence has changed the way network traffic can be analyzed. In traditional computer networks, administrators inspect logs, firewall alerts, and bandwidth graphs manually. This approach is useful but can be slow when the number of packets, flows, or users increases. Educational institutions often face an additional problem: they need monitoring capability, but they do not always have enterprise-grade security infrastructure.": "Artificial Intelligence telah mengubah cara trafik jaringan dianalisis. Pada jaringan komputer tradisional, administrator memeriksa log, peringatan firewall, dan grafik bandwidth secara manual. Pendekatan tersebut berguna, tetapi dapat menjadi lambat ketika jumlah paket, flow, atau pengguna meningkat. Institusi pendidikan menghadapi persoalan tambahan karena membutuhkan kemampuan monitoring, namun tidak selalu memiliki infrastruktur keamanan tingkat enterprise.",
    "The NetGuard AI project is designed for this context. It connects practical networking knowledge with supervised machine learning so that students can understand how network flow features are transformed into prediction results. The project is also designed as a research artifact: it includes dataset documentation, preprocessing scripts, training scripts, evaluation outputs, dashboard visualization, and a draft article that can be improved for undergraduate research.": "Proyek NetGuard AI dirancang untuk konteks tersebut. Sistem ini menghubungkan pengetahuan praktis jaringan dengan supervised machine learning sehingga mahasiswa dapat memahami bagaimana fitur network flow diubah menjadi hasil prediksi. Proyek ini juga dirancang sebagai artefak penelitian yang mencakup dokumentasi dataset, script preprocessing, script training, output evaluasi, visualisasi dashboard, dan draft artikel yang dapat dikembangkan untuk penelitian sarjana.",
    "The main research question is: how can a low-budget machine learning pipeline detect anomaly traffic and present the result in a form that is useful for beginner network administrators? This question is intentionally practical. The project does not only report numerical performance; it also converts anomaly ratio into Low, Medium, or High risk and provides recommended operational actions.": "Rumusan masalah utama adalah bagaimana pipeline machine learning berbiaya rendah dapat mendeteksi trafik anomali dan menyajikan hasilnya dalam bentuk yang berguna bagi administrator jaringan pemula. Rumusan ini sengaja bersifat praktis. Proyek tidak hanya melaporkan performa numerik, tetapi juga mengubah anomaly ratio menjadi risiko Low, Medium, atau High serta memberikan rekomendasi tindakan operasional.",
    "The proposed workflow consists of dataset acquisition, preprocessing, model training, model evaluation, prediction export, dashboard visualization, and research packaging.": "Alur kerja yang diusulkan terdiri dari akuisisi dataset, preprocessing, training model, evaluasi model, ekspor prediksi, visualisasi dashboard, dan pengemasan artefak penelitian.",
    "The dashboard is implemented using Flask, Bootstrap, and Chart.js [14], [15].": "Dashboard diimplementasikan menggunakan Flask, Bootstrap, dan Chart.js [14], [15].",
    "The project package is organized into fifteen main folders following the UAS AI artifact structure.": "Paket proyek disusun ke dalam lima belas folder utama sesuai struktur artefak UAS AI.",
    "Future research should extend the dataset beyond one DDoS subset.": "Penelitian selanjutnya perlu memperluas dataset melampaui satu subset DDoS.",
}


def set_props_and_style(doc: Document) -> None:
    props = doc.core_properties
    props.author = STUDENT
    props.last_modified_by = STUDENT
    props.comments = ""
    props.subject = ""
    props.keywords = ""
    for style_name in ["Normal", "Heading 1", "Heading 2", "Heading 3", "Title", "Subtitle"]:
        if style_name in doc.styles:
            style = doc.styles[style_name]
            style.font.name = FONT
            style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
            style.font.color.rgb = RGBColor(0, 0, 0)
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            run.font.name = FONT
            run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
            run.font.color.rgb = RGBColor(0, 0, 0)
            if run.font.size is None:
                run.font.size = Pt(9)


def replace_paragraph_text(paragraph, text: str) -> None:
    paragraph.clear()
    run = paragraph.add_run(text)
    run.font.name = FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    run.font.color.rgb = RGBColor(0, 0, 0)
    run.font.size = Pt(8 if text.startswith("[") else 9)
    if re.match(r"^[IVX]+\\.", text) or text in {"Referensi", "References"}:
        run.bold = True


def create_indonesian() -> None:
    doc = Document(SOURCE)
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text in PARAGRAPH_MAP:
            replace_paragraph_text(paragraph, PARAGRAPH_MAP[text])
            continue
        new_text = paragraph.text
        for old, new in PHRASE_MAP.items():
            new_text = new_text.replace(old, new)
        if new_text != paragraph.text:
            replace_paragraph_text(paragraph, new_text)
    set_props_and_style(doc)
    doc.core_properties.title = "Draft Artikel IEEE NetGuard AI Bahasa Indonesia"
    doc.save(ID_DOCX)


def export_pdf(docx: Path, pdf: Path) -> int | None:
    ps = f"""
$ErrorActionPreference = 'Stop'
$word = New-Object -ComObject Word.Application
$word.Visible = $false
$doc = $word.Documents.Open('{docx}')
$doc.Repaginate()
$pages = $doc.ComputeStatistics(2)
$doc.SaveAs([ref]'{pdf}', [ref]17)
$doc.Close($false)
$word.Quit()
Write-Output $pages
"""
    result = subprocess.run(["powershell", "-NoProfile", "-Command", ps], capture_output=True, text=True, timeout=120, check=True)
    try:
        return int(result.stdout.strip().splitlines()[-1])
    except Exception:
        return None


def clear_readonly(path: Path) -> None:
    if path.exists():
        subprocess.run(["attrib", "-R", str(path)], check=False)


def main() -> None:
    shutil.copy2(SOURCE, EN_DOCX)
    en = Document(EN_DOCX)
    set_props_and_style(en)
    en.core_properties.title = "Draft Artikel IEEE NetGuard AI English Version"
    en.save(EN_DOCX)
    create_indonesian()
    pages_en = export_pdf(EN_DOCX, EN_PDF)
    pages_id = export_pdf(ID_DOCX, ID_PDF)
    for p in [SOURCE, EN_DOCX, EN_PDF, ID_DOCX, ID_PDF]:
        clear_readonly(p)
    print(f"English pages: {pages_en}")
    print(f"Indonesian pages: {pages_id}")


if __name__ == "__main__":
    main()
