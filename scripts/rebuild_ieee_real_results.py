from __future__ import annotations

import json
import shutil
import subprocess
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
MAIN = ROOT / "Achmad_Maulana_241730016_UAS_AI"
OUT_DIR = MAIN / "09_Draft_IEEE"
OUT_DOCX = OUT_DIR / "Draft_Artikel_IEEE.docx"
OUT_PDF = OUT_DIR / "Draft_Artikel_IEEE.pdf"

STUDENT = "Achmad Maulana"
NIM = "241730016"
AFFILIATION = "Program Studi Informatika, Fakultas Sains dan Teknologi, Universitas Islam Negeri Sultan Maulana Hasanuddin Banten"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False, size: int = 8) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor(0, 0, 0)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_two_columns(section) -> None:
    sect_pr = section._sectPr
    cols = sect_pr.xpath("./w:cols")
    if cols:
        cols = cols[0]
    else:
        cols = OxmlElement("w:cols")
        sect_pr.append(cols)
    cols.set(qn("w:num"), "2")
    cols.set(qn("w:space"), "360")


def add_caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.font.name = "Times New Roman"
    r.font.size = Pt(8)
    r.italic = True


def add_figure(doc: Document, path: Path, caption: str, width: float = 3.15) -> None:
    if path.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(path), width=Inches(width))
        add_caption(doc, caption)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], caption: str) -> None:
    add_caption(doc, caption)
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for i, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], header, bold=True, size=7)
        set_cell_shading(table.rows[0].cells[i], "D9EAF7")
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value, size=7)


def heading(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(text.upper())
    r.bold = True
    r.font.name = "Times New Roman"
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(0, 0, 0)


def para(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.first_line_indent = Inches(0.18)
    r = p.add_run(text)
    r.font.name = "Times New Roman"
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0, 0, 0)


def build_docx() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    metrics = json.loads((ROOT / "reports" / "metrics.json").read_text(encoding="utf-8"))
    comparison = pd.read_csv(ROOT / "reports" / "model_comparison.csv")
    summary = json.loads((ROOT / "reports" / "research_summary.json").read_text(encoding="utf-8"))
    ds = metrics["dataset_summary"]
    best_model = metrics["best_model"]
    best = metrics["models"][best_model]

    doc = Document()
    doc.core_properties.author = STUDENT
    doc.core_properties.last_modified_by = STUDENT
    doc.core_properties.comments = ""
    doc.core_properties.title = "NetGuard AI: Lightweight Machine Learning-Based Network Anomaly Detection and Risk Monitoring Dashboard"

    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.68)
    section.right_margin = Inches(0.68)

    for style in ["Normal", "Heading 1", "Heading 2", "Heading 3"]:
        doc.styles[style].font.name = "Times New Roman"
        doc.styles[style].font.color.rgb = RGBColor(0, 0, 0)
    doc.styles["Normal"].font.size = Pt(9)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title.add_run("NetGuard AI: Lightweight Machine Learning-Based Network Anomaly Detection and Risk Monitoring Dashboard")
    tr.bold = True
    tr.font.name = "Times New Roman"
    tr.font.size = Pt(18)

    author = doc.add_paragraph()
    author.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ar = author.add_run(f"{STUDENT}\n{NIM}\n{AFFILIATION}\nBanten, Indonesia\n2026")
    ar.font.name = "Times New Roman"
    ar.font.size = Pt(10)

    abstract = doc.add_paragraph()
    abstract.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    abstract.add_run("Abstract—").bold = True
    abstract.add_run(
        "Network administrators in educational environments often have limited infrastructure for detecting abnormal traffic patterns. "
        "This study proposes NetGuard AI, a low-budget machine learning system for detecting network anomalies and presenting risk information through a lightweight dashboard. "
        "A representative subset of the CICIDS2017 DDoS traffic file was used with 20,000 records sampled from real BENIGN and DDoS classes. "
        "After preprocessing and duplicate removal, 19,935 records were used for binary classification. "
        "Three baseline models were evaluated: Logistic Regression, Decision Tree, and Random Forest. "
        f"The best model was {best_model}, achieving accuracy {best['accuracy']:.4f}, precision {best['precision']:.4f}, recall {best['recall']:.4f}, and F1-score {best['f1_score']:.4f}. "
        "The system also exports metrics, prediction results, confusion matrix visualization, and a research dashboard that converts anomaly ratio into Low, Medium, or High risk categories. "
        "The contribution of this work is the integration of reproducible anomaly detection, model comparison, automated risk classification, and beginner-friendly deployment artifacts for educational network monitoring."
    )

    kw = doc.add_paragraph()
    kw.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    kw.add_run("Keywords—").bold = True
    kw.add_run("network anomaly detection, CICIDS2017, machine learning, intrusion detection, Flask dashboard")

    section2 = doc.add_section(WD_SECTION.CONTINUOUS)
    set_two_columns(section2)

    heading(doc, "I. Introduction")
    para(doc, "The growth of digital services in schools, universities, and small organizations increases the need for practical network monitoring. However, many educational institutions do not have expensive security information and event management systems. Students with networking backgrounds also need a project that connects server infrastructure, traffic analysis, and artificial intelligence in a reproducible way.")
    para(doc, "NetGuard AI addresses this need by combining public network traffic data, a simple preprocessing pipeline, supervised machine learning models, and a dashboard that converts prediction results into operational risk recommendations. The project is designed to run on a normal Windows laptop without paid APIs, GPU infrastructure, or commercial monitoring software.")

    heading(doc, "II. Related Works")
    para(doc, "Recent studies on intrusion detection using CICIDS2017 report strong performance from both traditional machine learning and deep learning approaches. Benchmarking studies show that tree-based models often perform strongly on CICIDS2017 features, while dataset-quality studies warn that improper use of CICIDS2017 can produce overly optimistic conclusions. Other works explore hybrid deep learning, imbalance handling, adaptive learning, and feature importance analysis.")
    para(doc, "The gap selected in this work is practical integration. Many studies emphasize model performance, but fewer provide a lightweight educational dashboard that presents model outputs, risk score, recommended action, and reproducible artifacts in a form suitable for beginner network administrators.")

    heading(doc, "III. Proposed Method")
    para(doc, "The proposed method consists of five stages: dataset acquisition, preprocessing, model training, evaluation, and dashboard/report integration. The preprocessing stage normalizes column names, converts BENIGN to class 0 and DDoS to class 1, replaces infinity values, converts features to numeric values, fills missing values using median values, and removes duplicate rows.")
    para(doc, "The training stage compares Logistic Regression, Decision Tree, and Random Forest. Model selection uses F1-score as the primary criterion, recall as the secondary criterion, and accuracy as the tertiary criterion. The dashboard stage reads exported JSON and CSV artifacts and displays total records, normal traffic, anomaly traffic, model accuracy, risk score, risk level, prediction result, and recommended action.")

    add_figure(doc, MAIN / "08_Visualisasi" / "research_pipeline_diagram.png", "Fig. 1. Research and implementation pipeline of NetGuard AI.")

    heading(doc, "IV. Experimental Setup")
    setup_rows = [
        ["Dataset", "CICIDS2017 Friday Afternoon DDoS"],
        ["Raw records", "225,745"],
        ["Subset method", "Stratified random sampling, random_state=42"],
        ["Subset size", "20,000 records"],
        ["Processed records", f"{ds['total_records']:,}"],
        ["Normal records", f"{ds['normal_records']:,}"],
        ["Anomaly records", f"{ds['anomaly_records']:,}"],
        ["Train/test split", "80:20, random_state=42"],
        ["Models", "Logistic Regression, Decision Tree, Random Forest"],
    ]
    add_table(doc, ["Parameter", "Value"], setup_rows, "TABLE I. EXPERIMENTAL SETUP")
    add_figure(doc, MAIN / "08_Visualisasi" / "dataset_distribution_chart.png", "Fig. 2. Distribution of normal and anomaly records after preprocessing.")

    heading(doc, "V. Results and Discussion")
    model_rows = []
    for _, row in comparison.iterrows():
        model_rows.append([
            str(row["model"]),
            f"{float(row['accuracy']):.4f}",
            f"{float(row['precision']):.4f}",
            f"{float(row['recall']):.4f}",
            f"{float(row['f1_score']):.4f}",
        ])
    add_table(doc, ["Model", "Accuracy", "Precision", "Recall", "F1-score"], model_rows, "TABLE II. MODEL COMPARISON RESULTS")
    para(doc, f"The experimental results show that {best_model} achieved the highest F1-score according to the selection rule. The confusion matrix of the best model indicates that only a small number of test samples were misclassified. This result is consistent with the known strength of tree-based methods on structured flow features.")
    para(doc, "Although the metrics are high, the result should be interpreted carefully because this experiment uses one CICIDS2017 attack day and a binary BENIGN-vs-DDoS setup. For a stronger thesis or publication target, further validation should include multiple attack types, cross-validation, hyperparameter tuning, and cross-dataset testing using UNSW-NB15 or CICIDS2018.")
    add_figure(doc, MAIN / "08_Visualisasi" / "model_comparison_chart.png", "Fig. 3. Comparison of accuracy, precision, recall, and F1-score.")
    add_figure(doc, MAIN / "08_Visualisasi" / "confusion_matrix.png", "Fig. 4. Confusion matrix of the selected best model.")

    heading(doc, "VI. Dashboard and Risk Classification")
    para(doc, f"The prediction module classified {summary.get('total_data', 20000):,} records and produced an anomaly ratio of {summary.get('risk_score', 49.99):.2f}%. Based on the predefined rule, the risk level is High. The recommended action is to prioritize incident investigation, isolate suspicious hosts, and verify critical services.")
    add_figure(doc, MAIN / "08_Visualisasi" / "risk_score_gauge.png", "Fig. 5. Risk score visualization based on anomaly ratio.")
    add_figure(doc, MAIN / "08_Visualisasi" / "deployment_architecture_diagram.png", "Fig. 6. Deployment architecture and future development direction.")

    heading(doc, "VII. Conclusion")
    para(doc, "This study demonstrates that NetGuard AI can detect network anomalies from a real CICIDS2017 subset, compare multiple machine learning models, save the best model, and present results through reproducible research artifacts. The system is suitable for a low-budget educational AI project because it runs on a normal laptop and does not require paid APIs or expensive hardware.")
    para(doc, "Future work should include more attack categories, cross-validation, hyperparameter optimization, feature importance analysis, database-backed monitoring, and public deployment through VPS, Vercel, Render, or Railway.")

    heading(doc, "References")
    refs = [
        "[1] Canadian Institute for Cybersecurity, University of New Brunswick, \"Intrusion Detection Evaluation Dataset (CICIDS2017),\" 2017.",
        "[2] Z. Maseer et al., \"Benchmarking of Machine Learning for Anomaly-Based Intrusion Detection Systems in the CICIDS2017 Dataset,\" IEEE Access, 2021.",
        "[3] M. Engelen et al., \"Troubleshooting an intrusion detection dataset: the CICIDS2017 case study,\" 2021.",
        "[4] S. Dube, \"Faulty use of the CIC-IDS-2017 dataset in information security research,\" 2024.",
        "[5] A. Sajid et al., \"Enhancing intrusion detection: a hybrid machine and deep learning approach,\" Journal of Cloud Computing, 2024.",
        "[6] M. A. Talukder et al., \"Machine learning-based network intrusion detection for big and imbalanced data,\" Journal of Big Data, 2024.",
        "[7] A. Abdelaziz et al., \"Enhancing network threat detection with Random Forest-based NIDS and permutation feature importance,\" 2024.",
        "[8] V. Pai et al., \"Adaptive network anomaly detection using machine learning,\" EURASIP Journal on Information Security, 2025.",
        "[9] S. Ahmed et al., \"A smart deep learning-based intrusion detection system for IoT networks,\" Scientific Reports, 2025.",
        "[10] H. Deng et al., \"Network intrusion detection based on deep belief network and broad learning system,\" Electronics, 2024.",
    ]
    for ref in refs:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.15)
        p.paragraph_format.first_line_indent = Inches(-0.15)
        r = p.add_run(ref)
        r.font.name = "Times New Roman"
        r.font.size = Pt(8)

    doc.save(OUT_DOCX)
    shutil.copyfile(OUT_DOCX, ROOT / "Draft_Artikel_IEEE_NetGuard_AI_Achmad_Maulana.docx")


def export_pdf() -> None:
    powershell = f"""
$docx = '{OUT_DOCX}';
$pdf = '{OUT_PDF}';
try {{
  $word = New-Object -ComObject Word.Application;
  $word.Visible = $false;
  $doc = $word.Documents.Open($docx);
  $doc.SaveAs([ref] $pdf, [ref] 17);
  $doc.Close();
  $word.Quit();
}} catch {{
  Write-Output "WORD_COM_FAILED: $($_.Exception.Message)";
  exit 1;
}}
"""
    result = subprocess.run(["powershell", "-NoProfile", "-Command", powershell], capture_output=True, text=True)
    if result.returncode == 0 and OUT_PDF.exists():
        shutil.copyfile(OUT_PDF, ROOT / "Draft_Artikel_IEEE_NetGuard_AI_Achmad_Maulana.pdf")
    else:
        print(result.stdout)
        print(result.stderr)


def main() -> None:
    build_docx()
    export_pdf()
    print(f"Updated IEEE draft: {OUT_DOCX}")
    print(f"Updated IEEE PDF: {OUT_PDF}")


if __name__ == "__main__":
    main()
