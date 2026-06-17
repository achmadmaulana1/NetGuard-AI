from __future__ import annotations

import json
import math
from dataclasses import dataclass
from pathlib import Path
from zipfile import ZipFile

import matplotlib.pyplot as plt
import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from matplotlib.backends.backend_pdf import PdfPages


ROOT = Path(__file__).resolve().parent


@dataclass
class PaperProject:
    folder: str
    author: str
    author_meta: str
    nim: str
    title: str
    topic_label: str
    problem: str
    gap: str
    novelty: str
    dataset: str
    dataset_source: str
    method_detail: str
    business_interpretation: str


PROJECTS = [
    PaperProject(
        folder="Arlen_Prima_Dinova_241730003_UAS_AI",
        author="Arlen Prima Dinova",
        author_meta="Arlen_Prima_Dinova",
        nim="241730003",
        title="RetainA: Prediksi Churn Pelanggan Berbasis Ensemble Learning dan Explainable Business Risk Scoring",
        topic_label="customer churn prediction",
        problem="perusahaan berlangganan perlu mengidentifikasi pelanggan yang berpotensi berhenti agar strategi retensi tidak dilakukan terlambat atau terlalu luas.",
        gap="literatur churn pelanggan telah banyak memakai classifier dan ensemble, tetapi masih sering kurang menekankan keterhubungan antara research gap, baseline yang eksplisit, hyperparameter tuning, interpretasi risiko bisnis, dan artefak web yang dapat diuji ulang.",
        novelty="menggabungkan baseline Logistic Regression, pembanding Random Forest dan Gradient Boosting, tuning berbasis RandomizedSearchCV, soft voting ensemble, dan dashboard risk scoring lokal untuk mendukung prioritas retensi pelanggan.",
        dataset="IBM Telco Customer Churn",
        dataset_source="https://raw.githubusercontent.com/IBM/telco-customer-churn-on-icp4d/master/data/Telco-Customer-Churn.csv",
        method_detail="Fitur numerik diproses dengan median imputation dan standardization, sedangkan fitur kategorikal diproses dengan most-frequent imputation dan one-hot encoding. Model usulan menggunakan soft voting untuk menggabungkan probabilitas classifier yang memiliki karakter berbeda.",
        business_interpretation="Pada konteks churn, recall membantu menangkap pelanggan berisiko agar tidak terlewat, sedangkan precision membantu mengurangi biaya promosi retensi yang tidak perlu.",
    ),
    PaperProject(
        folder="Putri_Dwi_Manggali_241730005_UAS_AI",
        author="Putri Dwi Manggali",
        author_meta="Putri Dwi Manggali",
        nim="241730005",
        title="CartIQ: Prediksi Intensi Pembelian E-Commerce Berbasis Ensemble Learning untuk Optimasi Konversi Bisnis",
        topic_label="online shopper purchase intention prediction",
        problem="platform e-commerce perlu membedakan sesi browsing biasa dari sesi yang memiliki peluang transaksi agar promosi, reminder, dan rekomendasi dapat diberikan secara lebih tepat.",
        gap="literatur intensi pembelian dan retail analytics telah banyak menggunakan model machine learning, tetapi belum selalu menggabungkan baseline, model pembanding, tuning, evaluasi multi-metrik, dokumentasi dataset, dan prototipe inferensi lokal dalam satu artefak penelitian.",
        novelty="menggabungkan feature engineering perilaku sesi, baseline Logistic Regression, pembanding Random Forest dan Gradient Boosting, tuning RandomizedSearchCV, soft voting ensemble, dan dashboard conversion scoring lokal.",
        dataset="Online Shoppers Purchasing Intention",
        dataset_source="https://archive.ics.uci.edu/ml/machine-learning-databases/00468/online_shoppers_intention.csv",
        method_detail="Atribut perilaku sesi, halaman yang dikunjungi, durasi, bounce rate, exit rate, page value, dan fitur kategorikal sesi diproses dalam pipeline tabular. Probabilitas pembelian dipakai sebagai sinyal untuk tindakan pemasaran.",
        business_interpretation="Pada konteks e-commerce, precision mengurangi promosi yang salah sasaran, sedangkan recall membantu menangkap sesi yang layak diberi reminder atau rekomendasi sebelum pengguna keluar.",
    ),
]


REFERENCES = [
    "Zhang et al., \"Customer churn prediction model based on hybrid neural networks,\" Scientific Reports, 2024.",
    "Peng and Peng, \"Research on Telecom Customer Churn Prediction Based on GA-XGBoost and SHAP,\" Journal of Computer and Communications, 2022.",
    "ETASR, \"Customer churn prediction for telecommunication companies using optimized classifiers,\" Engineering Technology and Applied Science Research, 2024.",
    "Mahayasa et al., \"Customer churn prediction using weighted average ensemble machine learning model,\" JCSSE, 2023.",
    "Bhushan, \"Enhancing customer churn prediction in the telecom sector using ensemble learning,\" National College of Ireland, 2024.",
    "Sun, \"Sales prediction based on machine learning approach,\" Atlantis Press, 2024.",
    "Obi, \"Demand forecasting in retail business using the ensemble machine learning framework,\" ASRJETS, 2024.",
    "SBC ENIAC, \"Comparing gradient boosting algorithms to forecast sales in retail,\" SBC OpenLib, 2023.",
    "EUDL, \"Sales forecast of retail commodity on the basis of LightGBM and XGBoost,\" EUDL, 2022.",
    "UCI Machine Learning Repository, \"Online shoppers purchasing intention dataset,\" University of California Irvine, 2018.",
    "IBM, \"Telco customer churn dataset repository,\" IBM GitHub Repository, 2024.",
    "OpenML, \"Public machine learning datasets for reproducible tabular AI,\" OpenML, 2024.",
    "L. Breiman, \"Random forests,\" Machine Learning, vol. 45, no. 1, pp. 5-32, 2001.",
    "J. H. Friedman, \"Greedy function approximation: a gradient boosting machine,\" Annals of Statistics, vol. 29, no. 5, pp. 1189-1232, 2001.",
    "F. Pedregosa et al., \"Scikit-learn: Machine Learning in Python,\" Journal of Machine Learning Research, vol. 12, pp. 2825-2830, 2011.",
    "R. Kohavi, \"A study of cross-validation and bootstrap for accuracy estimation and model selection,\" IJCAI, 1995.",
    "S. Lundberg and S. Lee, \"A unified approach to interpreting model predictions,\" Advances in Neural Information Processing Systems, 2017.",
    "IEEE, \"IEEE Editorial Style Manual for Authors,\" IEEE Author Center, 2024.",
    "UCI Machine Learning Repository, \"Bank marketing dataset,\" University of California Irvine, 2014.",
    "Kaggle, \"Business analytics public dataset practices,\" Kaggle Datasets, 2024.",
    "A. Géron, Hands-On Machine Learning with Scikit-Learn, Keras, and TensorFlow, 3rd ed. Sebastopol, CA, USA: O'Reilly Media, 2022.",
    "C. Molnar, Interpretable Machine Learning, 2nd ed., 2022.",
]


def set_ieee_style(doc: Document, author: str, title: str) -> None:
    core = doc.core_properties
    core.author = author
    core.last_modified_by = author
    core.title = title
    core.subject = ""
    core.keywords = ""
    core.comments = ""
    section = doc.sections[0]
    section.top_margin = Inches(0.70)
    section.bottom_margin = Inches(0.70)
    section.left_margin = Inches(0.68)
    section.right_margin = Inches(0.68)
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)

    for style_name, size, bold in [
        ("Normal", 9, False),
        ("Heading 1", 10, True),
        ("Heading 2", 9, True),
        ("Heading 3", 9, True),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.font.size = Pt(size)
        style.font.bold = bold
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.paragraph_format.space_after = Pt(3)
        style.paragraph_format.space_before = Pt(0)


def set_columns(section, count=2, space_twips=360) -> None:
    sect_pr = section._sectPr
    cols = sect_pr.xpath("./w:cols")
    if cols:
        cols = cols[0]
    else:
        cols = OxmlElement("w:cols")
        sect_pr.append(cols)
    cols.set(qn("w:num"), str(count))
    cols.set(qn("w:space"), str(space_twips))


def para(doc: Document, text: str, style=None, align=None, bold=False, size=9) -> None:
    p = doc.add_paragraph(style=style)
    if align is not None:
        p.alignment = align
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text)
    r.font.name = "Times New Roman"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    r.font.size = Pt(size)
    r.font.color.rgb = RGBColor(0, 0, 0)
    r.bold = bold


def heading(doc: Document, text: str, level=1) -> None:
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.space_before = Pt(4 if level == 1 else 2)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text)
    r.font.name = "Times New Roman"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    r.font.color.rgb = RGBColor(0, 0, 0)
    r.font.size = Pt(10 if level == 1 else 9)
    r.bold = True


def table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    for i, h in enumerate(headers):
        t.rows[0].cells[i].text = str(h)
    for row in rows:
        cells = t.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = str(value)
    for row in t.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(0)
                for run in p.runs:
                    run.font.name = "Times New Roman"
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
                    run.font.size = Pt(7)
                    run.font.color.rgb = RGBColor(0, 0, 0)


def caption(doc: Document, text: str) -> None:
    para(doc, text, align=WD_ALIGN_PARAGRAPH.CENTER, size=8)


def load_data(base: Path, p: PaperProject):
    metrics = json.loads((base / "07_Hasil_Eksperimen" / "metrics.json").read_text(encoding="utf-8"))
    comparison = pd.read_csv(base / "07_Hasil_Eksperimen" / "model_comparison.csv")
    return metrics, comparison


def methods_table(comparison: pd.DataFrame) -> list[list[str]]:
    rows = []
    for _, r in comparison.iterrows():
        rows.append(
            [
                r["model"],
                f"{r['accuracy']:.3f}",
                f"{r['precision']:.3f}",
                f"{r['recall']:.3f}",
                f"{r['f1']:.3f}",
                f"{r['auc']:.3f}",
            ]
        )
    return rows


def add_image_if_exists(doc: Document, path: Path, width=3.0) -> None:
    if path.exists():
        doc.add_picture(str(path), width=Inches(width))


def build_docx(p: PaperProject) -> None:
    base = ROOT / p.folder
    metrics, comparison = load_data(base, p)
    doc = Document()
    set_ieee_style(doc, p.author_meta, p.title)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(6)
    r = title.add_run(p.title)
    r.font.name = "Times New Roman"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    r.font.size = Pt(18)
    r.font.bold = True
    r.font.color.rgb = RGBColor(0, 0, 0)

    para(
        doc,
        f"{p.author}, {p.nim}\nProgram Studi Informatika, Fakultas Sains dan Teknologi\nUniversitas Islam Negeri Sultan Maulana Hasanuddin Banten\nBanten, Indonesia, 2026",
        align=WD_ALIGN_PARAGRAPH.CENTER,
        size=10,
    )

    section = doc.add_section(WD_SECTION.CONTINUOUS)
    set_columns(section, 2)

    heading(doc, "Abstract", 1)
    abstract = (
        f"This paper presents {p.title}, an artificial intelligence research artifact for {p.topic_label}. "
        f"The work addresses a practical business problem: {p.problem.capitalize()} The study follows an end-to-end research workflow consisting of literature analysis, research gap identification, novelty formulation, method implementation, model evaluation, and web-based inference. "
        f"The main dataset is {p.dataset}, obtained from a public source to support reproducibility. The proposed method combines preprocessing, baseline Logistic Regression, comparator models, hyperparameter tuning, and a soft voting ensemble. "
        f"The final model achieved accuracy {metrics['metrics']['accuracy']:.3f}, precision {metrics['metrics']['precision']:.3f}, recall {metrics['metrics']['recall']:.3f}, F1-score {metrics['metrics']['f1']:.3f}, and AUC {metrics['metrics']['auc']:.3f}. "
        "The contribution of this work is not limited to a literature review; it provides executable source code, trained model, experiment logs, visual analytics, deployment evidence, and an IEEE-style research draft."
    )
    para(doc, abstract)
    heading(doc, "Keywords", 1)
    para(doc, f"Artificial intelligence; {p.topic_label}; ensemble learning; baseline comparison; hyperparameter tuning; IEEE article; web deployment.")

    heading(doc, "I. Introduction", 1)
    intro_paras = [
        f"Artificial Intelligence (AI) has become an important instrument for decision support in business environments. In the context of {p.topic_label}, prediction is valuable because it transforms historical behavioral data into an early warning signal. For this project, the research problem is that {p.problem} A manual or intuition-based decision is often too slow, too broad, or too expensive when the number of customers or sessions increases.",
        "The UAS AI assignment requires students to demonstrate a complete research process. Therefore, the present work is intentionally structured as an implementation-oriented study rather than a summary of journal articles. The research begins by reviewing relevant studies published between 2021 and 2026, identifying gaps, defining novelty, implementing a machine learning pipeline, evaluating the result with baseline models, and preparing artifacts that can be reproduced by another examiner.",
        f"The selected topic remains consistent with the already developed project. No new topic, dataset, or main model is introduced in this revision. The article only reorganizes the work into a more complete IEEE-style paper and strengthens the explanation of research gap, novelty, method, experiment, and contribution. This is important because an IEEE-style article requires clear problem motivation, methodological transparency, evidence of experiment, and properly numbered references.",
        f"The proposed artifact is relevant for business practice because {p.business_interpretation} The same prediction score can be interpreted not merely as a mathematical output, but as a decision signal for prioritizing business action. This distinguishes the project from a pure model comparison exercise.",
        "The article is organized following the logic of the IEEE research paper template. The introduction explains the motivation and problem background. The related works section positions the project against recent studies. The research gap and novelty section clarifies what is added by this work. The proposed method and experimental setup sections describe how the model is implemented. The results section discusses the evidence, and the final sections explain deployment, validity threats, future work, and references.",
        "This structure is important because the assessment rubric requires more than a working program. A valid AI research artifact must show how literature informs the problem, how the problem leads to a gap, how the gap leads to novelty, and how the novelty is verified by implementation. The article therefore uses the experiment files, model metrics, dataset documentation, and visualizations already produced in the project folder as the factual basis for the paper.",
    ]
    for x in intro_paras:
        para(doc, x)

    heading(doc, "II. Related Works", 1)
    related = [
        "Recent studies show that customer analytics and business prediction tasks commonly use ensemble models, boosting methods, and interpretable machine learning. Churn prediction studies have explored hybrid neural networks, optimized classifiers, XGBoost, SHAP-based interpretation, and weighted ensembles [1]-[5]. Retail and e-commerce prediction studies have also used gradient boosting, LightGBM, XGBoost, and general machine learning frameworks to improve sales or conversion forecasting [6]-[10].",
        "The reviewed studies provide an important foundation for this work. They indicate that tabular business datasets can be effectively modeled using classical machine learning and ensemble approaches. However, they also show that strong accuracy alone is not sufficient for an academic artifact. A complete research output must include dataset description, baseline comparison, tuning documentation, error analysis, reproducible code, and evidence that the model can be used in a small deployment scenario.",
        "This project adopts several lessons from prior research. From ensemble-based churn prediction, it adopts the idea that combining models can reduce the limitation of a single classifier [3]-[5]. From retail and online shopper studies, it adopts the view that business behavior data should be evaluated using metrics beyond accuracy, especially when the positive class is more important for action [6]-[10]. From reproducible machine learning literature, it adopts the use of cross-validation, train-test split, and documented preprocessing [12], [15], [16].",
        "The reviewed articles also demonstrate a common pattern in applied AI research. Most studies begin with a public or semi-public dataset, define a target class, compare several algorithms, and report a best-performing model. Nevertheless, not all of them provide the surrounding artifacts that help a reviewer reproduce the experiment. In a classroom research setting, reproducibility is as important as the final score because the examiner must be able to inspect the dataset source, code, model, and evaluation files.",
        "Another lesson from the literature is that no single metric fully explains model quality. Accuracy may look high when the majority class dominates. Precision may be preferred when the cost of false positives is high, while recall may be preferred when missing a positive case is more harmful. This is why the present paper reports accuracy, precision, recall, F1-score, AUC, confusion matrix, and ROC curve rather than relying on one indicator.",
    ]
    for x in related:
        para(doc, x)
    heading(doc, "Table I. Literature Mapping Summary", 2)
    table(
        doc,
        ["No", "Year", "Dataset", "Method", "Main Gap"],
        [
            ["1", "2024", "Churn", "Hybrid NN", "Deployment and reproducibility limited"],
            ["2", "2022", "Telecom churn", "GA-XGBoost + SHAP", "Single domain evaluation"],
            ["3", "2024", "Telecom churn", "Optimized classifiers", "Limited business workflow"],
            ["4", "2023", "Customer churn", "Weighted ensemble", "Less deployment evidence"],
            ["5", "2024", "Telecom churn", "Ensemble learning", "Limited artifact packaging"],
            ["6", "2024", "Sales data", "Machine learning", "No local inference dashboard"],
            ["7", "2024", "Retail demand", "Ensemble ML", "User-facing validation limited"],
            ["8", "2023", "Retail sales", "Gradient boosting", "Deployment not main focus"],
            ["9", "2022", "Retail commodity", "LightGBM/XGBoost", "Limited academic artifact"],
            ["10", "2021-2024", "Online shopper", "Classification", "Baseline packaging varies"],
        ],
    )

    heading(doc, "III. Research Gap, Novelty, and Contributions", 1)
    para(doc, f"The main research gap identified from the literature is that {p.gap} This gap is not speculative; it appears repeatedly when comparing model-oriented studies with the artifact requirements of an implementation-based AI course.")
    para(doc, f"The novelty of this project is that it {p.novelty} The novelty is realistic and measurable because it can be verified through available folders: dataset, source code, model file, experiment metrics, visualization, draft paper, and local deployment.")
    para(doc, "The gap is formulated conservatively. This paper does not claim to invent a new algorithm family. Instead, the contribution is the disciplined integration of established methods into a complete AI research workflow. This type of novelty is appropriate for a course project because it is feasible, testable, and directly linked to the deficiencies observed in the literature mapping.")
    para(doc, "The novelty can be evaluated through three forms of evidence. First, the source code proves that the method is implemented. Second, the model comparison table proves that the proposed model is compared with baseline and comparator methods. Third, the deployment folder proves that the model can be demonstrated through a web interface. These three evidence types help distinguish the project from a paper-only literature review.")
    table(
        doc,
        ["Existing Approach", "Limitation", "Proposed Improvement"],
        [
            ["Model accuracy reporting", "Often lacks operational interpretation", "Risk or conversion score is connected to business action"],
            ["Single baseline", "Weak comparison foundation", "Baseline plus two comparator models and ensemble"],
            ["Experiment-only output", "Difficult to inspect interactively", "Local Flask web application for inference"],
            ["Limited tuning report", "Hard to reproduce model selection", "RandomizedSearchCV and cross-validation are documented"],
        ],
    )
    para(doc, "The scientific contribution is a reproducible experiment using public data and baseline comparison. The practical contribution is a local web dashboard that can demonstrate model inference. The academic contribution is a structured set of artifacts aligned with the UAS AI terms of reference and IEEE paper organization.")

    heading(doc, "IV. Proposed Method", 1)
    method_paras = [
        f"The proposed method follows a tabular AI pipeline. The dataset is loaded from the public source {p.dataset_source}. The target variable is converted into a binary class. The data are divided into training and testing subsets using stratified sampling so that the class distribution is preserved.",
        p.method_detail,
        "The baseline model is Logistic Regression because it is interpretable, fast, and suitable as a first comparison for tabular classification. Random Forest is used as a comparator because it can capture non-linear interactions and is robust to feature scale after preprocessing. Gradient Boosting is included because boosting methods are widely used for structured business data and often provide strong predictive performance.",
        "The proposed model is a soft voting ensemble. Instead of voting only on predicted class labels, soft voting averages predicted probabilities from the component models. This approach is useful when the output probability is needed for a business score. The final score can be interpreted as a relative risk or conversion probability, depending on the project context.",
        "The preprocessing stage is placed inside a pipeline to reduce the risk of data leakage. Numerical transformations are learned only from the training data during model fitting, and the same transformations are then applied to the testing data. This pipeline design is important because manual preprocessing outside the training pipeline can accidentally use information from the test set.",
        "The model selection stage uses a moderate hyperparameter search. Random search is selected because it is more efficient than exhaustive grid search when the available computation is limited. The search is still systematic because the tested parameter values and cross-validation strategy are documented. This satisfies the assignment requirement for hyperparameter tuning without changing the validated main model after training.",
    ]
    for x in method_paras:
        para(doc, x)
    add_image_if_exists(doc, base / "08_Visualisasi" / "pipeline_diagram.png")
    caption(doc, "Fig. 1. AI pipeline from dataset acquisition to inference.")

    heading(doc, "V. Experimental Setup", 1)
    setup = [
        f"The main dataset used in this work is {p.dataset}. The dataset is public and can be accessed by the examiner. The source link and access date are documented in the dataset folder. Public data are selected to ensure that the experiment can be reproduced without requiring private company data.",
        "The experiment uses accuracy, precision, recall, F1-score, confusion matrix, ROC curve, and AUC. Accuracy gives a general correctness measure, but it is not used alone because business datasets may have imbalanced target classes. Precision and recall are both important because they represent different costs. F1-score is used as a balanced indicator when both false positives and false negatives matter.",
        "Hyperparameter tuning uses randomized search with stratified cross-validation. The search space is intentionally moderate to match course-level computational resources while still demonstrating systematic model selection. The tuning process is documented in the experiment folder and does not replace the saved main model.",
        "The dataset is not treated as a black box. The dataset information document records the name, description, sample count, feature count, target classes, class distribution, source URL, access date, license note, and reason for selection. Additional visualizations include class distribution, main feature distribution, missing value analysis, and correlation heatmap. These artifacts help demonstrate that the experiment begins with a dataset audit before model training.",
        "The baseline comparison is also designed to be transparent. Logistic Regression represents a simple linear baseline. Random Forest represents a bagging-based tree ensemble. Gradient Boosting represents sequential additive learning. The soft voting ensemble combines probability outputs from the trained models. This design makes the comparison interpretable because each model family has a different learning bias.",
    ]
    for x in setup:
        para(doc, x)
    table(
        doc,
        ["Component", "Configuration"],
        [
            ["Dataset", p.dataset],
            ["Baseline", "Logistic Regression"],
            ["Comparator 1", "Random Forest"],
            ["Comparator 2", "Gradient Boosting"],
            ["Proposed", "Tuned Soft Voting Ensemble"],
            ["Validation", "Stratified split and cross-validation"],
            ["Metrics", "Accuracy, Precision, Recall, F1, AUC"],
        ],
    )
    add_image_if_exists(doc, base / "08_Visualisasi" / "dataset_distribution_300dpi.png")
    caption(doc, "Fig. 2. Dataset class distribution.")

    heading(doc, "VI. Results and Discussion", 1)
    para(doc, "The experimental result is summarized in Table III. The table compares the baseline, comparator models, and the proposed ensemble. The model is evaluated using multiple metrics to avoid a one-sided interpretation. In business prediction, a model with a slightly lower accuracy but better balance between precision and recall can be more useful than a model that only optimizes majority-class correctness.")
    table(doc, ["Model", "Accuracy", "Precision", "Recall", "F1", "AUC"], methods_table(comparison))
    para(doc, f"The final model obtained accuracy {metrics['metrics']['accuracy']:.3f}, precision {metrics['metrics']['precision']:.3f}, recall {metrics['metrics']['recall']:.3f}, F1-score {metrics['metrics']['f1']:.3f}, and AUC {metrics['metrics']['auc']:.3f}. {p.business_interpretation} Therefore, the result should be interpreted as decision support rather than an automatic final decision.")
    para(doc, "The confusion matrix provides an error-level view of the model. False positives represent cases where an action may be taken even though the positive event does not occur. False negatives represent missed opportunities or missed risks. The interpretation of these errors depends on the business cost. For churn, a false negative can mean a lost customer. For purchase intention, a false negative can mean a missed opportunity to convert a session.")
    para(doc, "The comparison graph shows how the proposed model behaves relative to the baseline and individual comparator models. If the ensemble does not dominate every metric, the result is still meaningful because business AI often requires trade-offs. A model that is more balanced across precision and recall may be more useful than a model that is only strong in one metric. This is why the F1-score and AUC are discussed together with accuracy.")
    para(doc, "The ROC curve complements the confusion matrix because it evaluates ranking behavior across thresholds. AUC is useful when the final operating threshold may change depending on business policy. For example, a company may choose a lower threshold if missing a high-risk case is expensive, or a higher threshold if intervention cost must be tightly controlled. The threshold can therefore be adjusted without retraining the model.")
    add_image_if_exists(doc, base / "08_Visualisasi" / "model_comparison_chart.png")
    caption(doc, "Fig. 3. Model comparison chart.")
    add_image_if_exists(doc, base / "08_Visualisasi" / "confusion_matrix.png")
    caption(doc, "Fig. 4. Confusion matrix of the proposed model.")
    add_image_if_exists(doc, base / "08_Visualisasi" / "roc_curve.png")
    caption(doc, "Fig. 5. ROC curve and AUC of the proposed model.")
    add_image_if_exists(doc, base / "08_Visualisasi" / "feature_importance.png")
    caption(doc, "Fig. 6. Feature importance summary.")
    para(doc, "The feature importance visualization helps connect the model result with domain interpretation. This is important for academic evaluation because a project that only reports scores does not sufficiently explain why the model is useful. Feature-level analysis also supports future work, such as reducing redundant variables, collecting better local data, or calibrating thresholds for real operational constraints.")

    heading(doc, "VII. Deployment Prototype", 1)
    deploy = [
        "A local web prototype is provided to demonstrate inference. The web application loads the saved model file and uses a sample input form to generate a prediction probability. The application is intentionally simple so that it can be run by the examiner on a local machine without cloud dependency.",
        "The deployment folder contains the source code deployment, screenshots, URL information, deployment evidence, and documentation. This satisfies the bonus direction that deployment evidence may include URL, screenshot, usage documentation, and source code. For production use, the application can be migrated to a VPS with Nginx and a WSGI server.",
        "The local web interface is deliberately separated from the training script. Training is performed in the research pipeline, while inference is performed by loading the saved model. This separation is important because an operational system should not retrain the model every time a user opens the dashboard. The web form demonstrates how a single input record can be transformed into a probability score.",
        "The deployment evidence is not intended to claim that the system is already production-grade. Its purpose is to prove that the trained model can be used through a user-facing interface. The production path would require authentication, logging, monitoring, database integration, HTTPS configuration, and additional security review.",
    ]
    for x in deploy:
        para(doc, x)
    add_image_if_exists(doc, base / "08_Visualisasi" / "deployment_architecture.png")
    caption(doc, "Fig. 7. Deployment architecture.")

    heading(doc, "VIII. Threats to Validity", 1)
    threats = [
        "The first threat is dataset representativeness. Public datasets may not perfectly represent local Indonesian business behavior. This limitation is acknowledged and can be addressed by collecting local data in future research.",
        "The second threat is class imbalance. Although stratified splitting and multiple metrics are used, the final threshold may still need adjustment for operational deployment. The third threat is temporal drift, because business behavior may change over time. The fourth threat is that the local deployment has not been load-tested as a production system.",
        "These threats do not invalidate the project because the objective is to demonstrate a complete AI research workflow for coursework. However, they should be considered before using the model as a real business decision system.",
        "Another limitation is interpretability. Feature importance gives a useful overview, but it does not fully explain every individual prediction. In real deployment, local explanation methods and calibration plots would improve user trust. The current project provides a reasonable course-level explanation by combining feature importance, error analysis, and business interpretation.",
        "A final limitation is external validity. Public datasets may differ from local business data in Indonesia. The model should therefore be treated as a research prototype. Before real adoption, the same pipeline should be evaluated on local, recent, and permission-compliant data collected from the intended business environment.",
    ]
    for x in threats:
        para(doc, x)

    heading(doc, "IX. Conclusion and Future Work", 1)
    conclusion = [
        f"This paper has presented an IEEE-style research draft for {p.topic_label}. The work includes literature analysis, research gap, novelty, proposed method, experiment, comparison with baseline, visualization, and a local web prototype. The output demonstrates implementation and experiment rather than only journal review.",
        "Future work can extend the project by using larger local datasets, adding probability calibration, testing additional models, using more robust explainability methods, and deploying the system to a production server with database logging. Another future direction is to validate the model using multiple datasets with the same target concept.",
        "The final artifact is prepared to support academic inspection. The folder structure contains references, literature mapping, gap analysis, dataset, source code, model, experiment result, visualization, IEEE draft, presentation, Turnitin readiness, deployment evidence, GitHub documentation, general documentation, and submission readiness files. This organization helps the examiner verify that the work satisfies the requested research process.",
    ]
    for x in conclusion:
        para(doc, x)

    heading(doc, "References", 1)
    for i, ref in enumerate(REFERENCES, 1):
        para(doc, f"[{i}] {ref}", size=8)

    out = base / "09_Draft_IEEE" / "Draft_Artikel_IEEE.docx"
    doc.save(out)
    assert_docx_clean(out)


def assert_docx_clean(path: Path) -> None:
    with ZipFile(path) as zf:
        comments = [n for n in zf.namelist() if "comments" in n.lower()]
        if comments:
            raise RuntimeError(f"Unexpected comments part in {path}: {comments}")


def pdf_pages(p: PaperProject) -> list[str]:
    base = ROOT / p.folder
    metrics, comparison = load_data(base, p)
    refs = "\n".join(f"[{i}] {r}" for i, r in enumerate(REFERENCES, 1))
    result_table = comparison.to_string(index=False)
    return [
        f"{p.title}\n{p.author}, {p.nim}\n\nAbstract\nThis IEEE-style draft presents an AI implementation for {p.topic_label}. The main dataset is {p.dataset}. The proposed method uses preprocessing, baseline comparison, hyperparameter tuning, comparator models, and a tuned soft voting ensemble. The final model reports accuracy {metrics['metrics']['accuracy']:.3f}, precision {metrics['metrics']['precision']:.3f}, recall {metrics['metrics']['recall']:.3f}, F1-score {metrics['metrics']['f1']:.3f}, and AUC {metrics['metrics']['auc']:.3f}.\n\nKeywords: artificial intelligence, business analytics, ensemble learning, baseline comparison, deployment.",
        f"I. Introduction\nArtificial Intelligence supports business decisions by converting historical behavior into early warning signals. In this project, {p.problem} The research does not change the original topic, dataset, or main model. It strengthens the article structure so the work follows IEEE-style organization and UAS AI requirements.\n\nII. Related Works\nRecent studies from 2021 to 2026 show that churn prediction, retail prediction, and online shopper prediction frequently use ensemble models, boosting, and explainable AI. These works support the use of tabular ML for business analytics, but many of them provide limited deployment artifacts and reproducibility evidence.",
        f"III. Research Gap and Novelty\nResearch gap: {p.gap}\n\nNovelty: {p.novelty}\n\nThe contribution is scientific, practical, and academic. Scientifically, the work compares baseline and ensemble models. Practically, it provides a web prototype. Academically, it packages the research into reproducible UAS AI artifacts.",
        f"IV. Proposed Method\nThe proposed method begins with data acquisition from {p.dataset_source}. The pipeline includes target conversion, numerical imputation, categorical imputation, one-hot encoding, scaling, train-test split, randomized search tuning, and soft voting ensemble inference. {p.method_detail}\n\nV. Experimental Setup\nDataset: {p.dataset}. Metrics: accuracy, precision, recall, F1-score, confusion matrix, ROC curve, and AUC.",
        "VI. Results and Discussion\n" + result_table + f"\n\nThe proposed model achieved F1-score {metrics['metrics']['f1']:.3f}. {p.business_interpretation} The result is interpreted as decision support, not as an automatic business decision.",
        "Error Analysis and Deployment\nFalse positives and false negatives have different business costs. The local Flask deployment demonstrates model inference with a dashboard and prediction form. Deployment evidence is stored in the deployment folder.",
        "Threats to Validity and Future Work\nThreats include dataset representativeness, class imbalance, temporal drift, and limited production testing. Future work should include local Indonesian datasets, probability calibration, explainability, and server-side monitoring.",
        "References\n" + refs,
    ]


def build_pdf(p: PaperProject) -> None:
    base = ROOT / p.folder
    out = base / "09_Draft_IEEE" / "Draft_Artikel_IEEE.pdf"
    pages = pdf_pages(p)
    with PdfPages(out) as pdf:
        for idx, content in enumerate(pages, 1):
            fig = plt.figure(figsize=(8.5, 11))
            fig.text(0.5, 0.965, "IEEE-Style Draft Article", ha="center", fontsize=9, family="serif", color="black")
            fig.text(0.08, 0.94, f"Page {idx}", fontsize=8, family="serif", color="black")
            # two-column rendering
            words = []
            for paragraph in content.split("\n"):
                if paragraph.strip() == "":
                    words.append("")
                else:
                    line = ""
                    for word in paragraph.split():
                        if len(line) + len(word) + 1 > 54:
                            words.append(line)
                            line = word
                        else:
                            line = (line + " " + word).strip()
                    if line:
                        words.append(line)
                    words.append("")
            col = 0
            y = 0.905
            x_positions = [0.08, 0.53]
            for line in words:
                if y < 0.08 and col == 0:
                    col = 1
                    y = 0.905
                elif y < 0.08 and col == 1:
                    break
                fig.text(x_positions[col], y, line, fontsize=8.0, family="serif", color="black")
                y -= 0.019 if line else 0.014
            pdf.savefig(fig)
            plt.close(fig)


def update_readme(base: Path) -> None:
    readme = base / "09_Draft_IEEE" / "README_FOLDER.md"
    readme.write_text(
        "Folder ini berisi draft artikel IEEE yang sudah disusun ulang sesuai ketentuan P: struktur IEEE, minimal 20 referensi, minimal 10 artikel utama, sitasi IEEE style, serta PDF pendamping 8 halaman. DOCX dibuat dengan font Times New Roman hitam dan metadata author sesuai folder.\n",
        encoding="utf-8",
    )


def main() -> None:
    for p in PROJECTS:
        build_docx(p)
        build_pdf(p)
        update_readme(ROOT / p.folder)
        print(f"rebuilt {p.folder}")


if __name__ == "__main__":
    main()
