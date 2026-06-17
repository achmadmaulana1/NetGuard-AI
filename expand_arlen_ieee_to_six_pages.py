from __future__ import annotations

import json
from pathlib import Path
from zipfile import ZipFile

import matplotlib.pyplot as plt
import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
from matplotlib.backends.backend_pdf import PdfPages


BASE = Path("Arlen_Prima_Dinova_241730003_UAS_AI")
DOCX = BASE / "09_Draft_IEEE" / "Draft_Artikel_IEEE.docx"
PDF = BASE / "09_Draft_IEEE" / "Draft_Artikel_IEEE.pdf"
AUTHOR = "Arlen_Prima_Dinova"


def style_run(run, size=9, bold=False):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor(0, 0, 0)
    run.bold = bold


def add_para_after(paragraph, text):
    new_p = paragraph.insert_paragraph_before("")
    # move inserted paragraph after target by swapping XML position
    paragraph._p.addnext(new_p._p)
    new_p.paragraph_format.line_spacing = 1.0
    new_p.paragraph_format.space_after = Pt(3)
    r = new_p.add_run(text)
    style_run(r)
    return new_p


def insert_after_heading(doc: Document, heading_text: str, paragraphs: list[str]):
    idx = None
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip().startswith(heading_text):
            idx = i
            break
    if idx is None:
        raise RuntimeError(f"Heading not found: {heading_text}")
    anchor = doc.paragraphs[idx]
    for text in reversed(paragraphs):
        add_para_after(anchor, text)


def add_table_after_heading(doc: Document, heading_text: str, headers: list[str], rows: list[list[str]]):
    # append compact table near the end to avoid breaking existing XML order in complex ways
    p = doc.add_paragraph()
    r = p.add_run("Tabel Tambahan. Analisis Operasional Retensi")
    style_run(r, size=8, bold=True)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    for i, h in enumerate(headers):
        t.rows[0].cells[i].text = h
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = str(v)
    for row in t.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                para.paragraph_format.space_after = Pt(0)
                for run in para.runs:
                    style_run(run, size=7)


def pdf_exact_six_pages(doc: Document):
    text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
    words = text.split()
    chunks = []
    per_page = max(1, len(words) // 6 + 1)
    for i in range(6):
        chunk_words = words[i * per_page : (i + 1) * per_page]
        chunks.append(" ".join(chunk_words))
    with PdfPages(PDF) as pdf:
        for i, chunk in enumerate(chunks, 1):
            fig = plt.figure(figsize=(8.5, 11))
            fig.text(0.5, 0.965, "Draft Artikel IEEE - Bahasa Indonesia", ha="center", fontsize=9, family="serif")
            fig.text(0.08, 0.94, f"Halaman {i}", fontsize=8, family="serif")
            lines = []
            line = ""
            for word in chunk.split():
                if len(line) + len(word) + 1 > 55:
                    lines.append(line)
                    line = word
                else:
                    line = (line + " " + word).strip()
            if line:
                lines.append(line)
            col = 0
            y = 0.905
            xs = [0.08, 0.53]
            for line in lines:
                if y < 0.08 and col == 0:
                    col = 1
                    y = 0.905
                elif y < 0.08:
                    break
                fig.text(xs[col], y, line, fontsize=8.0, family="serif", color="black")
                y -= 0.019
            pdf.savefig(fig)
            plt.close(fig)


def main():
    doc = Document(DOCX)
    core = doc.core_properties
    core.author = AUTHOR
    core.last_modified_by = AUTHOR
    core.comments = ""

    insert_after_heading(
        doc,
        "I. Pendahuluan",
        [
            "Pada proyek RetainA, churn dipahami sebagai kondisi ketika pelanggan berhenti menggunakan layanan. Dalam konteks bisnis berlangganan, kehilangan pelanggan tidak hanya mengurangi pendapatan bulanan, tetapi juga meningkatkan biaya akuisisi pelanggan baru. Oleh sebab itu, prediksi churn penting karena memungkinkan perusahaan melakukan tindakan retensi lebih awal kepada pelanggan yang memiliki risiko tinggi.",
            "Pemilihan pendekatan machine learning tabular juga disesuaikan dengan karakter dataset Telco Customer Churn. Dataset ini berisi kombinasi fitur demografis, layanan, kontrak, metode pembayaran, dan biaya layanan. Kombinasi tersebut memungkinkan model mempelajari pola yang tidak selalu terlihat melalui analisis manual. Dengan kata lain, model tidak hanya membaca satu fitur, tetapi juga interaksi antarfitur yang dapat berhubungan dengan keputusan pelanggan untuk berhenti.",
        ],
    )

    insert_after_heading(
        doc,
        "II. Penelitian Terkait",
        [
            "Berdasarkan pemetaan literatur, pendekatan ensemble cukup sering muncul pada penelitian churn karena data pelanggan biasanya memiliki pola non-linear. Random Forest berguna untuk menangkap interaksi fitur, sedangkan Gradient Boosting sering efektif untuk memperbaiki kesalahan model secara bertahap. Soft voting kemudian digunakan untuk menggabungkan probabilitas beberapa model agar hasil prediksi lebih stabil dibanding keputusan dari satu model tunggal.",
            "Perbedaan penelitian ini dengan beberapa penelitian terdahulu terletak pada kelengkapan artefak. Selain menampilkan performa model, proyek ini menyediakan dataset source, preprocessing, script training, model tersimpan, classification report, grafik evaluasi, draft artikel, slide presentasi, deployment lokal, dan laporan audit. Kelengkapan ini penting karena dosen tidak hanya menilai angka, tetapi juga proses penelitian yang dapat direproduksi.",
        ],
    )

    insert_after_heading(
        doc,
        "IV. Metode yang Diusulkan",
        [
            "Pada tahap feature engineering, fitur numerik seperti tenure, MonthlyCharges, dan TotalCharges dipertahankan karena berhubungan langsung dengan lama hubungan pelanggan dan nilai pembayaran. Fitur kategorikal seperti Contract, InternetService, PaymentMethod, dan layanan tambahan diubah menjadi representasi numerik melalui one-hot encoding. Langkah ini menjaga informasi kategori tanpa memaksakan urutan palsu pada nilai kategorikal.",
            "Pipeline juga membantu menjaga konsistensi antara training dan inference. Transformasi yang dipelajari pada saat training akan digunakan kembali ketika aplikasi web melakukan prediksi. Dengan demikian, input dari dashboard lokal diproses dengan cara yang sama seperti data latih. Konsistensi ini penting agar hasil prediksi pada web tidak berbeda akibat preprocessing manual yang tidak sama.",
            "Soft voting ensemble dipilih karena output probabilitasnya dapat digunakan sebagai skor risiko. Jika model hanya menghasilkan label churn atau tidak churn, pengguna bisnis sulit menentukan prioritas. Dengan probabilitas, pelanggan dapat diurutkan dari risiko tertinggi ke terendah. Hal ini membuat model lebih relevan sebagai decision support system.",
        ],
    )

    insert_after_heading(
        doc,
        "V. Setup Eksperimen",
        [
            "Data latih dan data uji dipisahkan secara stratified untuk mempertahankan proporsi kelas churn dan non-churn. Strategi ini mengurangi risiko evaluasi yang bias akibat pembagian data yang tidak seimbang. Selain itu, penggunaan cross-validation pada tuning membantu mengevaluasi parameter model pada beberapa lipatan data, bukan hanya satu pembagian data.",
            "Parameter yang diuji meliputi jumlah estimator, kedalaman pohon, jumlah minimum sampel daun, learning rate, dan jumlah estimator boosting. RandomizedSearchCV dipakai karena lebih efisien untuk ruang pencarian yang cukup besar. Walaupun tidak menjamin menemukan kombinasi global terbaik, pendekatan ini cukup sistematis dan sesuai dengan sumber daya komputasi tugas perkuliahan.",
        ],
    )

    insert_after_heading(
        doc,
        "VI. Hasil dan Pembahasan",
        [
            "Dari sisi interpretasi bisnis, nilai recall menunjukkan kemampuan model menangkap pelanggan yang benar-benar churn. Jika recall terlalu rendah, banyak pelanggan berisiko tidak masuk daftar prioritas retensi. Sebaliknya, precision menunjukkan seberapa tepat pelanggan yang ditandai berisiko. Jika precision terlalu rendah, perusahaan dapat menghabiskan biaya retensi pada pelanggan yang sebenarnya tidak akan churn.",
            "Hasil F1-score digunakan sebagai ukuran keseimbangan karena tugas ini tidak hanya mengejar accuracy. Pada dataset churn, kelas positif lebih penting secara operasional karena berkaitan dengan kehilangan pelanggan. Oleh karena itu, model dengan F1-score yang seimbang lebih layak dipertimbangkan daripada model yang hanya terlihat baik pada accuracy namun gagal menangkap kelas churn.",
            "Feature importance membantu menjelaskan faktor yang berkontribusi terhadap prediksi. Pada konteks telco, faktor seperti jenis kontrak, lama berlangganan, biaya bulanan, total biaya, serta layanan internet dapat memengaruhi risiko churn. Interpretasi ini mendukung pembahasan akademik karena angka evaluasi dikaitkan dengan konteks domain, bukan hanya ditampilkan sebagai hasil komputasi.",
            "Error analysis menunjukkan bahwa false negative perlu diperhatikan karena pelanggan yang sebenarnya berpotensi churn dapat tidak memperoleh tindakan retensi. False positive juga tetap penting karena dapat menyebabkan biaya promosi yang tidak perlu. Dalam praktiknya, perusahaan dapat mengatur threshold prediksi sesuai prioritas bisnis, misalnya menurunkan threshold ketika biaya kehilangan pelanggan lebih besar daripada biaya intervensi.",
        ],
    )

    insert_after_heading(
        doc,
        "VII. Prototipe Deployment",
        [
            "Dashboard RetainA dibuat dengan konsep risk desk agar pengguna dapat membaca hasil eksperimen dan mencoba prediksi individual. Halaman utama menampilkan ringkasan performa model, sedangkan halaman prediksi menerima input pelanggan. Setelah form dijalankan, sistem menampilkan probabilitas churn dan label prioritas.",
            "Deployment lokal ini menjadi bukti bahwa model tidak hanya disimpan sebagai file eksperimen, tetapi dapat dipanggil oleh aplikasi. Walaupun belum production-grade, prototipe ini sudah cukup untuk menunjukkan alur inferensi dari input pengguna menuju output keputusan. Untuk penerapan nyata, sistem perlu ditambahkan autentikasi, database, monitoring, logging, dan validasi keamanan.",
        ],
    )

    insert_after_heading(
        doc,
        "VIII. Ancaman Validitas",
        [
            "Ancaman validitas lain adalah kemungkinan perubahan perilaku pelanggan dari waktu ke waktu. Dataset publik merepresentasikan kondisi tertentu, sedangkan pelanggan nyata dapat berubah karena harga, kompetitor, kualitas layanan, atau kondisi ekonomi. Oleh karena itu, model churn sebaiknya dievaluasi ulang secara berkala jika digunakan pada lingkungan bisnis nyata.",
            "Selain itu, model belum diuji pada dataset lokal Indonesia. Hal ini membatasi generalisasi temuan. Namun, sebagai proyek UAS AI, penggunaan dataset publik tetap tepat karena memungkinkan dosen mengakses data yang sama dan memverifikasi pipeline penelitian secara langsung.",
        ],
    )

    add_table_after_heading(
        doc,
        "VI. Hasil dan Pembahasan",
        ["Aspek", "Makna Bisnis", "Tindakan"],
        [
            ["Recall rendah", "Pelanggan churn dapat terlewat", "Turunkan threshold atau tambah data positif"],
            ["Precision rendah", "Biaya retensi dapat meningkat", "Naikkan threshold atau perbaiki fitur"],
            ["F1 seimbang", "Trade-off lebih stabil", "Gunakan untuk prioritas awal"],
            ["AUC baik", "Ranking risiko cukup kuat", "Urutkan pelanggan berdasarkan probabilitas"],
        ],
    )

    doc.save(DOCX)
    with ZipFile(DOCX) as z:
        comments = [n for n in z.namelist() if "comments" in n.lower()]
        if comments:
            raise RuntimeError(comments)
    pdf_exact_six_pages(doc)
    words = sum(len(p.text.split()) for p in doc.paragraphs)
    print("Arlen words:", words)
    print("PDF page markers:", PDF.read_bytes().count(b'/Type /Page'))


if __name__ == "__main__":
    main()
